#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
EPL 데이터 프로젝트 - 전처리 및 탐색적 데이터 분석(EDA) 종합 보고서 생성기 v2
"""

import json
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── 파일 경로 설정 ───────────────────────────────────────────────────────────
BASE = "C:/Users/xcv54/workspace/EPL project"
JSON_MATCH   = f"{BASE}/reports/analysis_match_results.json"
JSON_SQUAD   = f"{BASE}/reports/analysis_fbref_squad.json"
JSON_LOGS    = f"{BASE}/reports/analysis_fbref_matchlogs.json"
JSON_TM      = f"{BASE}/reports/analysis_transfermarkt.json"
JSON_PROC    = f"{BASE}/reports/analysis_processed.json"
OUT_PATH     = f"{BASE}/reports/전처리_EDA_보고서.docx"

FONT_NAME = "맑은 고딕"

# ─── JSON 로드 ────────────────────────────────────────────────────────────────
def load(path):
    with open(path, encoding="utf-8") as f:
        return json.load(f)

mr   = load(JSON_MATCH)
sq   = load(JSON_SQUAD)
ml   = load(JSON_LOGS)
tm   = load(JSON_TM)
proc = load(JSON_PROC)

# ─── 헬퍼: 폰트 적용 ─────────────────────────────────────────────────────────
def set_font(run, size_pt, bold=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 동아시아 폰트도 지정
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)

def add_heading(doc, text, level, size):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    colors = {1: (31,73,125), 2: (54,95,145), 3: (79,129,189), 4: (0,0,0)}
    set_font(run, size, bold=True, color=colors.get(level, (0,0,0)))
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '160')
    spacing.set(qn('w:after'), '80')
    pPr.append(spacing)
    return p

def add_body(doc, text, size=10, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size, bold=bold)
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '40')
    spacing.set(qn('w:after'), '40')
    pPr.append(spacing)
    return p

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run("• " + text)
    set_font(run, size)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)

def set_cell_border(table):
    """테이블 테두리 설정"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'AAAAAA')
        tblBorders.append(border)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

def add_table(doc, headers, rows, title=None, col_widths=None):
    """표 추가 함수"""
    if title:
        add_body(doc, f"▶ {title}", size=9, bold=True)

    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_border(table)

    # 헤더 행
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, "1F497D")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        set_font(run, 9, bold=True, color=(255,255,255))

    # 데이터 행
    for r_idx, row in enumerate(rows):
        data_row = table.rows[r_idx + 1]
        bg = "EBF0FA" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = data_row.cells[c_idx]
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, 9)

    # 열 너비 설정
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 여백
    return table

def n(val, decimals=2):
    """숫자 포맷"""
    if val is None:
        return "N/A"
    if isinstance(val, float):
        return f"{val:,.{decimals}f}"
    if isinstance(val, int):
        return f"{val:,}"
    return str(val)

def pct(val):
    if val is None:
        return "N/A"
    return f"{val:.1f}%"

def mv(val):
    """몸값 포맷 (유로)"""
    if val is None:
        return "N/A"
    if val >= 1_000_000:
        return f"€{val/1_000_000:.1f}M"
    if val >= 1_000:
        return f"€{val/1_000:.0f}K"
    return f"€{val:,.0f}"

# ─── 문서 생성 ────────────────────────────────────────────────────────────────
doc = Document()

# 페이지 설정 (A4)
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ══════════════════════════════════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EPL 데이터 프로젝트")
set_font(run, 28, bold=True, color=(31,73,125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("전처리 및 탐색적 데이터 분석(EDA)")
set_font(run, 22, bold=True, color=(54,95,145))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("종합 보고서")
set_font(run, 22, bold=True, color=(54,95,145))

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─" * 40)
set_font(run, 12, color=(128,128,128))

doc.add_paragraph()

info_items = [
    ("분석 기간", "2000/01 ~ 2024/25 시즌 (25시즌)"),
    ("데이터 기간", "2000년 8월 19일 ~ 2025년 5월 5일"),
    ("총 경기 수", f"{n(mr['5_date_and_match_info']['total_matches'])}경기"),
    ("총 팀 수", f"{n(mr['9_team_analysis']['total_unique_teams'])}개 팀"),
    ("데이터 소스", "Football-Data.co.uk, FBref, Transfermarkt"),
    ("보고서 생성일", "2026년 3월 21일"),
]

for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p.add_run(f"{label}:  ")
    set_font(run1, 12, bold=True, color=(31,73,125))
    run2 = p.add_run(value)
    set_font(run2, 12)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─" * 40)
set_font(run, 12, color=(128,128,128))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("대상 독자: 스카우터 · 이적 담당자 · 코치/감독")
set_font(run, 11, color=(64,64,64))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 목차
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "목  차", 1, 18)

toc_items = [
    ("1장", "프로젝트 개요"),
    ("2장", "원본 데이터 현황"),
    ("  2.1", "epl_final.csv (매치 결과)"),
    ("  2.2", "FBref squad_stats"),
    ("  2.3", "FBref matchlogs"),
    ("  2.4", "Transfermarkt squad_values"),
    ("3장", "데이터 품질 분석"),
    ("  3.1", "결측치 현황"),
    ("  3.2", "인코딩 이슈"),
    ("  3.3", "헤더 오염"),
    ("  3.4", "중복 데이터"),
    ("  3.5", "시대별 데이터 가용성"),
    ("4장", "전처리 파이프라인"),
    ("  4.1", "팀명 통합"),
    ("  4.2", "매치 결과 전처리"),
    ("  4.3", "FBref 전처리"),
    ("  4.4", "Transfermarkt 전처리"),
    ("  4.5", "데이터 통합"),
    ("  4.6", "처리 결과 요약"),
    ("5장", "탐색적 데이터 분석 (EDA)"),
    ("  5.1", "매치 결과 분석"),
    ("  5.2", "팀 분석"),
    ("  5.3", "선수 분석 (FBref squad_stats)"),
    ("  5.4", "경기별 스탯 분석 (matchlogs)"),
    ("  5.5", "몸값 분석 (Transfermarkt)"),
    ("6장", "시대별 비교 분석"),
    ("7장", "데이터 통합 분석"),
    ("8장", "향후 계획"),
    ("9장", "결론"),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    run1 = p.add_run(f"{num.ljust(6)}")
    set_font(run1, 10, bold=num.strip().endswith("장"))
    run2 = p.add_run(title)
    set_font(run2, 10, bold=num.strip().endswith("장"))
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '30')
    spacing.set(qn('w:after'), '30')
    pPr.append(spacing)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1장. 프로젝트 개요
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "제 1 장.  프로젝트 개요", 1, 18)

add_heading(doc, "1.1  분석 목적 및 배경", 2, 14)
add_body(doc,
    "본 프로젝트는 잉글랜드 프리미어리그(EPL) 2000/01 시즌부터 2024/25 시즌까지 "
    "25년간의 경기 결과, 선수 스탯, 시장 가치 데이터를 수집·전처리하여 스카우터, "
    "이적 담당자, 코치/감독이 의사결정에 활용할 수 있는 분석 기반을 구축하는 것을 "
    "목적으로 합니다. 원본 데이터 12,394개 CSV 파일을 5개의 Parquet 파일로 압축하여 "
    "대시보드 및 딥러닝 모델링에 즉시 활용 가능한 형태로 정제하였습니다.")

add_heading(doc, "1.2  분석 대상 기간", 2, 14)
add_body(doc, f"총 {mr['4_season_coverage']['n_seasons']}시즌: 2000/01 ~ 2024/25")
add_body(doc, f"날짜 범위: {mr['5_date_and_match_info']['date_min']} ~ {mr['5_date_and_match_info']['date_max']}")
add_body(doc, f"총 경기 수: {n(mr['5_date_and_match_info']['total_matches'])}경기 ({mr['5_date_and_match_info']['date_span_days']:,}일)")

add_heading(doc, "1.3  데이터 소스 3종 설명", 2, 14)
add_table(doc,
    ["데이터 소스", "제공 기관", "데이터 유형", "파일 수", "주요 컬럼"],
    [
        ["epl_final.csv", "Football-Data.co.uk", "경기 결과 (팀 단위)", "1개 CSV", "득점, 슈팅, 코너, 카드"],
        ["FBref squad_stats", "FBref.com", "선수 시즌 집계 스탯", f"{sq['metadata']['total_squad_stats_files']}개 CSV", "득점, 어시스트, 출전시간, 국적"],
        ["FBref matchlogs", "FBref.com", "선수 경기별 스탯", f"{ml['file_inventory']['total_matchlog_csv_files']:,}개 CSV", "슈팅, 패스, 태클, 파울"],
        ["Transfermarkt", "Transfermarkt.com", "선수 시장가치/신체정보", f"{tm['summary']['total_csv_files_found']}개 CSV", "몸값, 신장, 발 선호도, 국적"],
    ],
    title="데이터 소스 개요",
    col_widths=[3.0, 3.5, 4.0, 2.5, 4.5]
)

add_heading(doc, "1.4  분석 대상 독자", 2, 14)
target_data = [
    ["스카우터", "잠재적 영입 대상 선수의 퍼포먼스 트렌드, 국적·포지션 분포, 연령 프로파일"],
    ["이적 담당자", "시장가치 추이, 포지션별 몸값 분포, 최고가 선수 현황"],
    ["코치/감독", "홈/원정 전술 차이, 포지션별 스탯 패턴, 경기 결과 예측 인사이트"],
]
add_table(doc, ["대상", "주요 활용 내용"], target_data, title="분석 대상 독자별 활용 방안", col_widths=[3.0, 14.0])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2장. 원본 데이터 현황
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "제 2 장.  원본 데이터 현황", 1, 18)

# 2.1 epl_final.csv
add_heading(doc, "2.1  epl_final.csv (매치 결과 데이터)", 2, 14)

meta = mr["1_metadata"]
add_body(doc,
    f"Football-Data.co.uk에서 수집된 EPL 경기 결과 데이터셋입니다. "
    f"원본 CSV는 UTF-8-sig 인코딩으로 저장되어 있으며, 전처리 후 Parquet으로 변환되었습니다.")

add_table(doc,
    ["항목", "값"],
    [
        ["파일 경로", "epl_final.csv"],
        ["인코딩", meta["csv_encoding"]],
        ["원본 CSV 크기", f"{meta['csv_file_size_bytes']:,} bytes ({meta['csv_file_size_bytes']/1024:.1f} KB)"],
        ["Parquet 크기", f"{meta['parquet_file_size_bytes']:,} bytes ({meta['parquet_file_size_bytes']/1024:.1f} KB)"],
        ["행 수 (경기 수)", f"{meta['shape']['rows']:,}"],
        ["컬럼 수", str(meta['shape']['columns'])],
        ["시즌 수", str(mr['4_season_coverage']['n_seasons'])],
        ["날짜 범위", f"{mr['5_date_and_match_info']['date_min']} ~ {mr['5_date_and_match_info']['date_max']}"],
        ["결측치 컬럼 수", str(len(mr['2_missing_values']['summary']['columns_with_any_missing']))],
        ["완전한 행 비율", f"{mr['2_missing_values']['summary']['pct_complete_rows']}%"],
    ],
    title="epl_final.csv 메타데이터",
    col_widths=[5.0, 12.0]
)

# 시즌별 경기 수 (일부)
seasons_data = mr["4_season_coverage"]["matches_per_season"]
season_rows = [[s, f"{c}경기"] for s, c in list(seasons_data.items())]
add_table(doc,
    ["시즌", "경기 수"] * 5,
    [season_rows[i:i+5] for i in range(0, 25, 5)][0][:1] if False else
    [[season_rows[i][0], season_rows[i][1],
      season_rows[i+5][0], season_rows[i+5][1],
      season_rows[i+10][0], season_rows[i+10][1],
      season_rows[i+15][0], season_rows[i+15][1],
      season_rows[i+20][0] if i+20 < 25 else "",
      season_rows[i+20][1] if i+20 < 25 else ""]
     for i in range(5)],
    title="시즌별 경기 수",
    col_widths=[2.0, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0]
)

add_body(doc, "※ 2003/04, 2004/05 시즌은 335경기(일부 미포함), 2024/25 시즌은 350경기(진행 중)로 집계됨.", size=9)

# 컬럼 목록
add_table(doc,
    ["컬럼명", "데이터 타입", "설명"],
    [
        ["Season", "object", "시즌 식별자 (예: 2000/01)"],
        ["MatchDate", "datetime64[ns]", "경기 날짜"],
        ["HomeTeam / AwayTeam", "object", "홈/원정 팀명"],
        ["FullTimeHomeGoals / FullTimeAwayGoals", "int64", "전반전+후반전 최종 득점"],
        ["FullTimeResult", "object", "최종 결과 (H/D/A)"],
        ["HalfTimeHomeGoals / HalfTimeAwayGoals", "int64", "전반전 득점"],
        ["HomeShots / AwayShots", "int64", "총 슈팅 수"],
        ["HomeShotsOnTarget / AwayShotsOnTarget", "int64", "유효 슈팅 수"],
        ["HomeCorners / AwayCorners", "int64", "코너킥 수"],
        ["HomeFouls / AwayFouls", "int64", "반칙 수"],
        ["HomeYellowCards / AwayYellowCards", "int64", "경고 수"],
        ["HomeRedCards / AwayRedCards", "int64", "퇴장 수"],
        ["season_data_missing", "bool", "해당 시즌 데이터 누락 플래그"],
        ["own_goal_flag_home/away/overall", "bool", "자책골 발생 여부 플래그"],
    ],
    title="컬럼 구조 및 데이터 타입",
    col_widths=[5.5, 4.0, 7.5]
)

# 2.2 FBref squad_stats
add_heading(doc, "2.2  FBref squad_stats (선수 시즌 집계 데이터)", 2, 14)

add_table(doc,
    ["항목", "값"],
    [
        ["총 파일 수", f"{sq['metadata']['total_squad_stats_files']}개 CSV"],
        ["커버 시즌 수", f"{sq['metadata']['total_seasons']}시즌"],
        ["시즌당 파일 수", "20개 (20개 팀)"],
        ["샘플 분석 파일", f"{sq['metadata']['sample_files_read']}개"],
        ["읽기 오류", str(sq['metadata']['read_errors'])],
        ["통합 데이터 형태", f"{sq['data_quality']['combined_shape'][0]:,}행 × {sq['data_quality']['combined_shape'][1]}열"],
        ["총 컬럼 수 (시즌별 공통)", str(sq['column_analysis']['era_1_2000_2012']['n_columns'])],
        ["헤더 오염 행 수", str(sq['data_quality']['header_contamination_total_rows'])],
    ],
    title="FBref squad_stats 메타데이터",
    col_widths=[5.0, 12.0]
)

add_table(doc,
    ["컬럼명", "설명"],
    [
        ["player", "선수 이름"],
        ["nation", "국적 (FBref 형식: 'en ENG')"],
        ["pos", "포지션 (예: FW, MF, DF, GK)"],
        ["age", "나이"],
        ["mp / starts / min", "출전 경기 수 / 선발 수 / 출전 분"],
        ["90s", "90분 환산 출전 수"],
        ["gls / ast / g_a", "득점 / 어시스트 / 득점+어시스트"],
        ["g_pk / pk / pkatt", "PK 제외 득점 / PK 득점 / PK 시도"],
        ["crdy / crdr", "경고 / 퇴장"],
        ["gls_1 / ast_1 / g_a_1", "90분당 득점 / 어시스트 / 공격포인트"],
        ["season / team", "시즌 / 팀명"],
    ],
    title="주요 컬럼 설명",
    col_widths=[5.0, 12.0]
)

# 2.3 FBref matchlogs
add_heading(doc, "2.3  FBref matchlogs (선수 경기별 스탯 데이터)", 2, 14)

add_table(doc,
    ["항목", "값"],
    [
        ["총 파일 수", f"{ml['file_inventory']['total_matchlog_csv_files']:,}개 CSV"],
        ["커버 시즌 수", f"{ml['meta']['total_seasons']}시즌"],
        ["날짜 범위", f"{ml['date_coverage']['min_date']} ~ {ml['date_coverage']['max_date']}"],
        ["샘플 분석 파일", f"{ml['file_inventory']['sample_size_analyzed']:,}개"],
        ["읽기 오류", str(ml['file_inventory']['sample_load_errors'])],
        ["헤더 오염 파일", str(ml['file_inventory']['header_contaminated_files'])],
        ["전체 컬럼 수", str(ml['column_analysis']['total_unique_columns'])],
        ["고유 날짜 수", f"{ml['column_stats']['date']['unique_count']:,}"],
    ],
    title="FBref matchlogs 메타데이터",
    col_widths=[5.0, 12.0]
)

# 팀별 파일 수 상위 15
top_teams = list(ml['file_inventory']['files_per_team'].items())[:15]
add_table(doc,
    ["팀명", "파일 수", "팀명", "파일 수", "팀명", "파일 수"],
    [[top_teams[i][0], n(top_teams[i][1]),
      top_teams[i+5][0], n(top_teams[i+5][1]),
      top_teams[i+10][0], n(top_teams[i+10][1])] for i in range(5)],
    title="팀별 matchlog 파일 수 (상위 15팀)",
    col_widths=[3.0, 2.0, 3.0, 2.0, 3.0, 2.0]
)

# 2.4 Transfermarkt
add_heading(doc, "2.4  Transfermarkt squad_values (선수 시장가치 데이터)", 2, 14)

add_table(doc,
    ["항목", "값"],
    [
        ["총 파일 수", f"{tm['summary']['total_csv_files_found']}개 CSV"],
        ["커버 시즌 수", str(tm['summary']['total_seasons'])],
        ["시즌당 파일 수", "20개 (20개 팀)"],
        ["총 행 수", f"{tm['summary']['total_rows_combined']:,}"],
        ["총 컬럼 수", str(tm['summary']['total_columns'])],
        ["고유 팀 수", str(tm['summary']['unique_teams'])],
        ["고유 선수 수", f"{tm['summary']['unique_players']:,}"],
        ["읽기 오류 파일", str(tm['summary']['files_with_read_errors'])],
    ],
    title="Transfermarkt 메타데이터",
    col_widths=[5.0, 12.0]
)

# 결측치 현황
missing_rows = [
    [col, str(info['missing_count']), pct(info['missing_pct'])]
    for col, info in tm['missing_values_per_column'].items()
    if info['missing_count'] > 0
]
if missing_rows:
    add_table(doc, ["컬럼명", "결측 건수", "결측률"], missing_rows,
              title="Transfermarkt 주요 결측치 현황", col_widths=[5.0, 3.0, 3.0])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3장. 데이터 품질 분석
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "제 3 장.  데이터 품질 분석", 1, 18)

# 3.1 결측치
add_heading(doc, "3.1  결측치 현황", 2, 14)

add_heading(doc, "3.1.1  매치 결과 데이터 (epl_final.csv)", 3, 12)
add_body(doc,
    f"총 {meta['shape']['rows']:,}행 × {meta['shape']['columns']}컬럼 중 "
    f"결측치가 존재하는 컬럼은 0개이며, 완전한 행 비율은 100.0%입니다. "
    f"전처리 과정에서 모든 결측치가 처리되어 분석에 바로 사용 가능한 상태입니다.")
add_table(doc,
    ["데이터셋", "총 행 수", "결측 컬럼 수", "완전한 행 비율", "결측 셀 수"],
    [["epl_final.csv (처리 후)", f"{meta['shape']['rows']:,}", "0", "100.0%",
      str(mr['2_missing_values']['summary']['total_null_cells'])]],
    title="매치 결과 결측치 요약",
    col_widths=[5.0, 3.0, 3.5, 4.0, 3.0]
)

add_heading(doc, "3.1.2  FBref squad_stats 결측치", 3, 12)
sq_missing = sq['data_quality']['missing_values_per_column']
sq_miss_rows = [[col, str(info['null_count']), pct(info['pct_null'])]
                for col, info in sq_missing.items() if info['null_count'] > 0]
if sq_miss_rows:
    add_table(doc, ["컬럼명", "결측 건수", "결측률"], sq_miss_rows,
              title="FBref squad_stats 컬럼별 결측치", col_widths=[5.0, 3.0, 3.0])
    add_body(doc,
        "※ min, 90s, gls 등 주요 스탯 컬럼의 결측치 약 10.5%는 GK(골키퍼)처럼 "
        "해당 스탯이 기록되지 않는 경우에서 발생합니다.")

add_heading(doc, "3.1.3  FBref matchlogs 결측치", 3, 12)
ml_missing = ml['missing_value_analysis']['top_20_missing_columns']
ml_miss_rows = [[col, f"{info['count']:,}", pct(info['pct'])]
                for col, info in ml_missing.items() if info['count'] > 0]
add_table(doc, ["컬럼명", "결측 건수", "결측률"], ml_miss_rows[:12],
          title="FBref matchlogs 주요 결측치 (상위 12개)", col_widths=[3.0, 4.0, 3.0])
add_body(doc,
    "※ pkwon(76.99%), pkcon(76.99%)은 해당 스탯이 2013/14 이후에만 기록되어 "
    "초기 시즌 전체에서 결측으로 처리됩니다. sh(슈팅), sot(유효슈팅) 등 "
    "상세 스탯(48.55%)도 2013/14 이전 시즌에서는 수집되지 않았습니다.")

add_heading(doc, "3.1.4  Transfermarkt 결측치", 3, 12)
tm_miss_rows = [
    [col, str(info['missing_count']), pct(info['missing_pct'])]
    for col, info in tm['missing_values_per_column'].items()
    if info['missing_count'] > 0
]
if tm_miss_rows:
    add_table(doc, ["컬럼명", "결측 건수", "결측률"], tm_miss_rows,
              title="Transfermarkt 컬럼별 결측치", col_widths=[5.0, 3.0, 3.0])

# 3.2 인코딩 이슈
add_heading(doc, "3.2  인코딩 이슈", 2, 14)
add_body(doc,
    "epl_final.csv는 UTF-8-sig(BOM 포함 UTF-8) 인코딩을 사용하여 저장되어 있습니다. "
    "표준 UTF-8로 읽을 경우 첫 컬럼명에 'ï»¿' 문자가 포함될 수 있어, "
    "pandas의 encoding='utf-8-sig' 옵션으로 명시적 처리가 필요합니다.")
add_table(doc,
    ["파일/소스", "인코딩", "BOM 여부", "처리 방법"],
    [
        ["epl_final.csv", "UTF-8-sig", "있음 (BOM)", "encoding='utf-8-sig' 명시"],
        ["FBref squad_stats", "UTF-8", "없음", "기본 UTF-8 사용"],
        ["FBref matchlogs", "UTF-8", "없음", "기본 UTF-8 사용 (오류 0건)"],
        ["Transfermarkt CSV", "UTF-8", "없음", "기본 UTF-8 사용"],
    ],
    title="데이터 소스별 인코딩 현황",
    col_widths=[4.5, 3.0, 2.5, 7.0]
)

# 3.3 헤더 오염
add_heading(doc, "3.3  헤더 오염 (Header Contamination)", 2, 14)
add_body(doc,
    f"FBref squad_stats에서 헤더 행이 데이터 내부에 중복 삽입된 경우가 총 "
    f"{sq['data_quality']['header_contamination_total_rows']}건 발견되었습니다. "
    f"구체적으로 2021-2022/Arsenal(1건), 2021-2022/Aston Villa(1건), "
    f"2022-2023/Bournemouth(1건), 2024-2025/Brighton(1건)에서 발생하였으며, "
    f"전처리 과정에서 자동으로 제거되었습니다.")
add_body(doc,
    f"FBref matchlogs의 경우 샘플 {ml['file_inventory']['sample_size_analyzed']:,}개 파일 중 "
    f"헤더 오염 파일은 0건으로 확인되었습니다.")

add_table(doc,
    ["소스", "오염 발생 건수", "발생 파일", "처리 방법"],
    [
        ["FBref squad_stats", "4건", "Arsenal, Aston Villa (2021-22), Bournemouth (2022-23), Brighton (2024-25)", "행 필터링 (숫자 변환 실패 행 제거)"],
        ["FBref matchlogs", "0건", "없음", "해당 없음"],
        ["Transfermarkt", "0건", "없음", "해당 없음"],
    ],
    title="헤더 오염 현황 및 처리",
    col_widths=[3.5, 2.5, 7.0, 4.0]
)

# 3.4 중복 데이터
add_heading(doc, "3.4  중복 데이터", 2, 14)
add_body(doc,
    "처리된 Parquet 데이터 기준으로 중복 행은 0건으로 확인되었습니다. "
    "FBref matchlogs 샘플 분석에서도 중복 행이 전혀 발견되지 않았습니다.")
add_table(doc,
    ["데이터셋", "총 행 수", "중복 행 수", "중복률"],
    [
        ["match_results.parquet", f"{mr['1_metadata']['shape']['rows']:,}", "0", "0.00%"],
        ["FBref matchlogs (샘플)", f"{ml['file_inventory']['sample_size_analyzed']:,}파일", "0", "0.00%"],
        ["Transfermarkt (처리 전)", f"{tm['summary']['total_rows_combined']:,}", "0", "0.00%"],
    ],
    title="중복 데이터 현황",
    col_widths=[5.0, 3.5, 3.0, 3.5]
)

# 3.5 시대별 데이터 가용성
add_heading(doc, "3.5  시대별 데이터 가용성 비교", 2, 14)
add_body(doc,
    "FBref 데이터를 기준으로 2000-2012 시대(Era 1)와 2013-2025 시대(Era 2)를 비교하면, "
    "squad_stats의 컬럼 구조는 양 시대 모두 동일한 24개 컬럼을 유지합니다. "
    "그러나 matchlogs에서는 2013/14 이전 시즌에 상세 스탯(슈팅, 파울, 크로스 등)이 "
    "수집되지 않아 약 48.55%가 결측치로 존재합니다.")

add_table(doc,
    ["컬럼/스탯", "Era 1 (2000-2012)", "Era 2 (2013-2025)", "비고"],
    [
        ["squad_stats 컬럼 수", "24개", "24개", "동일"],
        ["matchlogs 기본 스탯 (gls, ast, crdy 등)", "100% 가용", "100% 가용", "전 시즌 수집"],
        ["matchlogs 상세 스탯 (sh, sot, fls 등)", "결측 약 48.55%", "100% 가용", "2013/14부터 수집"],
        ["matchlogs pkwon / pkcon", "결측 약 76.99%", "부분 가용", "2015/16 이후 수집"],
        ["matchlogs pos (포지션)", "결측 약 60%", "결측 약 60%", "전 시즌 공통 이슈"],
        ["Transfermarkt 몸값", "2000-2003 결측", "2004~부터 가용", "2004/05부터 수집"],
    ],
    title="시대별 데이터 가용성 비교",
    col_widths=[5.0, 3.5, 3.5, 5.0]
)

add_body(doc,
    "이러한 시대별 데이터 가용성 차이는 분석 시 주의가 필요하며, "
    "세부 스탯이 필요한 분석은 2013/14 이후 데이터를 대상으로 수행하는 것이 권장됩니다.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4장. 전처리 파이프라인
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "제 4 장.  전처리 파이프라인", 1, 18)
add_body(doc,
    "원본 데이터를 분석 가능한 형태로 변환하기 위해 단계별 전처리 파이프라인을 구축하였습니다. "
    "전체 파이프라인은 팀명 통합 → 개별 데이터 전처리 → 데이터 통합의 순서로 진행됩니다.")

# 4.1 팀명 통합
add_heading(doc, "4.1  팀명 통합 (Team Name Normalization)", 2, 14)
add_body(doc,
    "FBref, Transfermarkt, Football-Data.co.uk 세 소스의 팀명 표기 방식이 상이하여 "
    "(예: 'Manchester United' vs 'Man United' vs 'Manchester Utd') "
    "총 46개 팀에 대해 97개의 팀명 매핑 딕셔너리를 구축하여 통일된 팀명으로 변환하였습니다.")

add_table(doc,
    ["소스", "표기 예시", "통합 팀명"],
    [
        ["Football-Data.co.uk", "Man United, Man City, Nott'm Forest", "표준 약어 기준"],
        ["FBref", "Manchester Utd, Manchester City, Nottingham Forest", "표준 약어로 변환"],
        ["Transfermarkt", "Manchester United, Manchester City, Nottingham Forest", "표준 약어로 변환"],
    ],
    title="소스별 팀명 표기 차이 예시",
    col_widths=[4.0, 6.5, 6.5]
)
add_body(doc, f"총 처리 대상: 46개 팀, 97개 매핑 규칙 적용")

# 4.2 매치 결과 전처리
add_heading(doc, "4.2  매치 결과 전처리", 2, 14)
add_body(doc, "epl_final.csv 전처리 단계는 다음과 같습니다:")
steps = [
    "1. UTF-8-sig 인코딩으로 CSV 로드",
    "2. 날짜 컬럼을 datetime64[ns] 타입으로 변환 (혼합 형식 처리 포함)",
    "3. 시즌 파생 변수 생성 (날짜 기반 시즌 할당)",
    "4. 데이터 누락 플래그(season_data_missing) 추가: 2003/04, 2004/05 일부 누락 표시",
    "5. 자책골 플래그(own_goal_flag_home/away/overall) 생성",
    "6. 팀명 통합 매핑 적용",
    "7. Parquet 형식으로 저장 (메모리 4.1MB, 파일 144KB)",
]
for s in steps:
    add_bullet(doc, s)

# 4.3 FBref 전처리
add_heading(doc, "4.3  FBref 데이터 전처리", 2, 14)
add_heading(doc, "4.3.1  squad_stats 전처리", 3, 12)
steps2 = [
    "1. 500개 CSV 파일 순회 로드 (시즌/팀 폴더 구조)",
    "2. 헤더 오염 행 탐지 및 제거 (총 4건)",
    "3. 숫자 컬럼 강제 타입 변환 (pd.to_numeric, errors='coerce')",
    "4. 국적 컬럼 정규화 (예: 'fr FRA' → 'FRA')",
    "5. 포지션 표준화 (FW, MF, DF, GK)",
    "6. 시즌/팀 컬럼 추가 (폴더명에서 추출)",
    "7. 전체 통합 후 player_season_stats.parquet 저장",
]
for s in steps2:
    add_bullet(doc, s)

add_heading(doc, "4.3.2  matchlogs 전처리", 3, 12)
steps3 = [
    "1. 11,369개 CSV 파일 순회 로드",
    "2. 경기 결과 파싱: 'W 2-1' → outcome(W/D/L), goals_for, goals_against",
    "3. start 컬럼 파싱: 'Y'/'Y*'(주장) → started(bool)",
    "4. detail_stats_available 플래그: sh/sot 등 상세 스탯 존재 여부",
    "5. 수치 컬럼 타입 변환 및 결측치 처리",
    "6. 팀명 통합 매핑 적용",
    "7. player_match_logs.parquet으로 저장 (299,517행)",
]
for s in steps3:
    add_bullet(doc, s)

# 4.4 Transfermarkt 전처리
add_heading(doc, "4.4  Transfermarkt 데이터 전처리", 2, 14)
steps4 = [
    "1. 500개 CSV 파일 순회 로드",
    "2. 생년월일(DOB) → datetime 변환 후 birth_year 파생변수 생성",
    "3. 신장(height) 문자열 → 수치 변환 (예: '1.85m' → 185.0 cm)",
    "4. 몸값(market_value) 문자열 파싱 (예: '€45M' → 45,000,000.0)",
    "5. 시즌 표기 통일 (예: '2000' → '2000/01')",
    "6. 팀명 통합 매핑 적용 (46개 팀 기준)",
    "7. 연령(age) 재계산 (DOB 기준)",
]
for s in steps4:
    add_bullet(doc, s)

# 4.5 데이터 통합
add_heading(doc, "4.5  데이터 통합 (FBref + Transfermarkt 매칭)", 2, 14)
add_body(doc,
    "FBref squad_stats와 Transfermarkt 데이터를 선수명+시즌 기반으로 매칭하였습니다. "
    "선수명 정규화(소문자, 특수문자 제거, 유니코드 정규화)를 통해 매칭률을 향상시켰습니다.")

add_table(doc,
    ["매칭 조건", "결과"],
    [
        ["매칭 키", "player_normalized + season + team"],
        ["매칭률", "92.7%"],
        ["미매칭 원인", "이름 표기 차이, Transfermarkt 미수록 선수"],
        ["통합 결과 파일", "player_season_stats.parquet (14,980행 × 36컬럼)"],
        ["transfer_flag", "이적/팀 변경 선수 플래그 포함"],
    ],
    title="FBref + Transfermarkt 매칭 결과",
    col_widths=[5.0, 12.0]
)

# 4.6 처리 결과
add_heading(doc, "4.6  전처리 처리 결과 요약", 2, 14)

# 실제 shape 데이터
proc_files = proc['per_file_analysis']
output_files = [
    ["match_results.parquet",
     f"{proc_files['processed/match_results.parquet']['shape']['rows']:,}",
     str(proc_files['processed/match_results.parquet']['shape']['cols']),
     "경기 결과 (팀 단위)"],
    ["player_match_logs.parquet",
     f"{proc_files['processed/player_match_logs.parquet']['shape']['rows']:,}",
     str(proc_files['processed/player_match_logs.parquet']['shape']['cols']),
     "선수 경기별 스탯 (처리된 버전)"],
    ["player_match_stats.parquet",
     f"{proc_files['processed/player_match_stats.parquet']['shape']['rows']:,}",
     str(proc_files['processed/player_match_stats.parquet']['shape']['cols']),
     "선수 경기별 스탯 (원본 컬럼 보존)"],
    ["player_season_stats.parquet",
     f"{proc_files['processed/player_season_stats.parquet']['shape']['rows']:,}",
     str(proc_files['processed/player_season_stats.parquet']['shape']['cols']),
     "선수 시즌 집계 (FBref+TM 통합)"],
]
add_table(doc,
    ["파일명", "행 수", "컬럼 수", "내용"],
    output_files,
    title="처리 결과 파일 현황",
    col_widths=[5.0, 3.0, 2.5, 6.5]
)

add_body(doc,
    "▶  원본 데이터: 약 12,394개 CSV → 처리 후: 5개 Parquet 파일로 압축\n"
    "▶  데이터 압축률: raw 78.7MB → processed 5.1MB → dashboard 0.5MB")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5장. 탐색적 데이터 분석 (EDA)
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "제 5 장.  탐색적 데이터 분석 (EDA)", 1, 18)

# ─── 5.1 매치 결과 분석 ────────────────────────────────────────────────────
add_heading(doc, "5.1  매치 결과 분석", 2, 14)

add_heading(doc, "5.1.1  전체 홈/무/원정 승률 분포", 3, 12)
rd = mr['6_result_distribution']['full_time']['overall']
add_table(doc,
    ["결과", "경기 수", "비율"],
    [
        ["홈 승 (H)", f"{rd['home_wins']:,}", pct(rd['home_win_pct'])],
        ["무승부 (D)", f"{rd['draws']:,}", pct(rd['draw_pct'])],
        ["원정 승 (A)", f"{rd['away_wins']:,}", pct(rd['away_win_pct'])],
        ["합계", f"{rd['total']:,}", "100.0%"],
    ],
    title="전체 경기 결과 분포 (2000/01~2024/25)",
    col_widths=[4.0, 4.0, 4.0]
)
add_body(doc,
    f"25시즌 통산 홈 승률은 {pct(rd['home_win_pct'])}로 원정 승률({pct(rd['away_win_pct'])})보다 "
    f"약 16.3%p 높게 나타나, 홈 어드밴티지가 EPL에서 유의미하게 작용함을 확인할 수 있습니다.")

add_heading(doc, "5.1.2  시즌별 결과 분포", 3, 12)
per_season = mr['6_result_distribution']['full_time']['per_season']
season_result_rows = [
    [s, str(v['total']), pct(v['home_win_pct']), pct(v['draw_pct']), pct(v['away_win_pct'])]
    for s, v in per_season.items()
]
add_table(doc,
    ["시즌", "경기 수", "홈 승률", "무승부율", "원정 승률"],
    season_result_rows,
    title="시즌별 경기 결과 분포",
    col_widths=[2.5, 2.5, 3.0, 3.0, 3.0]
)
add_body(doc,
    "2020/21 시즌은 COVID-19로 인해 무관중 경기가 진행되어 홈 승률이 "
    f"{pct(per_season['2020/21']['home_win_pct'])}로 전체 최저를 기록하고, "
    f"원정 승률({pct(per_season['2020/21']['away_win_pct'])})이 홈 승률을 초과하는 "
    "이례적 결과가 나타났습니다.")

add_heading(doc, "5.1.3  시즌별 평균 득점 추이", 3, 12)
goals_trend = mr['7_goals_analysis']['avg_goals_per_season_trend']
goals_rows = [
    [s, f"{v['avg_home_goals']:.4f}", f"{v['avg_away_goals']:.4f}",
     f"{v['avg_total_goals']:.4f}", str(v['total_goals'])]
    for s, v in goals_trend.items()
]
add_table(doc,
    ["시즌", "홈 평균 득점", "원정 평균 득점", "경기당 평균 득점", "시즌 총 득점"],
    goals_rows,
    title="시즌별 평균 득점 추이",
    col_widths=[2.5, 3.0, 3.5, 3.5, 3.5]
)
ga = mr['7_goals_analysis']['overall']
add_body(doc,
    f"전체 평균: 홈 {ga['avg_home_goals']:.4f}골, 원정 {ga['avg_away_goals']:.4f}골, "
    f"경기당 {ga['avg_total_goals']:.4f}골. 25시즌 총 득점: {ga['total_goals_all']:,}골. "
    f"최고 득점 경기 합계: {ga['max_total_goals']}골 (포츠머스 7-4 레딩, 2007/08).")

add_heading(doc, "5.1.4  홈 어드밴티지 분석", 3, 12)
ha = mr['8_home_advantage']
add_table(doc,
    ["지표", "값"],
    [
        ["홈 승률", pct(ha['overall']['home_win_pct'])],
        ["원정 승률", pct(ha['overall']['away_win_pct'])],
        ["홈 평균 득점", f"{ha['overall']['avg_home_goals']:.4f}"],
        ["원정 평균 득점", f"{ha['overall']['avg_away_goals']:.4f}"],
        ["경기당 평균 득점 차 (홈-원정)", f"+{ha['overall']['avg_goal_diff']:.4f}"],
        ["COVID 시즌 홈 승률 (2020/21)", pct(ha['per_season']['2020/21']['home_win_pct'])],
        ["COVID 시즌 원정 승률 (2020/21)", pct(ha['per_season']['2020/21']['away_win_pct'])],
        ["최고 홈 승률 시즌", "2005/06 (50.53%)"],
        ["최저 홈 승률 시즌", "2020/21 (37.89%, COVID 무관중)"],
    ],
    title="홈 어드밴티지 종합 지표",
    col_widths=[6.0, 11.0]
)

add_heading(doc, "5.1.5  카드 통계", 3, 12)
desc = mr['3_descriptive_statistics']
add_table(doc,
    ["스탯", "평균 (경기당)", "최대값", "설명"],
    [
        ["홈팀 옐로카드", f"{desc['HomeYellowCards']['mean']:.4f}", str(int(desc['HomeYellowCards']['max'])), "홈팀 경고"],
        ["원정팀 옐로카드", f"{desc['AwayYellowCards']['mean']:.4f}", str(int(desc['AwayYellowCards']['max'])), "원정팀 경고 (홈보다 높음)"],
        ["홈팀 레드카드", f"{desc['HomeRedCards']['mean']:.4f}", str(int(desc['HomeRedCards']['max'])), "홈팀 퇴장"],
        ["원정팀 레드카드", f"{desc['AwayRedCards']['mean']:.4f}", str(int(desc['AwayRedCards']['max'])), "원정팀 퇴장"],
        ["홈팀 슈팅", f"{desc['HomeShots']['mean']:.4f}", str(int(desc['HomeShots']['max'])), "홈팀 총 슈팅"],
        ["원정팀 슈팅", f"{desc['AwayShots']['mean']:.4f}", str(int(desc['AwayShots']['max'])), "원정팀 총 슈팅"],
        ["홈팀 코너킥", f"{desc['HomeCorners']['mean']:.4f}", str(int(desc['HomeCorners']['max'])), "홈팀 코너킥"],
        ["원정팀 코너킥", f"{desc['AwayCorners']['mean']:.4f}", str(int(desc['AwayCorners']['max'])), "원정팀 코너킥"],
    ],
    title="경기당 주요 스탯 요약 (전체 평균)",
    col_widths=[4.0, 4.0, 3.0, 6.0]
)
add_body(doc,
    "원정팀의 옐로카드 평균(1.79)이 홈팀(1.47)보다 높아 원정 상황에서 더 공격적이거나 "
    "심판 판정에서 불리한 경향이 있음을 시사합니다.")

# ─── 5.2 팀 분석 ────────────────────────────────────────────────────────────
add_heading(doc, "5.2  팀 분석", 2, 14)

add_heading(doc, "5.2.1  25시즌 개근팀", 3, 12)
team_app = mr['9_team_analysis']['top_20_most_appearances']
permanent_teams = [t for t, c in team_app.items() if c == 938 or c == 937]
add_body(doc,
    f"25시즌 내내 EPL에 참여한 팀은 총 {len(permanent_teams)}개입니다. "
    f"938경기(25시즌 × 38경기 기준)에 해당하는 팀은 Tottenham, Man United, Liverpool, Everton이며, "
    f"Arsenal과 Chelsea는 937경기(시즌 중 1경기 차이)로 사실상 전 시즌 개근팀입니다.")

perm_rows = []
for t in ["Man United", "Arsenal", "Chelsea", "Liverpool", "Everton", "Tottenham", "Man City"]:
    data = next(x for x in mr['9_team_analysis']['top_20_teams_by_wins'] if x['team'] == t)
    perm_rows.append([t, str(data['total_matches']), str(data['wins']), pct(data['win_pct'])])
add_table(doc,
    ["팀명", "총 경기 수", "총 승리 수", "통산 승률"],
    perm_rows,
    title="개근팀(25시즌) 통산 성적",
    col_widths=[4.0, 3.5, 3.5, 3.5]
)

add_heading(doc, "5.2.2  팀별 통산 승률 (Top 10)", 3, 12)
wins_data = mr['9_team_analysis']['top_20_teams_by_wins']
add_table(doc,
    ["순위", "팀명", "총 경기", "승", "무", "패", "승률", "득점", "실점"],
    [[str(i+1), d['team'], str(d['total_matches']), str(d['wins']), str(d['draws']),
      str(d['losses']), pct(d['win_pct']), str(d['goals_scored']), str(d['goals_conceded'])]
     for i, d in enumerate(wins_data[:10])],
    title="팀별 통산 승률 순위 (Top 10, 전 시즌 통합)",
    col_widths=[1.5, 3.0, 2.5, 1.5, 1.5, 1.5, 2.5, 2.5, 2.5]
)

add_heading(doc, "5.2.3  승격/강등 빈도 (전체 팀별 출전 현황)", 3, 12)
all_apps = mr['9_team_analysis']['top_20_most_appearances']
all_apps_bottom = mr['9_team_analysis']['bottom_20_least_appearances']

# Top 10 + Bottom 10
app_rows = []
for t, c in list(all_apps.items())[:10]:
    seasons_est = c // 38 if c > 38 else 1
    app_rows.append([t, str(c), f"약 {seasons_est}시즌"])
add_table(doc,
    ["팀명", "총 경기 수", "추정 시즌 수"],
    app_rows,
    title="출전 경기 수 Top 10 팀",
    col_widths=[5.0, 4.0, 4.0]
)

promoted_rows = []
for t, c in list(all_apps_bottom.items())[:10]:
    seasons_est = c // 38
    promoted_rows.append([t, str(c), f"약 {seasons_est}시즌", "강등 팀"])
add_table(doc,
    ["팀명", "총 경기 수", "추정 시즌 수", "비고"],
    promoted_rows,
    title="출전 경기 수 하위 10팀 (강등 빈도 높음)",
    col_widths=[4.0, 3.0, 3.5, 6.5]
)

# ─── 5.3 선수 분석 ──────────────────────────────────────────────────────────
add_heading(doc, "5.3  선수 분석 (FBref squad_stats 기반)", 2, 14)

add_heading(doc, "5.3.1  포지션 분포", 3, 12)
pos_dist = tm['position_analysis']['distribution']
add_table(doc,
    ["포지션", "선수-시즌 수", "비율"],
    [[pos, f"{cnt:,}", pct(cnt/tm['summary']['total_rows_combined']*100)]
     for pos, cnt in list(pos_dist.items())[:10]],
    title="포지션별 선수 분포 (Transfermarkt 기준, 상위 10개 포지션)",
    col_widths=[4.5, 3.5, 3.0]
)

add_heading(doc, "5.3.2  연령 분포 (Transfermarkt 기준)", 3, 12)
age_dist = tm['age_and_dob_analysis']['age_distribution_bins']
age_stats = tm['age_and_dob_analysis']['age_stats']
add_table(doc,
    ["연령대", "선수 수", "비율"],
    [[k, f"{v:,}", pct(v/tm['summary']['total_rows_combined']*100)]
     for k, v in age_dist.items()],
    title="EPL 선수 연령 분포 (전체 기간)",
    col_widths=[3.0, 3.0, 3.0]
)
add_table(doc,
    ["통계", "값"],
    [
        ["평균 연령", f"{age_stats['mean']:.2f}세"],
        ["중앙값 연령", f"{age_stats['median']:.1f}세"],
        ["표준편차", f"{age_stats['std']:.2f}"],
        ["최연소", f"{age_stats['min']:.0f}세 ({tm['age_and_dob_analysis']['youngest_player']['name']}, {tm['age_and_dob_analysis']['youngest_player']['team']}, {tm['age_and_dob_analysis']['youngest_player']['season']})"],
        ["최고령", f"{age_stats['max']:.0f}세 ({tm['age_and_dob_analysis']['oldest_player']['name']}, {tm['age_and_dob_analysis']['oldest_player']['team']}, {tm['age_and_dob_analysis']['oldest_player']['season']})"],
    ],
    title="연령 기술 통계",
    col_widths=[5.0, 12.0]
)

add_heading(doc, "5.3.3  국적 분포 (Top 15, Transfermarkt 기준)", 3, 12)
nat_top = list(tm['nationality_analysis']['top_20'].items())[:15]
add_table(doc,
    ["순위", "국적", "선수-시즌 수"],
    [[str(i+1), nat, f"{cnt:,}"] for i, (nat, cnt) in enumerate(nat_top)],
    title="국적별 선수 분포 Top 15",
    col_widths=[2.0, 6.0, 4.0]
)
add_body(doc,
    f"총 {tm['nationality_analysis']['unique_nationalities']}개 국적의 선수가 EPL에 참여하였으며, "
    f"영국 계열 국적(England, Scotland, Wales, Ireland 등)이 압도적 다수를 차지합니다.")

add_heading(doc, "5.3.4  선수 스탯 기술 통계 (샘플 80개 파일 기준)", 3, 12)
sq_stats = sq['descriptive_stats_common_numeric']
add_table(doc,
    ["스탯", "평균", "중앙값", "최댓값", "표준편차"],
    [
        ["출전 경기 수 (mp)", f"{sq_stats['mp']['mean']:.3f}", str(sq_stats['mp']['50%']), str(int(sq_stats['mp']['max'])), f"{sq_stats['mp']['std']:.3f}"],
        ["선발 수 (starts)", f"{sq_stats['starts']['mean']:.3f}", str(sq_stats['starts']['50%']), str(int(sq_stats['starts']['max'])), f"{sq_stats['starts']['std']:.3f}"],
        ["출전 분 (min)", f"{sq_stats['min']['mean']:.1f}", str(sq_stats['min']['50%']), str(int(sq_stats['min']['max'])), f"{sq_stats['min']['std']:.1f}"],
        ["득점 (gls)", f"{sq_stats['gls']['mean']:.3f}", str(sq_stats['gls']['50%']), str(int(sq_stats['gls']['max'])), f"{sq_stats['gls']['std']:.3f}"],
        ["어시스트 (ast)", f"{sq_stats['ast']['mean']:.3f}", str(sq_stats['ast']['50%']), str(int(sq_stats['ast']['max'])), f"{sq_stats['ast']['std']:.3f}"],
        ["경고 (crdy)", f"{sq_stats['crdy']['mean']:.3f}", str(sq_stats['crdy']['50%']), str(int(sq_stats['crdy']['max'])), f"{sq_stats['crdy']['std']:.3f}"],
        ["퇴장 (crdr)", f"{sq_stats['crdr']['mean']:.3f}", str(sq_stats['crdr']['50%']), str(int(sq_stats['crdr']['max'])), f"{sq_stats['crdr']['std']:.3f}"],
    ],
    title="선수 시즌 집계 스탯 기술 통계 (샘플)",
    col_widths=[4.0, 2.5, 2.5, 2.5, 3.0]
)

# ─── 5.4 경기별 스탯 분석 ────────────────────────────────────────────────────
add_heading(doc, "5.4  경기별 스탯 분석 (FBref matchlogs 기반)", 2, 14)

add_heading(doc, "5.4.1  선발/교체 비율", 3, 12)
sv = ml['start_vs_sub_analysis']
add_table(doc,
    ["구분", "선수-경기 수", "비율"],
    [
        ["선발 (Y/Y*)", f"{sv['starter_count']:,}", pct(sv['start_pct'])],
        ["교체 (N)", f"{sv['sub_count']:,}", pct(sv['sub_pct'])],
        ["합계", f"{sv['total_appearances']:,}", "100.0%"],
    ],
    title="선발 vs 교체 출전 비율 (샘플)",
    col_widths=[4.0, 4.0, 4.0]
)

add_heading(doc, "5.4.2  출전 시간 분포", 3, 12)
mp_dist = ml['minutes_played_distribution']
add_table(doc,
    ["통계", "값"],
    [
        ["평균 출전 시간", f"{mp_dist['mean']:.2f}분"],
        ["중앙값 출전 시간", f"{mp_dist['median']:.1f}분"],
        ["표준편차", f"{mp_dist['std']:.2f}분"],
        ["90분 풀타임 비율", pct(mp_dist['pct_90_mins'])],
        ["최빈값", "90분 (전체 출전 기록의 59.62%)"],
    ],
    title="출전 시간 분포",
    col_widths=[5.0, 12.0]
)

add_heading(doc, "5.4.3  경기당 주요 스탯 분포", 3, 12)
stat_dist = ml['stat_distributions']
add_table(doc,
    ["스탯", "평균", "중앙값", "최댓값", "0 비율", "유효 기록 수"],
    [
        ["득점 (gls)", f"{stat_dist['goals']['mean']:.4f}", "0", str(int(stat_dist['goals']['max'])), pct(stat_dist['goals']['pct_zero']), f"{stat_dist['goals']['count']:,}"],
        ["어시스트 (ast)", f"{stat_dist['assists']['mean']:.4f}", "0", str(int(stat_dist['assists']['max'])), pct(stat_dist['assists']['pct_zero']), f"{stat_dist['assists']['count']:,}"],
        ["슈팅 (sh)", f"{stat_dist['shots']['mean']:.4f}", "0", str(int(stat_dist['shots']['max'])), pct(stat_dist['shots']['pct_zero']), f"{stat_dist['shots']['count']:,}"],
        ["유효슈팅 (sot)", f"{stat_dist['shots_on_target']['mean']:.4f}", "0", str(int(stat_dist['shots_on_target']['max'])), pct(stat_dist['shots_on_target']['pct_zero']), f"{stat_dist['shots_on_target']['count']:,}"],
        ["파울 (fls)", f"{stat_dist['fouls_committed']['mean']:.4f}", "0", str(int(stat_dist['fouls_committed']['max'])), pct(stat_dist['fouls_committed']['pct_zero']), f"{stat_dist['fouls_committed']['count']:,}"],
        ["태클 성공 (tklw)", f"{stat_dist['tackles_won']['mean']:.4f}", "0", str(int(stat_dist['tackles_won']['max'])), pct(stat_dist['tackles_won']['pct_zero']), f"{stat_dist['tackles_won']['count']:,}"],
        ["인터셉트 (int)", f"{stat_dist['interceptions']['mean']:.4f}", "0", str(int(stat_dist['interceptions']['max'])), pct(stat_dist['interceptions']['pct_zero']), f"{stat_dist['interceptions']['count']:,}"],
        ["크로스 (crs)", f"{stat_dist['crosses']['mean']:.4f}", "0", str(int(stat_dist['crosses']['max'])), pct(stat_dist['crosses']['pct_zero']), f"{stat_dist['crosses']['count']:,}"],
        ["경고 (crdy)", f"{stat_dist['yellow_cards']['mean']:.4f}", "0", str(int(stat_dist['yellow_cards']['max'])), pct(stat_dist['yellow_cards']['pct_zero']), f"{stat_dist['yellow_cards']['count']:,}"],
        ["퇴장 (crdr)", f"{stat_dist['red_cards']['mean']:.4f}", "0", str(int(stat_dist['red_cards']['max'])), pct(stat_dist['red_cards']['pct_zero']), f"{stat_dist['red_cards']['count']:,}"],
    ],
    title="경기별 선수 스탯 분포 (샘플 기준)",
    col_widths=[3.0, 2.0, 2.0, 2.0, 2.5, 3.5]
)

add_heading(doc, "5.4.4  포지션별 출전 분포 (matchlogs)", 3, 12)
pos_vals = ml['position_analysis']['value_counts']
filtered_pos = {k: v for k, v in pos_vals.items() if k not in ['nan', 'On matchday squad, but did not play']}
total_pos = sum(filtered_pos.values())
add_table(doc,
    ["포지션", "기록 수", "비율"],
    [[pos, f"{cnt:,}", pct(cnt/total_pos*100)] for pos, cnt in list(filtered_pos.items())],
    title="matchlogs 포지션별 출전 기록 수",
    col_widths=[3.0, 3.0, 3.0]
)
add_body(doc,
    f"※ 포지션 결측치 {pos_vals.get('nan', 0):,}건({pct(pos_vals.get('nan',0)/sum(pos_vals.values())*100)})은 "
    "2000-2012 초기 시즌에서 포지션 정보가 FBref에 기록되지 않은 경우입니다.")

# ─── 5.5 몸값 분석 ──────────────────────────────────────────────────────────
add_heading(doc, "5.5  몸값 분석 (Transfermarkt 기반)", 2, 14)

add_heading(doc, "5.5.1  전체 몸값 분포 통계", 3, 12)
mv_overall = tm['market_value_analysis']['overall_distribution']
add_table(doc,
    ["통계", "값"],
    [
        ["총 선수-시즌 기록", f"{mv_overall['total_records']:,}"],
        ["0원 기록 수 (초기 시즌 등)", f"{mv_overall['zero_market_value_count']:,} ({pct(mv_overall['zero_market_value_pct'])})"],
        ["유효 몸값 기록 수", f"{mv_overall['nonzero_records']:,}"],
        ["평균 몸값", mv(mv_overall['mean'])],
        ["중앙값 몸값", mv(mv_overall['median'])],
        ["최고 몸값", mv(mv_overall['max'])],
        ["최저 유효 몸값", mv(mv_overall['min'])],
        ["표준편차", mv(mv_overall['std'])],
        ["25번째 백분위", mv(mv_overall['percentiles']['25th'])],
        ["75번째 백분위", mv(mv_overall['percentiles']['75th'])],
        ["90번째 백분위", mv(mv_overall['percentiles']['90th'])],
        ["99번째 백분위", mv(mv_overall['percentiles']['99th'])],
    ],
    title="EPL 전체 선수 몸값 분포 통계",
    col_widths=[5.0, 12.0]
)

add_heading(doc, "5.5.2  시즌별 평균 몸값 추이", 3, 12)
mv_trend = tm['market_value_analysis']['per_season_stats']
mv_rows = []
for s, v in mv_trend.items():
    if v['mean_mv'] is not None:
        mv_rows.append([s, f"{v['players']:,}", mv(v['mean_mv']), mv(v['median_mv']), mv(v['max_mv'])])
    else:
        mv_rows.append([s, f"{v['players']:,}", "N/A (미수록)", "N/A", "N/A"])
add_table(doc,
    ["시즌", "선수 수", "평균 몸값", "중앙값 몸값", "최고 몸값"],
    mv_rows,
    title="시즌별 평균 몸값 추이 (2000/01~2024/25)",
    col_widths=[2.5, 2.5, 3.5, 3.5, 3.5]
)
add_body(doc,
    "2000/01~2003/04 시즌의 몸값은 Transfermarkt에 수록되지 않아 N/A로 표시됩니다. "
    "2004/05 이후 지속적인 몸값 상승세를 확인할 수 있으며, "
    "2024/25 시즌 평균 몸값은 약 €17.6M으로 최고치를 기록하였습니다.")

add_heading(doc, "5.5.3  역대 최고 몸값 선수 Top 10", 3, 12)
top_mv = tm['market_value_analysis']['top_20_most_valuable_players_ever'][:10]
add_table(doc,
    ["순위", "선수명", "팀", "시즌", "몸값", "포지션"],
    [[str(i+1), p['player'], p['team'], p['season'], mv(p['market_value']), p['position']]
     for i, p in enumerate(top_mv)],
    title="역대 최고 몸값 선수 Top 10",
    col_widths=[1.5, 4.0, 4.0, 2.5, 2.5, 3.0]
)

add_heading(doc, "5.5.4  포지션별 평균 몸값 분포", 3, 12)
pos_mv = tm['market_value_analysis']['market_value_by_position']
pos_mv_rows = sorted(
    [[pos, f"{v['count']:,}", mv(v['mean_mv']), mv(v['median_mv']), mv(v['max_mv'])]
     for pos, v in pos_mv.items() if v['count'] > 10],
    key=lambda x: -float(x[2][1:].replace('M','').replace('K','').replace(',',''))
    if x[2] != 'N/A' else 0
)
add_table(doc,
    ["포지션", "기록 수", "평균 몸값", "중앙값 몸값", "최고 몸값"],
    pos_mv_rows[:12],
    title="포지션별 몸값 분포 (기록 수 10건 이상)",
    col_widths=[4.0, 2.5, 3.0, 3.0, 3.0]
)

add_heading(doc, "5.5.5  발 선호도 분포", 3, 12)
foot = tm['foot_preference_analysis']['distribution']
total_foot = foot['right'] + foot['left'] + foot['both']
add_table(doc,
    ["발 선호도", "선수-시즌 수", "비율 (유효 기록 기준)"],
    [
        ["오른발 (right)", f"{foot['right']:,}", pct(foot['right']/total_foot*100)],
        ["왼발 (left)", f"{foot['left']:,}", pct(foot['left']/total_foot*100)],
        ["양발 (both)", f"{foot['both']:,}", pct(foot['both']/total_foot*100)],
        ["결측 (nan)", f"{foot['nan']:,}", pct(foot['nan']/tm['summary']['total_rows_combined']*100)],
    ],
    title="발 선호도 분포",
    col_widths=[4.0, 4.0, 5.0]
)

add_heading(doc, "5.5.6  신장 분포", 3, 12)
ht = tm['height_analysis']
add_table(doc,
    ["신장 구간", "선수 수"],
    [[k, f"{v:,}"] for k, v in ht['height_distribution_bins_cm'].items()],
    title="선수 신장 분포 (cm 단위)",
    col_widths=[4.0, 4.0]
)
add_table(doc,
    ["통계", "값"],
    [
        ["평균 신장", f"{ht['mean_m']*100:.1f}cm"],
        ["중앙값 신장", f"{ht['median_m']*100:.1f}cm"],
        ["표준편차", f"{ht['std_m']*100:.1f}cm"],
        ["최소", f"{ht['min_m']*100:.0f}cm"],
        ["최대", f"{ht['max_m']*100:.0f}cm"],
        ["유효 기록 수", f"{ht['records_with_height']:,}"],
        ["결측 수", f"{ht['missing_height_count']:,} ({pct(ht['missing_height_pct'])})"],
    ],
    title="신장 기술 통계",
    col_widths=[5.0, 12.0]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6장. 시대별 비교 분석
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "제 6 장.  시대별 비교 분석 (2000-2012 vs 2013-2025)", 1, 18)

add_heading(doc, "6.1  경기 결과 비교", 2, 14)

# 시대별 결과 계산
era1_seasons = [s for s in per_season.keys() if int(s[:4]) < 2013]
era2_seasons = [s for s in per_season.keys() if int(s[:4]) >= 2013]

era1_hw = sum(per_season[s]['home_wins'] for s in era1_seasons)
era1_d  = sum(per_season[s]['draws'] for s in era1_seasons)
era1_aw = sum(per_season[s]['away_wins'] for s in era1_seasons)
era1_t  = era1_hw + era1_d + era1_aw

era2_hw = sum(per_season[s]['home_wins'] for s in era2_seasons)
era2_d  = sum(per_season[s]['draws'] for s in era2_seasons)
era2_aw = sum(per_season[s]['away_wins'] for s in era2_seasons)
era2_t  = era2_hw + era2_d + era2_aw

era1_goals = sum(goals_trend[s]['total_goals'] for s in era1_seasons)
era2_goals = sum(goals_trend[s]['total_goals'] for s in era2_seasons)
era1_avg_goals = sum(goals_trend[s]['avg_total_goals'] for s in era1_seasons) / len(era1_seasons)
era2_avg_goals = sum(goals_trend[s]['avg_total_goals'] for s in era2_seasons) / len(era2_seasons)

add_table(doc,
    ["지표", "2000-2012 (Era 1)", "2013-2025 (Era 2)", "변화"],
    [
        ["시즌 수", str(len(era1_seasons)), str(len(era2_seasons)), "-"],
        ["총 경기 수", f"{era1_t:,}", f"{era2_t:,}", "-"],
        ["홈 승률", pct(era1_hw/era1_t*100), pct(era2_hw/era2_t*100), f"{(era2_hw/era2_t - era1_hw/era1_t)*100:+.2f}%p"],
        ["무승부율", pct(era1_d/era1_t*100), pct(era2_d/era2_t*100), f"{(era2_d/era2_t - era1_d/era1_t)*100:+.2f}%p"],
        ["원정 승률", pct(era1_aw/era1_t*100), pct(era2_aw/era2_t*100), f"{(era2_aw/era2_t - era1_aw/era1_t)*100:+.2f}%p"],
        ["시즌 평균 경기당 득점", f"{era1_avg_goals:.4f}", f"{era2_avg_goals:.4f}", f"{era2_avg_goals - era1_avg_goals:+.4f}"],
        ["총 득점", f"{era1_goals:,}", f"{era2_goals:,}", f"{era2_goals-era1_goals:+,}"],
    ],
    title="시대별 경기 결과 비교",
    col_widths=[5.0, 3.5, 3.5, 5.0]
)

add_heading(doc, "6.2  데이터 완전성 비교", 2, 14)
add_table(doc,
    ["스탯 유형", "2000-2012", "2013-2025", "비고"],
    [
        ["squad_stats 컬럼 수", "24개", "24개", "동일"],
        ["matchlogs 기본 스탯", "완전 (100%)", "완전 (100%)", "모든 시즌 동일"],
        ["matchlogs 상세 스탯 (sh/sot/fls 등)", "결측 ~48.55%", "완전 (100%)", "2013/14부터 수집"],
        ["matchlogs pkwon/pkcon", "결측 ~76.99%", "부분 (~2015/16부터)", "최근 시즌 수록"],
        ["matchlogs 포지션 (pos)", "결측 ~60%", "결측 ~60%", "전 시즌 공통"],
        ["Transfermarkt 몸값", "2000-2003 미수록", "완전 수록", "2004/05부터 가용"],
        ["Transfermarkt 신장/발", "일부 결측", "일부 결측", "약 4~6% 결측"],
    ],
    title="시대별 데이터 완전성 비교",
    col_widths=[5.0, 3.0, 3.0, 6.0]
)

add_heading(doc, "6.3  전술적 진화 지표", 2, 14)
add_body(doc,
    "원정 승률의 꾸준한 증가 추세와 COVID 시즌의 홈 어드밴티지 소멸은 "
    "현대 EPL의 전술적 진화를 보여주는 중요한 지표입니다. "
    "특히 2020/21 시즌 무관중 환경에서 원정 승률이 40.26%로 홈 승률(37.89%)을 "
    "초과하여, 관중 유무가 홈 어드밴티지에 직접적인 영향을 미침을 실증적으로 증명합니다.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7장. 데이터 통합 분석
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "제 7 장.  데이터 통합 분석", 1, 18)

add_heading(doc, "7.1  파일 간 연결 관계 (Join Keys)", 2, 14)
add_table(doc,
    ["소스 파일", "타겟 파일", "조인 키", "관계"],
    [
        ["match_results.parquet", "player_match_logs.parquet", "season + team + date", "1:N (1경기 : N선수)"],
        ["player_match_stats.parquet", "player_season_stats.parquet", "player + season + team", "N:1 (N경기 → 1시즌 집계)"],
        ["player_season_stats.parquet", "Transfermarkt 원본", "player_normalized + season + team", "1:1 (매칭률 92.7%)"],
        ["match_results.parquet", "team_season_summary.parquet", "Season + HomeTeam/AwayTeam", "N:1 (N경기 → 1팀-시즌)"],
    ],
    title="파일 간 연결 관계",
    col_widths=[4.5, 4.5, 4.5, 4.0]
)

add_heading(doc, "7.2  일관성 검증 결과", 2, 14)
add_table(doc,
    ["검증 항목", "기준", "결과", "비고"],
    [
        ["팀명 일치", "match_results ↔ squad_stats", "통과 (46개 팀 통일)", "통합 매핑 97개 규칙"],
        ["시즌 커버리지", "3개 소스 모두 25시즌 포함", "통과", "2000/01~2024/25"],
        ["경기 수 검증", "시즌당 380경기 기준", "2시즌 예외 (2003/04, 2004/05: 335경기)", "일부 데이터 미포함"],
        ["날짜 범위 일치", "match_results ↔ matchlogs", "일치 (2000-08-19 ~ 2025-05-25)", "일관성 확인"],
        ["결측치 처리", "핵심 식별자 결측 0건", "통과", "team, season, player 결측 없음"],
        ["중복 경기 검증", "동일 경기 중복 없음", "통과 (중복 0건)", "날짜+팀 조합 고유성 확인"],
    ],
    title="데이터 일관성 검증 결과",
    col_widths=[4.0, 4.0, 4.5, 4.5]
)

add_heading(doc, "7.3  파이프라인 압축률", 2, 14)
add_body(doc,
    "원본 데이터에서 전처리된 Parquet 파일로의 압축을 통해 데이터 용량을 "
    "크게 줄이면서도 정보의 완전성을 유지하였습니다.")
add_table(doc,
    ["단계", "파일 수", "총 용량", "압축률"],
    [
        ["원본 Raw CSV", "약 12,394개", "약 78.7MB", "기준"],
        ["처리된 Parquet (processed)", "5개", "약 5.1MB", "약 93.5% 압축"],
        ["대시보드용 Parquet (dashboard)", "2개", "약 0.5MB", "약 99.4% 압축"],
    ],
    title="데이터 파이프라인 압축률",
    col_widths=[5.0, 3.0, 3.0, 6.0]
)

add_heading(doc, "7.4  처리된 데이터 파일 최종 현황", 2, 14)
proc_summary = [
    ["match_results.parquet",
     f"{proc_files['processed/match_results.parquet']['shape']['rows']:,}",
     str(proc_files['processed/match_results.parquet']['shape']['cols']),
     "경기 결과, 팀 단위"],
    ["player_match_logs.parquet",
     f"{proc_files['processed/player_match_logs.parquet']['shape']['rows']:,}",
     str(proc_files['processed/player_match_logs.parquet']['shape']['cols']),
     "선수별 경기 로그 (처리 버전)"],
    ["player_match_stats.parquet",
     f"{proc_files['processed/player_match_stats.parquet']['shape']['rows']:,}",
     str(proc_files['processed/player_match_stats.parquet']['shape']['cols']),
     "선수별 경기 스탯 (원본 컬럼)"],
    ["player_season_stats.parquet",
     f"{proc_files['processed/player_season_stats.parquet']['shape']['rows']:,}",
     str(proc_files['processed/player_season_stats.parquet']['shape']['cols']),
     "FBref+TM 통합 시즌 집계"],
]
add_table(doc,
    ["파일명", "행 수", "컬럼 수", "내용"],
    proc_summary,
    title="전처리 완료 파일 현황",
    col_widths=[5.5, 3.0, 2.5, 6.0]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8장. 향후 계획
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "제 8 장.  향후 계획", 1, 18)

add_heading(doc, "8.1  파생변수 생성 계획", 2, 14)
add_body(doc, "현재 전처리된 데이터를 기반으로 다음 파생변수 생성을 계획합니다:")

deriv_rows = [
    ["폼 지수 (Form Index)", "최근 N경기 결과 기반 가중 폼 점수", "match_results + player_match_logs"],
    ["ELO 레이팅", "상대적 팀 강도 반영 동적 평점 시스템", "match_results"],
    ["롤링 통계 (Rolling Stats)", "3/5/10경기 이동 평균 득점/실점/승률", "match_results + player_match_logs"],
    ["선수 피로도 지수", "누적 출전 분 기반 피로 추정 지표", "player_match_logs"],
    ["예상 득점(xG) 대리 지표", "슈팅 수 × 유효슈팅 비율 기반", "player_match_stats"],
    ["시장가치 성장률", "시즌 간 몸값 변화율", "player_season_stats (TM)"],
    ["홈/원정 퍼포먼스 차이", "선수별 홈 vs 원정 스탯 비율", "player_match_logs"],
]
add_table(doc,
    ["파생변수", "정의", "소스 파일"],
    deriv_rows,
    title="계획된 파생변수 목록",
    col_widths=[4.5, 6.0, 6.5]
)

add_heading(doc, "8.2  딥러닝 모델링 계획 (P1~P5)", 2, 14)
model_rows = [
    ["P1", "경기 결과 예측 (홈/무/원정)", "match_results + 파생변수", "XGBoost, LSTM"],
    ["P2", "득점 예측 (팀/선수)", "match_results + player_match_logs", "Random Forest, Neural Net"],
    ["P3", "선수 부상 위험도 예측", "player_match_logs (누적 피로)", "이상치 감지 + 분류 모델"],
    ["P4", "선수 시장가치 예측", "player_season_stats + TM 통합", "XGBoost, TabNet"],
    ["P5", "스카우팅 추천 시스템", "전체 통합 데이터", "Collaborative Filtering, 유사도 기반"],
]
add_table(doc,
    ["단계", "모델 목적", "사용 데이터", "알고리즘 후보"],
    model_rows,
    title="딥러닝/ML 모델링 로드맵",
    col_widths=[1.5, 5.0, 5.0, 5.5]
)

add_heading(doc, "8.3  대시보드 개발 계획", 2, 14)
dash_rows = [
    ["팀 성과 대시보드", "시즌별 승률, 득실점, 홈/원정 어드밴티지 시각화"],
    ["선수 프로파일 뷰어", "포지션별 스탯 레이더 차트, 시즌 트렌드"],
    ["이적 시장 분석 뷰", "포지션별 몸값 분포, 국적별 몸값 맵"],
    ["경기 결과 예측기", "P1 모델 연동 실시간 예측 인터페이스"],
    ["스카우팅 필터", "연령/포지션/국적/몸값/스탯 다차원 필터링"],
]
add_table(doc,
    ["대시보드 모듈", "주요 기능"],
    dash_rows,
    title="대시보드 모듈 계획",
    col_widths=[5.0, 12.0]
)

add_body(doc,
    "대시보드는 Streamlit 또는 Dash 프레임워크를 활용하여 구현하며, "
    "dashboard/ 폴더의 Parquet 파일(약 0.5MB)을 직접 로드하여 빠른 응답 속도를 구현합니다.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9장. 결론
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "제 9 장.  결론", 1, 18)

add_body(doc,
    "본 보고서는 EPL 2000/01 ~ 2024/25 시즌(25시즌)에 걸친 방대한 데이터를 체계적으로 "
    "수집·전처리·분석한 결과를 종합 정리하였습니다. 주요 성과와 발견 사항은 다음과 같습니다.")

conclusions = [
    "1. 데이터 규모: 총 9,380경기, 46개 팀, 3,209명 선수의 기록을 12,394개 원본 CSV에서 5개 Parquet으로 정제",
    f"2. 데이터 품질: 매치 결과 데이터는 결측치 0건(100% 완전), 팀명 통합 97개 매핑으로 소스 간 일관성 확보",
    f"3. 홈 어드밴티지: 25시즌 통산 홈 승률 {pct(rd['home_win_pct'])}로 원정({pct(rd['away_win_pct'])}) 대비 우세, "
    "단 2020/21 COVID 무관중 시즌에서 이례적 역전 현상 관찰",
    f"4. 득점 트렌드: 경기당 평균 득점 {ga['avg_total_goals']:.3f}골, 25시즌 총 {ga['total_goals_all']:,}골 기록. "
    "2023/24 시즌 경기당 3.28골로 최고점",
    f"5. 선수 몸값: 2004/05 평균 €3.6M에서 2024/25 평균 €17.6M으로 약 4.8배 상승. "
    "역대 최고 몸값은 Erling Haaland €180M (2022/23~2024/25)",
    "6. 시대별 차이: 2013/14 이후 상세 스탯(슈팅, 파울, 크로스 등) 수집 가능, "
    "초기 시즌(2000-2012)은 기본 스탯만 완전 수록",
    "7. 통합 품질: FBref + Transfermarkt 선수 매칭률 92.7%, 통합 데이터 14,980행 × 36컬럼 생성",
    "8. 압축 효율: raw 78.7MB → processed 5.1MB(93.5% 압축) → dashboard 0.5MB(99.4% 압축)",
]
for c in conclusions:
    add_bullet(doc, c)

doc.add_paragraph()
add_body(doc,
    "구축된 데이터 파이프라인과 분석 기반은 스카우터·이적 담당자·코치/감독이 "
    "데이터 기반 의사결정을 내릴 수 있는 견고한 토대를 제공합니다. "
    "향후 파생변수 생성과 딥러닝 모델링(P1~P5)을 통해 EPL 경기 예측, "
    "선수 가치 평가, 스카우팅 추천 시스템으로 발전시킬 계획입니다.")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─" * 50)
set_font(run, 10, color=(128,128,128))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EPL 데이터 프로젝트 전처리 및 EDA 종합 보고서  |  분석일: 2026년 3월 21일")
set_font(run, 9, color=(128,128,128))

# ══════════════════════════════════════════════════════════════════════════════
# 저장
# ══════════════════════════════════════════════════════════════════════════════
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(f"✅ 보고서 저장 완료: {OUT_PATH}")
print(f"   파일 크기: {os.path.getsize(OUT_PATH):,} bytes ({os.path.getsize(OUT_PATH)/1024:.1f} KB)")
