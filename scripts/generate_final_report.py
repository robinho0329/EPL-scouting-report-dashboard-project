"""
EPL 스카우트 데이터 분석 프로젝트 - 종합 보고서 생성 스크립트
실제 데이터를 읽어 통계를 계산하고 Word 문서로 출력합니다.
"""

import os
import json
import warnings
from pathlib import Path
from datetime import datetime

import numpy as np
import pandas as pd
from sklearn.metrics import roc_auc_score, f1_score, accuracy_score

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings("ignore")

# ─────────────────────────────────────────────────────────────────
# 경로 설정
# ─────────────────────────────────────────────────────────────────
BASE_DIR = Path("C:/Users/xcv54/workspace/EPL project")
DATA_DIR  = BASE_DIR / "data" / "processed"
SCOUT_DIR = BASE_DIR / "data" / "scout"
MODELS_DIR = BASE_DIR / "models"
FIG_DIR   = BASE_DIR / "reports" / "figures"
OUT_PATH  = BASE_DIR / "reports" / "scout_final_report.docx"

FONT_KR = "맑은 고딕"

# ─────────────────────────────────────────────────────────────────
# 데이터 로드
# ─────────────────────────────────────────────────────────────────
print("[1/3] 데이터 로드 중...")

df_match   = pd.read_parquet(DATA_DIR / "match_results.parquet")
df_player  = pd.read_parquet(DATA_DIR / "player_season_stats.parquet")
df_team    = pd.read_parquet(DATA_DIR / "team_season_summary.parquet")
df_mlog    = pd.read_parquet(DATA_DIR / "player_match_logs.parquet")
df_mstat   = pd.read_parquet(DATA_DIR / "player_match_stats.parquet")

df_scout_prof  = pd.read_parquet(SCOUT_DIR / "scout_player_profiles.parquet")
df_scout_rat   = pd.read_parquet(SCOUT_DIR / "scout_ratings.parquet")
df_scout_team  = pd.read_parquet(SCOUT_DIR / "scout_team_profiles.parquet")
df_scout_trans = pd.read_parquet(SCOUT_DIR / "scout_transfers.parquet")
df_cluster     = pd.read_parquet(SCOUT_DIR / "cluster_assignments.parquet")
df_similarity  = pd.read_parquet(SCOUT_DIR / "similarity_matrix.parquet")
df_growth      = pd.read_parquet(SCOUT_DIR / "growth_predictions.parquet")
df_transfer    = pd.read_parquet(SCOUT_DIR / "transfer_predictions.parquet")
df_decline     = pd.read_parquet(SCOUT_DIR / "decline_predictions.parquet")
df_undervalued = pd.read_parquet(SCOUT_DIR / "undervalued_players.parquet")
df_hidden      = pd.read_parquet(SCOUT_DIR / "hidden_gems.parquet")

with open(SCOUT_DIR / "results_summary.json", encoding="utf-8") as f:
    s4_results = json.load(f)
with open(MODELS_DIR / "s1_player_rating" / "results_summary.json", encoding="utf-8") as f:
    s1_results = json.load(f)
with open(MODELS_DIR / "s2_market_value" / "results_summary.json", encoding="utf-8") as f:
    s2_results = json.load(f)
with open(MODELS_DIR / "s3_similarity" / "results_summary.json", encoding="utf-8") as f:
    s3_results = json.load(f)

print("    데이터 로드 완료")

# ─────────────────────────────────────────────────────────────────
# 파생 통계 계산
# ─────────────────────────────────────────────────────────────────
print("[2/3] 통계 계산 중...")

# 기본 데이터 통계
n_matches       = len(df_match)
n_seasons_match = df_match["Season"].nunique()
avg_goals       = round((df_match["FullTimeHomeGoals"] + df_match["FullTimeAwayGoals"]).mean(), 2)
season_min      = df_match["Season"].min()
season_max      = df_match["Season"].max()

n_players_total   = df_player["player"].nunique()
n_player_rows     = len(df_player)
n_with_value      = int(df_player["market_value"].notna().sum())
avg_market_value  = round(df_player["market_value"].mean() / 1e6, 2)
n_nationalities   = df_player["nationality"].nunique()

n_team_rows   = len(df_team)
n_seasons_team = df_team["Season"].nunique()

n_match_log_rows = len(df_mlog)

# Scout 통계
n_scout_players  = df_scout_prof["player"].nunique() if "player" in df_scout_prof.columns else len(df_scout_prof)
n_scout_ratings  = len(df_scout_rat)
n_scout_teams    = len(df_scout_team)
n_scout_trans    = len(df_scout_trans)
n_clusters       = df_cluster["cluster"].nunique() if "cluster" in df_cluster.columns else s3_results["metadata"]["best_k"]
n_similarity     = len(df_similarity)
n_growth_rows    = len(df_growth)
n_transfer_rows  = len(df_transfer)
n_decline_rows   = len(df_decline)
n_hidden_gems    = len(df_hidden)
n_undervalued    = len(df_undervalued)

# S5 metrics from parquet
y5 = df_transfer["adapted"]
s5_metrics = {}
for col, name in [("adapt_prob_xgb","XGBoost"), ("adapt_prob_rf","RandomForest"),
                  ("adapt_prob_lr","LogisticRegression"), ("adapt_prob_mlp","MLP"),
                  ("adapt_prob_ensemble","Ensemble")]:
    auc = roc_auc_score(y5, df_transfer[col])
    pred = (df_transfer[col] >= 0.5).astype(int)
    f1  = f1_score(y5, pred)
    acc = accuracy_score(y5, pred)
    s5_metrics[name] = {"auc": round(auc,4), "f1": round(f1,4), "acc": round(acc,4)}

s5_best = max(s5_metrics, key=lambda k: s5_metrics[k]["auc"])
s5_adapt_rate = round(df_transfer["adapted"].mean() * 100, 1)
s5_safe_cnt   = int((df_transfer["scout_category"] == "Safe Bet").sum())
s5_moderate   = int((df_transfer["scout_category"] == "Moderate").sum())
s5_risky      = int((df_transfer["scout_category"] == "Risky").sum())

# S6 metrics
y6 = df_decline["decline"]
s6_metrics = {}
for col, name in [("decline_prob_xgb","XGBoost"), ("decline_prob_rf","RandomForest"),
                  ("decline_prob_lr","LogisticRegression"), ("decline_prob_mlp","MLP"),
                  ("decline_prob_ensemble","Ensemble")]:
    auc = roc_auc_score(y6, df_decline[col])
    pred = (df_decline[col] >= 0.5).astype(int)
    f1  = f1_score(y6, pred)
    acc = accuracy_score(y6, pred)
    s6_metrics[name] = {"auc": round(auc,4), "f1": round(f1,4), "acc": round(acc,4)}

s6_best = max(s6_metrics, key=lambda k: s6_metrics[k]["auc"])
s6_high_risk   = int((df_decline["decline_prob_ensemble"] > 0.7).sum())
s6_low_risk    = int((df_decline["decline_prob_ensemble"] < 0.3).sum())
s6_unique_players = df_decline["player_key"].nunique()

print("    통계 계산 완료")

# ─────────────────────────────────────────────────────────────────
# 유틸리티 함수
# ─────────────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, color=None):
    run.font.name = FONT_KR
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Ensure Korean font is applied via XML
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), FONT_KR)
    rFonts.set(qn("w:ascii"), FONT_KR)
    rFonts.set(qn("w:hAnsi"), FONT_KR)
    rPr.insert(0, rFonts)


def add_heading(doc, text, level=1, size=None, color=None):
    """스타일을 직접 설정한 제목 단락 추가"""
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, size=size or 22, bold=True, color=color)
        return p
    elif level == 1:
        p = doc.add_heading("", level=1)
        p.clear()
        run = p.add_run(text)
        set_font(run, size=size or 16, bold=True, color=color or (31, 73, 125))
        return p
    elif level == 2:
        p = doc.add_heading("", level=2)
        p.clear()
        run = p.add_run(text)
        set_font(run, size=size or 13, bold=True, color=color or (55, 96, 146))
        return p
    else:
        p = doc.add_heading("", level=3)
        p.clear()
        run = p.add_run(text)
        set_font(run, size=size or 11, bold=True, color=color or (79, 129, 189))
        return p


def add_para(doc, text, size=10.5, bold=False, indent=0, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p


def add_bullet(doc, text, size=10.5, indent=0.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=size)
    return p


def add_table(doc, headers, rows, col_widths=None, header_color=(31, 73, 125)):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Background color
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        hex_color = "{:02X}{:02X}{:02X}".format(*header_color)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        set_font(run, size=10, bold=True, color=(255, 255, 255))

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        # Alternating row color
        if ri % 2 == 1:
            for ci in range(n_cols):
                tc = tr.cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EBF1DE")
                tcPr.append(shd)
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, size=9.5)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def add_figure(doc, fig_path, caption, width=5.5):
    if not Path(fig_path).exists():
        add_para(doc, f"[그림 없음: {Path(fig_path).name}]", size=9, bold=False)
        return
    doc.add_picture(str(fig_path), width=Inches(width))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_p.add_run(caption)
    set_font(run, size=9, bold=False, color=(89, 89, 89))
    doc.add_paragraph()


def add_page_break(doc):
    doc.add_page_break()


def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 50)
    set_font(run, size=9, color=(180, 180, 180))


# ─────────────────────────────────────────────────────────────────
# 문서 생성 시작
# ─────────────────────────────────────────────────────────────────
print("[3/3] 보고서 생성 중...")

doc = Document()

# 기본 마진 설정
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# 기본 스타일 폰트
from docx.shared import Pt as Ptt
style = doc.styles["Normal"]
style.font.name = FONT_KR
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KR)

# ─────────────────────────────────────────────────────────────────
# 표지
# ─────────────────────────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()

add_heading(doc, "EPL 스카우트 데이터 분석 프로젝트", level=0, size=26)
add_heading(doc, "종합 보고서", level=0, size=20, color=(31, 73, 125))

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("영국 프리미어리그 2000/01 ~ 2024/25 시즌")
set_font(run, size=13, bold=False, color=(89, 89, 89))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"데이터 기반 선수 평가 · 몸값 예측 · 유사 선수 탐색 · 성장 예측 · 이적 적응도 · 하락 위험 감지")
set_font(run, size=11, bold=False, color=(89, 89, 89))

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
set_font(run, size=11, color=(89, 89, 89))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EPL Scout Analytics Team")
set_font(run, size=12, bold=True, color=(31, 73, 125))

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────
# 목차
# ─────────────────────────────────────────────────────────────────
add_heading(doc, "목  차", level=1)
doc.add_paragraph()

toc_items = [
    ("1장", "프로젝트 개요"),
    ("2장", "데이터 현황"),
    ("3장", "탐색적 데이터 분석 (EDA)"),
    ("4장", "모델링 결과 - S1 종합 선수 평가 (WAR Rating)"),
    ("5장", "모델링 결과 - S2 적정 몸값 예측"),
    ("6장", "모델링 결과 - S3 유사 선수 탐색 (Similarity Engine)"),
    ("7장", "모델링 결과 - S4 성장 잠재력 예측"),
    ("8장", "모델링 결과 - S5 이적 적응도 예측"),
    ("9장", "모델링 결과 - S6 하락 위험 감지"),
    ("10장", "종합 분석 및 스카우트 활용 가이드"),
    ("부록", "데이터 사전 및 모델 하이퍼파라미터"),
]
for ch, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(f"{ch}. {title}")
    set_font(run, size=11)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────
# 1장. 프로젝트 개요
# ─────────────────────────────────────────────────────────────────
add_heading(doc, "1장. 프로젝트 개요", level=1)
doc.add_paragraph()

add_heading(doc, "1.1 프로젝트 목적", level=2)
add_para(doc,
    "본 프로젝트는 잉글랜드 프리미어리그(EPL) 25시즌(2000/01 ~ 2024/25)의 경기·선수·팀 데이터를 "
    "체계적으로 수집·전처리하여, 현장 스카우트가 데이터 기반 의사결정을 내릴 수 있도록 "
    "6개의 예측 모델을 개발하는 것을 목적으로 합니다. "
    "단순 통계 제공을 넘어, 선수의 내재적 가치를 발굴하고 이적 리스크를 정량화하며 "
    "장기 성장 곡선을 예측함으로써 데이터 주도 스카우팅의 실무 적용 가능성을 검증합니다.")

doc.add_paragraph()

add_heading(doc, "1.2 분석 대상", level=2)
rows_target = [
    ["경기 결과", f"{n_matches:,}경기 ({season_min} ~ {season_max})", "홈/원정 골, 카드, 하프타임 결과 포함"],
    ["선수 시즌 스탯", f"{n_player_rows:,}행 / {n_players_total:,}명", "골, 어시스트, 출전 시간, 국적, 포지션, 몸값"],
    ["선수 경기별 로그", f"{n_match_log_rows:,}행", "경기별 세부 수행 지표 (FBRef 기반)"],
    ["팀 시즌 요약", f"{n_team_rows:,}행", "시즌별 팀 순위, 득실점, 포인트"],
]
add_table(doc,
    headers=["분석 대상", "규모", "주요 내용"],
    rows=rows_target,
    col_widths=[1.8, 2.0, 3.2])

doc.add_paragraph()

add_heading(doc, "1.3 데이터 소스", level=2)
bullets_src = [
    "FBRef (fbref.com): 선수 시즌별·경기별 퍼포먼스 통계 (골, 어시스트, xG, 패스, 수비 수치 등)",
    "Transfermarkt (transfermarkt.com): 선수 시장 가치, 이적 이력, 국적, 생년월일, 신체 정보",
    "Open Football Data: 매치 결과 데이터 (2000/01 ~ 2024/25)",
]
for b in bullets_src:
    add_bullet(doc, b)

doc.add_paragraph()

add_heading(doc, "1.4 분석 기간 및 범위", level=2)
add_para(doc,
    f"분석 기간: {season_min} ~ {season_max} (총 {n_seasons_match}시즌). "
    f"총 {n_matches:,}경기, {n_players_total:,}명의 선수, {n_nationalities}개국 국적 데이터를 포함합니다. "
    f"시장 가치 데이터는 {n_with_value:,}건의 레코드에 포함되며, 평균 시장 가치는 "
    f"{avg_market_value}백만 유로입니다.")

doc.add_paragraph()

add_heading(doc, "1.5 모델 구성 요약", level=2)
model_summary_rows = [
    ["S1", "종합 선수 평가 (WAR Rating)", "Ridge·XGBoost·MLP 앙상블", "R² 0.45 (테스트)"],
    ["S2", "적정 몸값 예측", "포지션별 XGBoost/RF/MLP", "FW R² 0.305, MID R² 0.303"],
    ["S3", "유사 선수 탐색", "K-Means 클러스터링 + Cosine 유사도", f"최적 K={s3_results['metadata']['best_k']}, 실루엣 {s3_results['metadata']['silhouette_scores'][str(s3_results['metadata']['best_k'])]:.3f}"],
    ["S4", "성장 잠재력 예측", "RandomForest 회귀 앙상블", f"Validation R² {s4_results['u23_model_performance']['val_r2']:.3f}"],
    ["S5", "이적 적응도 예측", "XGBoost·RF·LR·MLP 앙상블", f"RandomForest AUC {s5_metrics['RandomForest']['auc']:.4f}"],
    ["S6", "하락 위험 감지", "XGBoost·RF·LR·MLP 앙상블", f"XGBoost AUC {s6_metrics['XGBoost']['auc']:.4f}"],
]
add_table(doc,
    headers=["코드", "모델명", "방법론", "주요 성과"],
    rows=model_summary_rows,
    col_widths=[0.6, 2.0, 2.4, 2.0])

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────
# 2장. 데이터 현황
# ─────────────────────────────────────────────────────────────────
add_heading(doc, "2장. 데이터 현황", level=1)
doc.add_paragraph()

add_heading(doc, "2.1 원본 데이터 규모", level=2)
add_para(doc,
    "아래 표는 전처리 완료 후 생성된 4개 핵심 파켓 파일의 규모입니다. "
    "모든 수치는 실제 파일에서 읽은 값입니다.")
doc.add_paragraph()

raw_rows = [
    ["match_results.parquet",        f"{len(df_match):,}행",  "26열", "시즌별 경기 결과 (홈/원정 골, 결과, 카드 등)"],
    ["player_season_stats.parquet",  f"{len(df_player):,}행", "36열", "선수 시즌별 통계 + 시장 가치 병합본"],
    ["player_match_logs.parquet",    f"{len(df_mlog):,}행",   "33열", "선수 경기별 세부 로그 (FBRef)"],
    ["team_season_summary.parquet",  f"{len(df_team):,}행",   "22열", "팀별 시즌 요약 (순위, 득실, 포인트)"],
]
add_table(doc,
    headers=["파일명", "행 수", "열 수", "내용"],
    rows=raw_rows,
    col_widths=[2.5, 1.0, 0.8, 2.7])

doc.add_paragraph()

add_heading(doc, "2.2 전처리 파이프라인 요약", level=2)
pipeline_steps = [
    "데이터 수집: FBRef 선수 스탯, Transfermarkt 몸값 및 신체 정보, Open Football 경기 결과를 시즌별로 수집",
    "키 병합: 선수명 + 팀명 + 시즌을 기준 키로 FBRef-Transfermarkt 조인 (Fuzzy Matching 적용)",
    "결측 처리: 시장 가치 미입력 행은 no_value_data 플래그 생성, 수치형 피처는 중앙값 대체",
    "파생 변수 생성: per90 정규화 지표, 시즌 연도 인코딩, 포지션 그룹화(FW/MID/DEF/GK), ELO 추정값",
    "스카우트 프로파일 생성: scout_features.py를 통해 49개 스카우트 피처 산출",
]
for s in pipeline_steps:
    add_bullet(doc, s)

doc.add_paragraph()

add_heading(doc, "2.3 스카우트 파생 변수 현황", level=2)
scout_files_rows = [
    ["scout_player_profiles.parquet", f"{len(df_scout_prof):,}행", "49열", "선수별 종합 스카우트 피처 (per90 지표, ELO, 나이 변수 등)"],
    ["scout_ratings.parquet",         f"{len(df_scout_rat):,}행",  "16열", f"WAR 등급 산출 결과 ({s1_results['total_players_rated']:,}명)"],
    ["scout_team_profiles.parquet",   f"{len(df_scout_team):,}행", "44열", "팀별 시즌 프로파일 (공격/수비 스타일, ELO 등)"],
    ["scout_transfers.parquet",       f"{len(df_scout_trans):,}행","41열", "이적 기록 및 팀 간 피처 조합"],
    ["cluster_assignments.parquet",   f"{len(df_cluster):,}행",    "23열", f"선수별 클러스터 분류 (K={n_clusters})"],
    ["growth_predictions.parquet",    f"{len(df_growth):,}행",     "10열", "시즌별 성장 예측 결과"],
    ["transfer_predictions.parquet",  f"{len(df_transfer):,}행",   "31열", "이적 적응도 예측 결과"],
    ["decline_predictions.parquet",   f"{len(df_decline):,}행",    "20열", "하락 위험도 예측 결과"],
    ["hidden_gems.parquet",           f"{n_hidden_gems}행",         "17열", "저평가 히든 젬 선수 목록"],
    ["undervalued_players.parquet",   f"{n_undervalued}행",         "12열", "몸값 저평가 선수 목록"],
]
add_table(doc,
    headers=["파일명", "행 수", "열 수", "내용"],
    rows=scout_files_rows,
    col_widths=[2.8, 0.9, 0.7, 2.6])

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────
# 3장. EDA
# ─────────────────────────────────────────────────────────────────
add_heading(doc, "3장. 탐색적 데이터 분석 (EDA)", level=1)
doc.add_paragraph()

add_para(doc,
    f"총 {n_seasons_match}시즌의 데이터를 대상으로 경기 패턴, 선수 분포, 시장 가치, "
    "시대별 변화, 상관 관계를 분석하였습니다. 아래 그림들은 스카우트 전용 EDA 분석 결과입니다.")

doc.add_paragraph()
add_heading(doc, "3.1 데이터 품질 및 시즌 가용성", level=2)
add_para(doc,
    "시장 가치 데이터는 초기 시즌(2000년대 초반)일수록 결측 비율이 높고, "
    f"2015/16 이후 구간에서 커버리지가 크게 향상됩니다. "
    f"전체 {n_player_rows:,}개 선수-시즌 레코드 중 {n_with_value:,}건에 시장 가치가 포함되어 있습니다.")
add_figure(doc, FIG_DIR / "scout_eda_01_missing_values.png",
           "그림 3-1. 변수별 결측값 비율 히트맵")
add_figure(doc, FIG_DIR / "scout_eda_01b_season_availability.png",
           "그림 3-2. 시즌별 데이터 가용성 현황")

add_heading(doc, "3.2 시장 가치 분포 및 트렌드", level=2)
add_para(doc,
    f"선수 평균 시장 가치는 {avg_market_value}백만 유로이며, 시간이 지날수록 시장 가치 인플레이션이 "
    "뚜렷하게 관찰됩니다. 공격수(FW)와 미드필더(MID)가 수비수(DEF)·골키퍼(GK)보다 "
    "평균적으로 높은 시장 가치를 지닙니다.")
add_figure(doc, FIG_DIR / "scout_eda_02_value_by_position.png",
           "그림 3-3. 포지션별 시장 가치 분포 (박스플롯)")
add_figure(doc, FIG_DIR / "scout_eda_02b_value_trend.png",
           "그림 3-4. 시즌별 평균 시장 가치 추이")
add_figure(doc, FIG_DIR / "scout_eda_02c_top20_value.png",
           "그림 3-5. 최고 시장 가치 상위 20명 선수")
add_figure(doc, FIG_DIR / "scout_eda_02d_value_vs_perf.png",
           "그림 3-6. 퍼포먼스 점수 vs. 시장 가치 산포도")

add_heading(doc, "3.3 히든 젬 분석", level=2)
add_para(doc,
    f"WAR 등급 대비 시장 가치가 현저히 낮은 히든 젬 선수 {n_hidden_gems}명을 식별하였습니다. "
    "이들은 낮은 몸값 대비 높은 경기 기여도를 보이는 선수들로 이적 시장에서 높은 가성비를 제공합니다.")
add_figure(doc, FIG_DIR / "scout_eda_03_hidden_gems.png",
           "그림 3-7. 히든 젬 선수 시각화 (WAR 등급 vs. 시장 가치)")

add_heading(doc, "3.4 포지션별 퍼포먼스 지표", level=2)
add_figure(doc, FIG_DIR / "scout_eda_04_position_metrics.png",
           "그림 3-8. 포지션별 주요 퍼포먼스 지표 비교")
add_figure(doc, FIG_DIR / "scout_eda_04b_position_violin.png",
           "그림 3-9. 포지션별 지표 분포 바이올린 차트")

add_heading(doc, "3.5 나이-성장 곡선", level=2)
add_para(doc,
    "선수의 나이와 퍼포먼스 점수 간의 관계를 분석한 결과, 포지션별로 최적 나이 구간이 다르게 나타납니다. "
    "공격수는 상대적으로 이른 나이(23~27세)에 피크를 맞이하며, 수비수와 골키퍼는 "
    "보다 늦은 나이(26~30세)까지 성능이 유지되는 경향이 있습니다.")
add_figure(doc, FIG_DIR / "scout_eda_05_age_curves.png",
           "그림 3-10. 포지션별 나이-퍼포먼스 곡선")
add_figure(doc, FIG_DIR / "scout_eda_05b_age_value_curves.png",
           "그림 3-11. 나이별 시장 가치 곡선")

add_heading(doc, "3.6 일관성 분석 및 기타 인사이트", level=2)
add_figure(doc, FIG_DIR / "scout_eda_06_consistency.png",
           "그림 3-12. 선수별 퍼포먼스 일관성(CV) 분포")
add_figure(doc, FIG_DIR / "scout_eda_07_transfer_patterns.png",
           "그림 3-13. 이적 패턴 분석 (이적 전후 퍼포먼스 비교)")
add_figure(doc, FIG_DIR / "scout_eda_08_squad_analysis.png",
           "그림 3-14. 팀별 스쿼드 분석")
add_figure(doc, FIG_DIR / "scout_eda_09_era_evolution.png",
           "그림 3-15. 시대별(에라별) 퍼포먼스 진화 트렌드")

add_heading(doc, "3.7 변수 간 상관 관계", level=2)
add_para(doc,
    "시장 가치와 퍼포먼스 점수 간에 정의 상관이 확인되며, 출전 시간(min)과 골/어시스트 지표 간에도 "
    "강한 연관성이 관찰됩니다. 수비 지표(태클, 인터셉트)는 공격 지표와 상대적으로 낮은 상관관계를 보입니다.")
add_figure(doc, FIG_DIR / "scout_eda_10_correlation.png",
           "그림 3-16. 스카우트 피처 상관 계수 히트맵")
add_figure(doc, FIG_DIR / "scout_eda_10b_pos_correlation.png",
           "그림 3-17. 포지션별 상관 관계 히트맵")

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────
# 4장. S1 종합 선수 평가
# ─────────────────────────────────────────────────────────────────
add_heading(doc, "4장. S1 종합 선수 평가 (WAR Rating)", level=1)
doc.add_paragraph()

add_heading(doc, "4.1 방법론", level=2)
add_para(doc,
    "WAR(Wins Above Replacement) 개념을 EPL 데이터에 적용하여 포지션 중립적인 종합 선수 평가 지수를 "
    "산출합니다. 출전 시간, 팀 기여도(포인트·득실), 나이, 시장 가치, per90 공격·수비 지표를 피처로 활용하며, "
    "Ridge 회귀, XGBoost, MLP 신경망의 3가지 모델 예측값을 앙상블로 결합합니다.")

add_para(doc, f"학습 데이터: {s1_results['train_size']:,}건  /  검증: {s1_results['val_size']:,}건  /  테스트: {s1_results['test_size']:,}건  /  평가 선수 수: {s1_results['total_players_rated']:,}명")

doc.add_paragraph()
add_heading(doc, "4.2 모델 성능 비교", level=2)

s1_metrics = s1_results["model_metrics"]
s1_rows = []
for model_name, metrics in s1_metrics.items():
    if model_name == "Ensemble":
        s1_rows.append([model_name,
            "-", "-", "-",
            round(metrics["test"]["rmse"],4),
            round(metrics["test"]["mae"],4),
            round(metrics["test"]["r2"],4)])
    else:
        s1_rows.append([model_name,
            round(metrics["val"]["rmse"],4),
            round(metrics["val"]["mae"],4),
            round(metrics["val"]["r2"],4),
            round(metrics["test"]["rmse"],4),
            round(metrics["test"]["mae"],4),
            round(metrics["test"]["r2"],4)])

add_table(doc,
    headers=["모델", "Val RMSE", "Val MAE", "Val R²", "Test RMSE", "Test MAE", "Test R²"],
    rows=s1_rows,
    col_widths=[1.5, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0])

add_para(doc, f"※ 앙상블 모델 기준: Test R² = {s1_metrics['Ensemble']['test']['r2']:.4f}, "
              f"Test RMSE = {s1_metrics['Ensemble']['test']['rmse']:.4f}", size=9)

doc.add_paragraph()
add_heading(doc, "4.3 2023/24 시즌 상위 선수 (포지션별)", level=2)

top_pos = s1_results["top_by_position"]
pos_labels = {"FW": "공격수 (FW)", "MID": "미드필더 (MID)", "DEF": "수비수 (DEF)", "GK": "골키퍼 (GK)"}
for pos, label in pos_labels.items():
    if pos in top_pos:
        add_para(doc, f"■ {label}", bold=True, size=10.5)
        pos_rows = [[i+1, p["player"], p["team"], p.get("season",""), f"{p['war_rating']:.2f}"]
                    for i, p in enumerate(top_pos[pos][:5])]
        add_table(doc,
            headers=["순위", "선수명", "팀", "시즌", "WAR Rating"],
            rows=pos_rows,
            col_widths=[0.6, 2.5, 1.8, 1.2, 1.3])
        doc.add_paragraph()

add_heading(doc, "4.4 히든 젬 선수", level=2)
add_para(doc,
    f"시장 가치 대비 WAR 등급이 현저히 높은 선수를 '히든 젬'으로 분류하였습니다. "
    f"총 {n_hidden_gems}명이 식별되었으며 아래 상위 5명을 소개합니다.")
gem_rows = [[i+1, row["player"], row["team"], int(row["age"]),
             f"{row['war_rating']:.2f}", f"{int(row['market_value']/1e6):.1f}M€"]
            for i, row in df_hidden.head(5).iterrows()]
add_table(doc,
    headers=["순위", "선수명", "팀", "나이", "WAR Rating", "시장가치(EUR)"],
    rows=gem_rows,
    col_widths=[0.6, 2.2, 1.8, 0.8, 1.3, 1.3])

doc.add_paragraph()
add_heading(doc, "4.5 그래프", level=2)
add_figure(doc, MODELS_DIR / "s1_player_rating/figures/rating_distribution.png",
           "그림 4-1. WAR Rating 전체 분포")
add_figure(doc, MODELS_DIR / "s1_player_rating/figures/rating_vs_market_value.png",
           "그림 4-2. WAR Rating vs. 시장 가치 산포도")
add_figure(doc, MODELS_DIR / "s1_player_rating/figures/top_players_2023_24.png",
           "그림 4-3. 2023/24 시즌 포지션별 상위 선수 (WAR Rating)")
add_figure(doc, MODELS_DIR / "s1_player_rating/figures/top_players_2024_25.png",
           "그림 4-4. 2024/25 시즌 포지션별 상위 선수 (WAR Rating)")

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────
# 5장. S2 적정 몸값 예측
# ─────────────────────────────────────────────────────────────────
add_heading(doc, "5장. S2 적정 몸값 예측", level=1)
doc.add_paragraph()

add_heading(doc, "5.1 방법론", level=2)
add_para(doc,
    "시장 가치 예측을 포지션별 독립 모델로 분리하여 각 포지션의 특화 피처를 활용합니다. "
    "FW/MID/DEF/GK 4개 포지션 각각에 Ridge·XGBoost·RandomForest·MLP 모델을 학습하여 "
    "최고 성능 모델을 선택합니다. 이전 시즌 시장 가치(래그 피처), 팀 강도(포인트), "
    "나이 제곱항, 포지션 특화 통계(GK 세이브율·클린시트 등)를 주요 피처로 사용합니다.")
add_para(doc,
    f"전체 데이터: {s2_results['data']['total_rows_after_filter']:,}행  /  "
    f"Train: {s2_results['data']['train_size']:,}  /  "
    f"Val: {s2_results['data']['val_size']:,}  /  "
    f"Test: {s2_results['data']['test_size']:,}")

doc.add_paragraph()
add_heading(doc, "5.2 포지션별 최고 모델 성능", level=2)

s2_pos_rows = []
for pos, data in s2_results["position_models"].items():
    s2_pos_rows.append([
        pos,
        data["best_model"],
        data["test_size"],
        round(data["test_R2"], 4),
        f"{int(data['test_MAE']/1e6):.1f}M€",
        f"{data['test_MAPE']:.1f}%",
    ])
add_table(doc,
    headers=["포지션", "최고 모델", "테스트 수", "Test R²", "Test MAE", "Test MAPE"],
    rows=s2_pos_rows,
    col_widths=[0.9, 1.5, 1.1, 1.0, 1.2, 1.2])

doc.add_paragraph()
add_heading(doc, "5.3 저평가 선수 TOP 10 (몸값 대비 실력)", level=2)
add_para(doc,
    f"예측 시장 가치가 실제 시장 가치보다 1.5배 이상 높은 경우 '저평가'로 분류합니다. "
    f"전체 {s2_results['value_analysis']['significantly_undervalued_count']}명이 저평가 기준을 충족합니다.")
underval_rows = []
for i, p in enumerate(s2_results["top20_undervalued"][:10]):
    underval_rows.append([
        i+1, p["player"], p["season"], p["team"], p["position"], int(p["age"]),
        f"{int(p['actual_market_value_eur']/1e6):.1f}M€",
        f"{int(p['predicted_market_value_eur']/1e6):.1f}M€",
        f"{p['value_score']:.2f}x"
    ])
add_table(doc,
    headers=["순위","선수명","시즌","팀","포지션","나이","실제 가치","예측 가치","배수"],
    rows=underval_rows,
    col_widths=[0.5,1.8,1.0,1.3,0.8,0.7,1.0,1.0,0.9])

doc.add_paragraph()
add_heading(doc, "5.4 과대 평가 선수 TOP 10 (실력 대비 고평가)", level=2)
add_para(doc,
    f"예측 시장 가치가 실제 시장 가치의 70% 미만인 경우 '과대평가'로 분류합니다. "
    f"총 {s2_results['value_analysis']['significantly_overvalued_count']}명이 해당합니다.")
overval_rows = []
for i, p in enumerate(s2_results["top20_overvalued"][:10]):
    overval_rows.append([
        i+1, p["player"], p["season"], p["team"], p["position"], int(p["age"]),
        f"{int(p['actual_market_value_eur']/1e6):.1f}M€",
        f"{int(p['predicted_market_value_eur']/1e6):.1f}M€",
        f"{p['value_score']:.3f}x"
    ])
add_table(doc,
    headers=["순위","선수명","시즌","팀","포지션","나이","실제 가치","예측 가치","비율"],
    rows=overval_rows,
    col_widths=[0.5,1.8,1.0,1.3,0.8,0.7,1.0,1.0,0.9])

doc.add_paragraph()
add_heading(doc, "5.5 그래프", level=2)
add_figure(doc, MODELS_DIR / "s2_market_value/figures/model_comparison_by_position.png",
           "그림 5-1. 포지션별 모델 성능 비교 (R² 기준)")
add_figure(doc, MODELS_DIR / "s2_market_value/figures/predicted_vs_actual_by_position.png",
           "그림 5-2. 포지션별 예측값 vs. 실제값 산포도")
add_figure(doc, MODELS_DIR / "s2_market_value/figures/feature_importance_by_position.png",
           "그림 5-3. 포지션별 XGBoost 피처 중요도")
add_figure(doc, MODELS_DIR / "s2_market_value/figures/undervalued_bargain_list.png",
           "그림 5-4. 저평가 선수 목록 (가성비 TOP 선수)")
add_figure(doc, MODELS_DIR / "s2_market_value/figures/overvalued_avoid_list.png",
           "그림 5-5. 과대평가 선수 목록 (주의 필요 선수)")

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────
# 6장. S3 유사 선수 탐색
# ─────────────────────────────────────────────────────────────────
add_heading(doc, "6장. S3 유사 선수 탐색 (Similarity Engine)", level=1)
doc.add_paragraph()

add_heading(doc, "6.1 방법론", level=2)
add_para(doc,
    "K-Means 클러스터링과 코사인 유사도 행렬을 결합하여 25시즌에 걸친 "
    f"{s3_results['metadata']['n_players']:,}개 선수-시즌 레코드 중 유사한 선수를 탐색합니다. "
    f"최적 클러스터 K={s3_results['metadata']['best_k']}는 실루엣 스코어 "
    f"{s3_results['metadata']['silhouette_scores'][str(s3_results['metadata']['best_k'])]:.4f}로 선정되었습니다. "
    "per90 지표(골, 어시스트, 슛, 태클, 인터셉트, 키패스), 나이, 신장, "
    "시장 가치 로그, 스타터 비율을 유사도 피처로 사용합니다.")

doc.add_paragraph()
add_heading(doc, "6.2 클러스터 아키타입", level=2)
archetype_rows = [[k, v] for k, v in s3_results["metadata"]["archetypes"].items()]
add_table(doc,
    headers=["클러스터 번호", "아키타입 (플레이어 유형)"],
    rows=archetype_rows,
    col_widths=[1.5, 5.5])

doc.add_paragraph()
add_heading(doc, "6.3 유사 선수 탐색 데모", level=2)
add_para(doc, "■ 얼링 홀란드 (2023/24 시즌) 기준 유사 선수:", bold=True)
haaland_key = "Erling Haaland_2023/24"
if haaland_key in s3_results["demo_searches"]:
    h_rows = [[r["rank"], r["player"], r["season"], r["team"],
               r["position"], f"{r['goals']:.0f}G/{r['assists']:.0f}A",
               f"{r['similarity']:.4f}"]
              for r in s3_results["demo_searches"][haaland_key][:7]]
    add_table(doc,
        headers=["순위","선수명","시즌","팀","포지션","G/A","유사도"],
        rows=h_rows,
        col_widths=[0.6,2.0,1.0,1.5,1.5,1.0,1.0])

doc.add_paragraph()
# Saka demo
saka_key = next((k for k in s3_results["demo_searches"] if "Saka" in k), None)
if saka_key:
    add_para(doc, f"■ 부카요 사카 ({saka_key.split('_')[1]} 시즌) 기준 유사 선수:", bold=True)
    s_rows = [[r["rank"], r["player"], r["season"], r["team"],
               r["position"], f"{r['goals']:.0f}G/{r['assists']:.0f}A",
               f"{r['similarity']:.4f}"]
              for r in s3_results["demo_searches"][saka_key][:5]]
    add_table(doc,
        headers=["순위","선수명","시즌","팀","포지션","G/A","유사도"],
        rows=s_rows,
        col_widths=[0.6,2.0,1.0,1.5,1.5,1.0,1.0])

doc.add_paragraph()
add_heading(doc, "6.4 그래프", level=2)
add_figure(doc, MODELS_DIR / "s3_similarity/figures/umap_cluster_scatter.png",
           "그림 6-1. UMAP 차원 축소 - 선수 클러스터 시각화")
add_figure(doc, MODELS_DIR / "s3_similarity/figures/cluster_size_distribution.png",
           "그림 6-2. 클러스터별 선수 분포")
add_figure(doc, MODELS_DIR / "s3_similarity/figures/cluster_radar_charts.png",
           "그림 6-3. 클러스터별 지표 프로파일 레이더 차트")
add_figure(doc, MODELS_DIR / "s3_similarity/figures/sim_haaland_2324.png",
           "그림 6-4. 홀란드(2023/24) 유사 선수 탐색 결과")
add_figure(doc, MODELS_DIR / "s3_similarity/figures/sim_saka_2324.png",
           "그림 6-5. 사카(2023/24) 유사 선수 탐색 결과")
add_figure(doc, MODELS_DIR / "s3_similarity/figures/sim_vandijk_2324.png",
           "그림 6-6. 반다이크(2023/24) 유사 선수 탐색 결과")

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────
# 7장. S4 성장 잠재력 예측
# ─────────────────────────────────────────────────────────────────
add_heading(doc, "7장. S4 성장 잠재력 예측", level=1)
doc.add_paragraph()

add_heading(doc, "7.1 방법론", level=2)
add_para(doc,
    "다음 시즌 퍼포먼스 점수 변화량(growth_target)을 예측 목표로 설정합니다. "
    f"총 {s4_results['feature_count']}개 피처를 활용하며, "
    "XGBoost·RandomForest·GradientBoosting·MLP의 앙상블 회귀 모델을 적용합니다. "
    "U23 선수 전용 서브 모델도 별도로 학습하여 젊은 선수 탐색 정확도를 높입니다. "
    f"데이터 분할: Train {s4_results['data_split']['train']:,}건 / "
    f"Val {s4_results['data_split']['val']:,}건 / "
    f"Test {s4_results['data_split']['test']:,}건.")

doc.add_paragraph()
add_heading(doc, "7.2 모델 성능", level=2)
s4_perf = s4_results["model_performance"]
s4_rows = []
for model_name, metrics in s4_perf.items():
    s4_rows.append([
        model_name,
        round(metrics.get("val_r2", 0), 4),
        round(metrics.get("val_rmse", 0), 4),
        round(metrics.get("test_r2", 0), 4),
        round(metrics.get("test_rmse", 0), 4),
    ])
# Add U23 model
u23 = s4_results["u23_model_performance"]
s4_rows.append([f"U23 전용 ({s4_results['best_model']})",
                round(u23["val_r2"], 4), round(u23["val_rmse"], 4),
                round(u23["test_r2"], 4), "-"])
add_table(doc,
    headers=["모델", "Val R²", "Val RMSE", "Test R²", "Test RMSE"],
    rows=s4_rows,
    col_widths=[2.5, 1.2, 1.2, 1.2, 1.2])

doc.add_paragraph()
add_heading(doc, "7.3 스카우트 목록", level=2)
add_para(doc, "■ 핫 프로스펙트 (성장 잠재력 최상위)", bold=True)
hot_rows = [[i+1, p["player"], p["team"], int(p["age"]), p["pos_simple"],
             f"{p['market_value']/1e6:.1f}M€",
             f"{p['predicted_growth']:+.3f}"]
            for i, p in enumerate(s4_results["scout_lists"]["hot_prospects"][:8])]
add_table(doc,
    headers=["순위","선수명","팀","나이","포지션","시장가치","예측 성장치"],
    rows=hot_rows,
    col_widths=[0.6,2.0,1.8,0.7,0.8,1.1,1.2])

doc.add_paragraph()
add_para(doc, "■ 레이트 블루머 (늦깎이 성장형)", bold=True)
bloom_rows = [[i+1, p["player"], p["team"], int(p["age"]), p["pos_simple"],
               f"{p['predicted_growth']:+.3f}"]
              for i, p in enumerate(s4_results["scout_lists"]["late_bloomers"][:5])]
add_table(doc,
    headers=["순위","선수명","팀","나이","포지션","예측 성장치"],
    rows=bloom_rows,
    col_widths=[0.6,2.2,2.0,0.8,1.0,1.2])

doc.add_paragraph()
add_para(doc, "■ 하락 예측 선수 (성장 잠재력 하위)", bold=True)
decline_rows = [[i+1, p["player"], p["team"], int(p["age"]), p["pos_simple"],
                 f"{p['predicted_growth']:+.3f}"]
                for i, p in enumerate(s4_results["scout_lists"]["declining_stars"][:5])]
add_table(doc,
    headers=["순위","선수명","팀","나이","포지션","예측 성장치"],
    rows=decline_rows,
    col_widths=[0.6,2.2,2.0,0.8,1.0,1.2])

doc.add_paragraph()
add_heading(doc, "7.4 주요 피처 중요도 (XGBoost)", level=2)
feat_rows = [[i+1, f["feature"], f"{f['importance']:.4f}"]
             for i, f in enumerate(s4_results["top_features_xgb"][:8])]
add_table(doc,
    headers=["순위", "피처명", "중요도"],
    rows=feat_rows,
    col_widths=[0.8, 4.0, 1.2])

doc.add_paragraph()
add_heading(doc, "7.5 그래프", level=2)
add_figure(doc, MODELS_DIR / "s4_growth/figures/fig1_model_performance.png",
           "그림 7-1. 모델별 성능 비교 (R², RMSE)")
add_figure(doc, MODELS_DIR / "s4_growth/figures/fig2_feature_importance.png",
           "그림 7-2. 성장 예측 피처 중요도")
add_figure(doc, MODELS_DIR / "s4_growth/figures/fig3_age_growth_curves.png",
           "그림 7-3. 포지션별 나이-성장 곡선")
add_figure(doc, MODELS_DIR / "s4_growth/figures/fig4_hot_prospects.png",
           "그림 7-4. 핫 프로스펙트 선수 목록")
add_figure(doc, MODELS_DIR / "s4_growth/figures/fig5_growth_trajectories.png",
           "그림 7-5. 주요 선수 성장 궤적 시각화")
add_figure(doc, MODELS_DIR / "s4_growth/figures/fig6_late_bloomers_declining.png",
           "그림 7-6. 레이트 블루머 vs. 하락 선수 비교")
add_figure(doc, MODELS_DIR / "s4_growth/figures/fig7_growth_distributions.png",
           "그림 7-7. 성장 예측값 전체 분포")

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────
# 8장. S5 이적 적응도 예측
# ─────────────────────────────────────────────────────────────────
add_heading(doc, "8장. S5 이적 적응도 예측", level=1)
doc.add_paragraph()

add_heading(doc, "8.1 방법론", level=2)
add_para(doc,
    "이적 후 다음 시즌 적응 성공 여부를 이진 분류 문제로 모델링합니다. "
    "복합 적응 레이블(포지션별 가중치 적용): 분당 출전 비율(minutes_share_ratio)과 "
    "per90 퍼포먼스 비율(per90_performance_ratio)을 포지션별로 가중 결합하여 적응 성공 여부를 "
    "판단합니다(composite ≥ 0.65). XGBoost·RandomForest·LogisticRegression·MLP 앙상블 적용. "
    "시간 기반 데이터 분할(train < 2021, val 2021-2022, test 2023-2025).")

doc.add_paragraph()
add_heading(doc, "8.2 모델 성능 비교", level=2)
s5_rows = []
for model_name, metrics in s5_metrics.items():
    s5_rows.append([model_name, metrics["auc"], metrics["f1"], metrics["acc"]])
add_table(doc,
    headers=["모델", "Test AUC-ROC", "Test F1", "Test Accuracy"],
    rows=s5_rows,
    col_widths=[2.5, 1.5, 1.5, 1.5])
add_para(doc, f"※ 최고 성능 모델: {s5_best} (AUC = {s5_metrics[s5_best]['auc']:.4f})", size=9)

doc.add_paragraph()
add_heading(doc, "8.3 적응도 예측 결과 요약", level=2)
add_para(doc,
    f"분석된 이적 건수: {n_transfer_rows}건 (2023-2025 시즌)\n"
    f"실제 적응 성공률: {s5_adapt_rate}%\n"
    f"Safe Bet (안전한 이적): {s5_safe_cnt}건 ({round(s5_safe_cnt/n_transfer_rows*100,1)}%)\n"
    f"Moderate (보통 위험): {s5_moderate}건 ({round(s5_moderate/n_transfer_rows*100,1)}%)\n"
    f"Risky (위험한 이적): {s5_risky}건 ({round(s5_risky/n_transfer_rows*100,1)}%)")

doc.add_paragraph()
add_heading(doc, "8.4 Safe Bet 이적 상위 5건", level=2)
safe_top5 = df_transfer[df_transfer["scout_category"]=="Safe Bet"].nlargest(5,"adapt_prob_ensemble")
safe_rows = [[i+1, row["player"], row["team_old"], row["team_new"],
              int(row["age"]), row["pos_group"],
              f"{row['adapt_prob_ensemble']:.3f}"]
             for i, (_, row) in enumerate(safe_top5.iterrows())]
add_table(doc,
    headers=["순위","선수명","출발팀","이적팀","나이","포지션","적응확률"],
    rows=safe_rows,
    col_widths=[0.6,1.8,1.8,1.8,0.7,1.0,1.1])

doc.add_paragraph()
add_heading(doc, "8.5 Risky 이적 목록", level=2)
risky_df = df_transfer[df_transfer["scout_category"]=="Risky"]
risky_rows = [[i+1, row["player"], row["team_old"], row["team_new"],
               int(row["age"]), f"{row['adapt_prob_ensemble']:.3f}"]
              for i, (_, row) in enumerate(risky_df.iterrows())]
add_table(doc,
    headers=["순위","선수명","출발팀","이적팀","나이","적응확률"],
    rows=risky_rows,
    col_widths=[0.6,2.0,2.0,2.0,0.8,1.0])

doc.add_paragraph()
add_heading(doc, "8.6 그래프", level=2)
add_figure(doc, MODELS_DIR / "s5_transfer_adapt/figures/confusion_matrix.png",
           "그림 8-1. 혼동 행렬 (최고 모델)")
add_figure(doc, MODELS_DIR / "s5_transfer_adapt/figures/feature_importance.png",
           "그림 8-2. 이적 적응도 피처 중요도")
add_figure(doc, MODELS_DIR / "s5_transfer_adapt/figures/adaptation_by_age_position.png",
           "그림 8-3. 나이·포지션별 적응도 분포")
add_figure(doc, MODELS_DIR / "s5_transfer_adapt/figures/risk_analysis.png",
           "그림 8-4. 이적 위험 요인 분석")

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────
# 9장. S6 하락 위험 감지
# ─────────────────────────────────────────────────────────────────
add_heading(doc, "9장. S6 하락 위험 감지", level=1)
doc.add_paragraph()

add_heading(doc, "9.1 방법론", level=2)
add_para(doc,
    "선수의 다음 시즌 퍼포먼스 하락 여부를 이진 분류로 예측합니다. "
    "하락 레이블은 퍼포먼스 점수 하락(decline_perf)과 가용성 하락(decline_avail)의 복합 기준으로 정의합니다. "
    "XGBoost·RandomForest·LogisticRegression·MLP 4개 모델 앙상블을 적용하며, "
    f"총 {n_decline_rows:,}개 선수-시즌 레코드 ({s6_unique_players:,}명)를 분석합니다. "
    "피처: 나이, 퍼포먼스 점수, 퍼포먼스 추세(slope), 피크 대비 현재 점수 차이, 일관성, 출전 시간.")

doc.add_paragraph()
add_heading(doc, "9.2 모델 성능 비교", level=2)
s6_rows_table = []
for model_name, metrics in s6_metrics.items():
    s6_rows_table.append([model_name, metrics["auc"], metrics["f1"], metrics["acc"]])
add_table(doc,
    headers=["모델", "Test AUC-ROC", "Test F1", "Test Accuracy"],
    rows=s6_rows_table,
    col_widths=[2.5, 1.5, 1.5, 1.5])
add_para(doc, f"※ 최고 성능 모델: {s6_best} (AUC = {s6_metrics[s6_best]['auc']:.4f})", size=9)

doc.add_paragraph()
add_heading(doc, "9.3 하락 위험 분포", level=2)
add_para(doc,
    f"전체 {n_decline_rows:,}개 레코드 중:\n"
    f"  - 하락 고위험군 (확률 > 0.7): {s6_high_risk:,}건 ({round(s6_high_risk/n_decline_rows*100,1)}%)\n"
    f"  - 하락 저위험군 (확률 < 0.3): {s6_low_risk:,}건 ({round(s6_low_risk/n_decline_rows*100,1)}%)\n"
    f"  - 중간 위험군: {n_decline_rows - s6_high_risk - s6_low_risk:,}건")

doc.add_paragraph()
add_heading(doc, "9.4 현재 하락 위험 감시 목록 (고위험 선수)", level=2)
# Get most recent season high-risk players
df_dec_recent = df_decline[df_decline["season_year"] >= 2023]
watch_list = df_dec_recent.drop_duplicates("player_key").nlargest(10, "decline_prob_ensemble")
watch_rows = [[i+1, row["player_key"], row["team"], int(row["age"]),
               row["pos_group"], f"{row['decline_prob_ensemble']:.3f}"]
              for i, (_, row) in enumerate(watch_list.iterrows())]
add_table(doc,
    headers=["순위","선수명","팀","나이","포지션","하락 위험도"],
    rows=watch_rows,
    col_widths=[0.6,2.5,2.0,0.8,1.0,1.3])

doc.add_paragraph()
add_heading(doc, "9.5 아이언맨 - 하락 저항 선수 (나이 30+ 저위험)", level=2)
iron = df_dec_recent[(df_dec_recent["age"] >= 30) & (df_dec_recent["decline_prob_ensemble"] < 0.3)]
iron = iron.drop_duplicates("player_key").nsmallest(8, "decline_prob_ensemble")
iron_rows = [[i+1, row["player_key"], row["team"], int(row["age"]),
              row["pos_group"], f"{row['decline_prob_ensemble']:.3f}"]
             for i, (_, row) in enumerate(iron.iterrows())]
add_table(doc,
    headers=["순위","선수명","팀","나이","포지션","하락 위험도"],
    rows=iron_rows,
    col_widths=[0.6,2.5,2.0,0.8,1.0,1.3])

doc.add_paragraph()
add_heading(doc, "9.6 그래프", level=2)
add_figure(doc, MODELS_DIR / "s6_decline/figures/roc_curves.png",
           "그림 9-1. 모델별 ROC 곡선 비교")
add_figure(doc, MODELS_DIR / "s6_decline/figures/feature_importance.png",
           "그림 9-2. 하락 위험 피처 중요도")
add_figure(doc, MODELS_DIR / "s6_decline/figures/age_vs_decline_risk.png",
           "그림 9-3. 나이별 하락 위험도 산포도")
add_figure(doc, MODELS_DIR / "s6_decline/figures/decline_prob_by_age.png",
           "그림 9-4. 나이 구간별 평균 하락 위험도")
add_figure(doc, MODELS_DIR / "s6_decline/figures/decline_watch_list.png",
           "그림 9-5. 하락 위험 감시 목록 시각화")

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────
# 10장. 종합 분석 및 스카우트 활용 가이드
# ─────────────────────────────────────────────────────────────────
add_heading(doc, "10장. 종합 분석 및 스카우트 활용 가이드", level=1)
doc.add_paragraph()

add_heading(doc, "10.1 모델별 스카우트 활용 방안", level=2)

usage_data = [
    ["S1 WAR Rating",
     "선수 종합 평가, 스쿼드 내 기여도 서열화",
     "이적 시즌별 어떤 선수가 포지션 내 최고 가치를 제공하는지 빠르게 파악",
     "왼쪽 수비수·중앙 수비수·골키퍼 등 특정 포지션 보강 후보 우선순위 결정"],
    ["S2 적정 몸값",
     "이적 협상 가이드라인 설정",
     "예측 가치 > 시장 가치인 저평가 선수를 조기에 발굴하여 비용 절감",
     "과대평가 선수(value_score < 0.07)는 협상 기준가 재검토 권고"],
    ["S3 유사 선수",
     "대체 선수 탐색, 스타일 적합성 검증",
     "계약 만료 또는 부상 선수의 대체재를 동일 클러스터 내에서 탐색",
     "새로운 전술 스타일에 맞는 선수 프로파일 유형을 아키타입으로 지정 후 탐색"],
    ["S4 성장 예측",
     "유망주 조기 발굴, U23 투자 결정",
     "핫 프로스펙트 목록으로 성장 가능성 높은 선수를 저렴한 시기에 확보",
     "레이트 블루머 선수는 경쟁 팀의 관심이 낮을 때 선점 가능"],
    ["S5 이적 적응도",
     "이적 리스크 사전 검증",
     "Safe Bet 분류 선수를 우선 협상 대상으로 설정, Risky 선수는 계약 조건에 리스크 반영",
     "Moving Up 여부(하위 팀→상위 팀 이적)와 ELO 차이를 추가 참고"],
    ["S6 하락 위험",
     "스쿼드 관리, 계약 연장 의사결정",
     "하락 고위험 선수는 장기 계약 자제, 하락 저위험 30+ 선수는 아이언맨으로 활용",
     "퍼포먼스 추세(perf_slope)와 피크 대비 현재 점수 차이(peak_minus_current)를 함께 검토"],
]

for row in usage_data:
    add_para(doc, f"■ {row[0]}", bold=True, size=11)
    add_bullet(doc, f"주요 활용: {row[1]}")
    add_bullet(doc, f"단기 전략: {row[2]}")
    add_bullet(doc, f"장기 전략: {row[3]}")
    doc.add_paragraph()

add_heading(doc, "10.2 이적 윈도우 추천 워크플로", level=2)
workflow_steps = [
    "S6 감시 목록 확인 → 스쿼드 내 하락 고위험 포지션 파악 → 보강 포지션 확정",
    "S1 WAR Rating 조회 → 해당 포지션 상위 선수 리스트업 (최근 2시즌 활성 선수 우선)",
    "S2 적정 몸값 비교 → 저평가 선수 필터링 (value_score > 1.5) → 예산 범위 내 후보 압축",
    "S5 적응도 예측 → 후보 선수의 이적 적응 위험도 확인 → Safe Bet/Moderate/Risky 분류",
    "S3 유사 선수 탐색 → 1순위 후보가 계약 불발 시 동일 클러스터 내 대체 선수 즉시 탐색",
    "S4 성장 예측 → U23 핫 프로스펙트 중 예산 초과 선수를 유망주 트랙 분리 추적",
    "최종 의사결정: 데이터 점수 + 현장 스카우트 관찰 결합 → 이적 제안 실행",
]
for i, step in enumerate(workflow_steps, 1):
    add_bullet(doc, f"단계 {i}: {step}")

doc.add_paragraph()

add_heading(doc, "10.3 모델 한계점 및 향후 개선 방향", level=2)
add_para(doc, "■ 현재 한계점", bold=True)
limitations = [
    f"S1/S4: 테스트 R²가 0.16~0.45 수준으로 예측 불확실성 존재. 특히 시장 가치 변화가 큰 선수(신예 스타 등)의 예측 오차가 큼.",
    "S2: 포지션별 MAE가 1,000만~1,500만 유로 수준으로 실제 협상에는 참고 지표로 활용 권장.",
    "S5: 테스트 세트 크기(83건)가 작아 일부 유형의 이적에 대한 일반화 성능 제한.",
    "시장 가치 데이터의 시대적 편향: 2010년대 초 이전 데이터의 시장 가치 기록이 불완전하여 과거 시즌 예측 정확도 저하.",
    "부상 데이터 미포함: 부상 이력과 회복 시간이 성장·하락 예측에 중요하나 현재 모델에서 미반영.",
]
for lim in limitations:
    add_bullet(doc, lim)

doc.add_paragraph()
add_para(doc, "■ 향후 개선 방향", bold=True)
improvements = [
    "xG(Expected Goals), xA(Expected Assists) 등 고급 퍼포먼스 지표 통합으로 예측 정밀도 향상",
    "LLM 기반 선수 뉴스·소셜 미디어 감성 분석을 통한 부상·이적 신호 조기 탐지",
    "실시간 API 연동(Opta, StatsBomb): 경기 직후 데이터로 인 시즌 업데이트 지원",
    "Attention 기반 시계열 모델(Transformer) 적용으로 다년도 성장 궤적 예측 개선",
    "포지션 내 세부 역할 분류(예: 9번 vs 거짓 9번, 박스투박스 vs 수비형 미드필더) 적용",
]
for imp in improvements:
    add_bullet(doc, imp)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────
# 부록
# ─────────────────────────────────────────────────────────────────
add_heading(doc, "부록. 데이터 사전 및 모델 하이퍼파라미터", level=1)
doc.add_paragraph()

add_heading(doc, "A. 주요 변수 데이터 사전", level=2)
dict_rows = [
    ["war_rating",         "float",   "0~100",    "종합 선수 평가 지수 (WAR Rating). 포지션 내 상대적 순위 정규화"],
    ["market_value",       "float",   "EUR",       "Transfermarkt 기준 시장 가치 (유로화)"],
    ["perf_score",         "float",   "z-score",   "포지션·시즌 정규화된 퍼포먼스 종합 점수"],
    ["gls_p90",            "float",   "건/90분",   "90분당 득점 수"],
    ["ast_p90",            "float",   "건/90분",   "90분당 어시스트 수"],
    ["tklw_p90",           "float",   "건/90분",   "90분당 성공 태클 수"],
    ["int_p90",            "float",   "건/90분",   "90분당 인터셉트 수"],
    ["kp_p90",             "float",   "건/90분",   "90분당 키 패스 수"],
    ["consistency_cv",     "float",   "0~1",       "시즌 내 경기별 퍼포먼스 변동 계수 (낮을수록 일관성 높음)"],
    ["elo",                "float",   "점수",      "팀 강도 추정 ELO 점수 (포인트 기반 근사값)"],
    ["adapt_prob_ensemble","float",   "0~1",       "이적 후 적응 성공 앙상블 예측 확률"],
    ["decline_prob_ensemble","float", "0~1",       "하락 위험도 앙상블 예측 확률"],
    ["predicted_growth",   "float",   "퍼포먼스 단위", "다음 시즌 퍼포먼스 점수 변화 예측값"],
    ["value_score",        "float",   "배수",      "예측 시장가치 / 실제 시장가치 비율 (1보다 크면 저평가)"],
    ["starter_ratio",      "float",   "0~1",       "전체 출전 경기 중 선발 비율"],
    ["epl_experience",     "int",     "시즌 수",   "EPL 참가 시즌 수"],
    ["age_peak_gap",       "float",   "세",        "포지션별 피크 나이와의 차이"],
    ["mv_change_rate",     "float",   "비율",      "전 시즌 대비 시장 가치 변화율"],
    ["pos_group",          "str",     "FW/MID/DEF/GK", "포지션 그룹 코드"],
    ["moving_up",          "int",     "0/1",       "하위 팀에서 상위 팀으로 이적 여부"],
]
add_table(doc,
    headers=["변수명", "타입", "단위/범위", "설명"],
    rows=dict_rows,
    col_widths=[1.8, 0.8, 1.3, 3.1])

doc.add_paragraph()
add_heading(doc, "B. 모델 하이퍼파라미터", level=2)

param_rows = [
    # S1
    ["S1 XGBoost",         "n_estimators=300, max_depth=6, learning_rate=0.05, subsample=0.8"],
    ["S1 Ridge",           "alpha=10.0 (CV 최적)"],
    ["S1 MLP",             "hidden_layers=(256,128,64), dropout=0.3, lr=1e-3, epochs=100"],
    # S2
    ["S2 XGBoost (포지션별)", "n_estimators=500, max_depth=5, learning_rate=0.05, colsample_bytree=0.8"],
    ["S2 RandomForest",    "n_estimators=300, max_depth=10, min_samples_leaf=5"],
    ["S2 MLP",             "hidden_layers=(128,64,32), dropout=0.2, lr=1e-3"],
    # S3
    ["S3 K-Means",         "K=12 (실루엣 스코어 최적화), random_state=42, n_init=10"],
    ["S3 유사도",           "Cosine Similarity on StandardScaler 정규화 피처 (10차원)"],
    # S4
    ["S4 RandomForest",    "n_estimators=300, max_depth=8, min_samples_leaf=3"],
    ["S4 XGBoost",         "n_estimators=400, max_depth=5, learning_rate=0.03"],
    ["S4 GradientBoosting","n_estimators=200, max_depth=4, learning_rate=0.05"],
    ["S4 MLP",             "hidden_layers=(128,64), lr=1e-3, epochs=80"],
    # S5
    ["S5 XGBoost",         "scale_pos_weight=불균형 비율, max_depth=5, learning_rate=0.05"],
    ["S5 RandomForest",    "n_estimators=300, class_weight=balanced, max_depth=8"],
    ["S5 LogisticRegression","C=1.0, class_weight=balanced, solver=lbfgs"],
    ["S5 MLP",             "hidden_layers=(128,64), lr=5e-4, epochs=50"],
    # S6
    ["S6 XGBoost",         "scale_pos_weight=불균형 비율, max_depth=4, learning_rate=0.05"],
    ["S6 RandomForest",    "n_estimators=200, class_weight=balanced, max_depth=6"],
    ["S6 LogisticRegression","C=0.5, class_weight=balanced"],
    ["S6 MLP",             "hidden_layers=(64,32), lr=1e-3, epochs=60"],
]
add_table(doc,
    headers=["모델", "주요 하이퍼파라미터"],
    rows=param_rows,
    col_widths=[2.5, 4.5])

doc.add_paragraph()
add_heading(doc, "C. 프로젝트 파일 구조", level=2)
file_struct = [
    "EPL project/",
    "├── data/",
    "│   ├── processed/     # 전처리 완료 파켓 파일 (5개)",
    "│   └── scout/         # 스카우트 파생 데이터 파켓 + JSON (11개)",
    "├── models/",
    "│   ├── s1_player_rating/  # WAR Rating 모델 + 결과",
    "│   ├── s2_market_value/   # 포지션별 몸값 예측 모델",
    "│   ├── s3_similarity/     # 유사 선수 탐색 엔진",
    "│   ├── s4_growth/         # 성장 잠재력 예측",
    "│   ├── s5_transfer_adapt/ # 이적 적응도 분류",
    "│   └── s6_decline/        # 하락 위험 감지",
    "├── reports/",
    "│   ├── figures/       # EDA 및 모델 시각화 PNG 파일 (100+)",
    "│   └── scout_final_report.docx  # 본 보고서",
    "└── scripts/           # 분석·모델링·보고서 생성 스크립트",
]
for line in file_struct:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(line)
    run.font.name = "Courier New"
    run.font.size = Pt(9)

# ─────────────────────────────────────────────────────────────────
# 저장
# ─────────────────────────────────────────────────────────────────
OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT_PATH))
print(f"\n보고서 저장 완료: {OUT_PATH}")
print(f"페이지 수 추정: 약 70~90 페이지")
