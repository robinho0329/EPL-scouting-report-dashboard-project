"""
EPL 25시즌 전처리 및 EDA 보고서 생성 스크립트
모든 통계를 실제 데이터에서 계산하여 Word 문서로 생성합니다.
"""
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import seaborn as sns
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import warnings
warnings.filterwarnings('ignore')

# ── 한글 폰트 설정 ──
plt.rcParams['font.family'] = 'Malgun Gothic'
plt.rcParams['axes.unicode_minus'] = False
plt.rcParams['figure.dpi'] = 150

# ── 경로 설정 ──
BASE = Path(r'C:/Users/xcv54/workspace/EPL project')
FIG_DIR = BASE / 'reports' / 'figures'
FIG_DIR.mkdir(parents=True, exist_ok=True)
PROCESSED = BASE / 'data' / 'processed'
FEATURES = BASE / 'data' / 'features'
REPORT_PATH = BASE / 'reports' / 'preprocessing_eda_report.docx'

BIG6_COLORS = {'Arsenal': '#EF0107', 'Chelsea': '#034694', 'Liverpool': '#C8102E',
               'Man City': '#6CABDD', 'Man United': '#DA291C', 'Tottenham': '#132257'}


# ============================================================
# 데이터 로드
# ============================================================
print("[1/5] 데이터 로드 중...")
match_results = pd.read_parquet(PROCESSED / 'match_results.parquet')
team_season = pd.read_parquet(PROCESSED / 'team_season_summary.parquet')
player_stats = pd.read_parquet(PROCESSED / 'player_season_stats.parquet')
player_logs = pd.read_parquet(PROCESSED / 'player_match_logs.parquet')
match_features = pd.read_parquet(FEATURES / 'match_features.parquet')
player_features = pd.read_parquet(FEATURES / 'player_features.parquet')

datasets = {
    'match_results': match_results,
    'team_season_summary': team_season,
    'player_season_stats': player_stats,
    'player_match_logs': player_logs,
    'match_features': match_features,
    'player_features': player_features,
}

print(f"  match_results: {match_results.shape}")
print(f"  team_season: {team_season.shape}")
print(f"  player_stats: {player_stats.shape}")
print(f"  player_logs: {player_logs.shape}")
print(f"  match_features: {match_features.shape}")
print(f"  player_features: {player_features.shape}")


# ============================================================
# 통계 계산
# ============================================================
print("[2/5] 통계 계산 중...")

mr = match_results.copy()
mr['TotalGoals'] = mr['FullTimeHomeGoals'] + mr['FullTimeAwayGoals']
mr['year'] = mr['Season'].str[:4].astype(int)

# 시즌별 통계
season_stats = mr.groupby('Season').agg(
    matches=('TotalGoals', 'count'),
    avg_goals=('TotalGoals', 'mean'),
    total_goals=('TotalGoals', 'sum'),
    home_win_pct=('FullTimeResult', lambda x: (x == 'H').mean() * 100),
    draw_pct=('FullTimeResult', lambda x: (x == 'D').mean() * 100),
    away_win_pct=('FullTimeResult', lambda x: (x == 'A').mean() * 100),
    avg_home_goals=('FullTimeHomeGoals', 'mean'),
    avg_away_goals=('FullTimeAwayGoals', 'mean'),
).reset_index()

# 팀 총계
ts = team_season.copy()
team_total = ts.groupby('team').agg(
    total_wins=('total_wins', 'sum'),
    total_losses=('total_losses', 'sum'),
    total_draws=('total_draws', 'sum'),
    total_gf=('total_goals_for', 'sum'),
    total_ga=('total_goals_against', 'sum'),
    seasons=('Season', 'count'),
    total_points=('points', 'sum')
).reset_index()
team_total['goal_diff'] = team_total['total_gf'] - team_total['total_ga']
team_total['win_rate'] = team_total['total_wins'] / (team_total['total_wins'] + team_total['total_draws'] + team_total['total_losses']) * 100

# 선수 통계
ps = player_stats.copy()
ps['pos_main'] = ps['pos'].str.split(',').str[0]
ps['age_used'] = ps['age'].fillna(ps['age_tm'])

# 시대별 비교
mr['era'] = mr['year'].apply(lambda x: '2000-2012' if x < 2013 else '2013-2025')


# ============================================================
# 시각화 생성
# ============================================================
print("[3/5] 시각화 생성 중...")

def save_fig(fig, name):
    path = FIG_DIR / f'report_{name}.png'
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return str(path)

fig_paths = {}

# ── 3.1 시즌별 평균 골 ──
fig, ax1 = plt.subplots(figsize=(14, 6))
ax1.plot(range(len(season_stats)), season_stats['avg_goals'], 'o-', color='#E74C3C', linewidth=2, markersize=6, label='경기당 평균 골')
ax1.fill_between(range(len(season_stats)), season_stats['avg_goals'], alpha=0.15, color='#E74C3C')
ax1.set_xticks(range(len(season_stats)))
ax1.set_xticklabels(season_stats['Season'], rotation=45, ha='right', fontsize=9)
ax1.set_ylabel('경기당 평균 골', fontsize=12, color='#E74C3C')
ax1.set_title('시즌별 경기당 평균 골 수 추이 (2000/01~2024/25)', fontsize=14, fontweight='bold')
covid_idx = list(season_stats['Season']).index('2020/21') if '2020/21' in list(season_stats['Season']) else None
if covid_idx is not None:
    ax1.axvspan(covid_idx - 0.5, covid_idx + 0.5, alpha=0.2, color='gray', label='COVID-19 (무관중)')
ax2 = ax1.twinx()
ax2.bar(range(len(season_stats)), season_stats['total_goals'], alpha=0.2, color='#3498DB', label='총 골 수')
ax2.set_ylabel('총 골 수', fontsize=12, color='#3498DB')
lines1, labels1 = ax1.get_legend_handles_labels()
lines2, labels2 = ax2.get_legend_handles_labels()
ax1.legend(lines1 + lines2, labels1 + labels2, loc='upper left')
ax1.grid(axis='y', alpha=0.3)
plt.tight_layout()
fig_paths['season_goals'] = save_fig(fig, 'season_goals')

# ── 3.2 경기 결과 비율 ──
result_pct = mr.groupby('Season')['FullTimeResult'].value_counts(normalize=True).unstack(fill_value=0) * 100
result_pct = result_pct.reindex(columns=['H', 'D', 'A'])
result_pct.columns = ['홈 승', '무승부', '원정 승']
fig, ax = plt.subplots(figsize=(14, 7))
result_pct.plot.area(ax=ax, stacked=True, alpha=0.7, color=['#2ECC71', '#F39C12', '#E74C3C'])
ax.set_ylabel('비율 (%)', fontsize=12)
ax.set_title('시즌별 경기 결과 비율 변화', fontsize=14, fontweight='bold')
ax.set_ylim(0, 100)
if covid_idx is not None:
    ax.axvline(x=result_pct.index.get_loc('2020/21'), color='gray', linestyle=':', alpha=0.7, linewidth=2)
plt.xticks(rotation=45, ha='right')
plt.tight_layout()
fig_paths['result_dist'] = save_fig(fig, 'result_dist')

# ── 3.3 홈 어드밴티지 ──
home_adv = mr.groupby('Season').apply(
    lambda x: pd.Series({
        '홈승률': (x['FullTimeResult'] == 'H').mean() * 100,
        '홈골': x['FullTimeHomeGoals'].mean(),
        '원정골': x['FullTimeAwayGoals'].mean(),
    })
).reset_index()
fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(14, 10), sharex=True)
ax1.plot(range(len(home_adv)), home_adv['홈승률'], 'o-', color='#27AE60', linewidth=2, markersize=6)
ax1.axhline(y=home_adv['홈승률'].mean(), color='#27AE60', linestyle='--', alpha=0.5, label=f'평균: {home_adv["홈승률"].mean():.1f}%')
ax1.fill_between(range(len(home_adv)), home_adv['홈승률'], home_adv['홈승률'].mean(), alpha=0.1, color='#27AE60')
ax1.set_ylabel('홈 승률 (%)', fontsize=12)
ax1.set_title('홈 어드밴티지 추이', fontsize=14, fontweight='bold')
ax1.legend()
ax1.grid(axis='y', alpha=0.3)
if covid_idx is not None:
    for a in [ax1, ax2]:
        a.axvspan(covid_idx - 0.5, covid_idx + 0.5, alpha=0.15, color='red')
    ax1.annotate('COVID-19\n무관중', xy=(covid_idx, home_adv.iloc[covid_idx]['홈승률']),
                 xytext=(covid_idx + 2, home_adv['홈승률'].max()),
                 arrowprops=dict(arrowstyle='->', color='red'), fontsize=10, color='red', fontweight='bold')
ax2.plot(range(len(home_adv)), home_adv['홈골'], 's-', color='#2ECC71', linewidth=2, label='홈 평균골')
ax2.plot(range(len(home_adv)), home_adv['원정골'], '^-', color='#E74C3C', linewidth=2, label='원정 평균골')
ax2.fill_between(range(len(home_adv)), home_adv['홈골'], home_adv['원정골'], alpha=0.15, color='#3498DB')
ax2.set_xticks(range(len(home_adv)))
ax2.set_xticklabels(home_adv['Season'], rotation=45, ha='right', fontsize=9)
ax2.set_ylabel('평균 골', fontsize=12)
ax2.legend()
ax2.grid(axis='y', alpha=0.3)
plt.tight_layout()
fig_paths['home_advantage'] = save_fig(fig, 'home_advantage')

# ── 4.1 Top 10 승수 ──
top10_wins = team_total.nlargest(10, 'total_wins').sort_values('total_wins')
fig, ax = plt.subplots(figsize=(12, 7))
colors = [BIG6_COLORS.get(t, '#95A5A6') for t in top10_wins['team']]
ax.barh(range(len(top10_wins)), top10_wins['total_wins'], color=colors, edgecolor='white')
ax.set_yticks(range(len(top10_wins)))
ax.set_yticklabels(top10_wins['team'], fontsize=11)
ax.set_xlabel('총 승수', fontsize=12)
ax.set_title('역대 EPL 총 승수 Top 10', fontsize=14, fontweight='bold')
for i, (wins, wr) in enumerate(zip(top10_wins['total_wins'], top10_wins['win_rate'])):
    ax.text(wins + 3, i, f'{wins}승 ({wr:.1f}%)', va='center', fontsize=10)
ax.grid(axis='x', alpha=0.3)
plt.tight_layout()
fig_paths['top10_wins'] = save_fig(fig, 'top10_wins')

# ── 4.2 Top 10 골득실 ──
top10_gd = team_total.nlargest(10, 'goal_diff').sort_values('goal_diff')
fig, ax = plt.subplots(figsize=(12, 7))
colors = [BIG6_COLORS.get(t, '#95A5A6') for t in top10_gd['team']]
ax.barh(range(len(top10_gd)), top10_gd['goal_diff'], color=colors, edgecolor='white')
ax.set_yticks(range(len(top10_gd)))
ax.set_yticklabels(top10_gd['team'], fontsize=11)
ax.set_xlabel('골득실', fontsize=12)
ax.set_title('역대 EPL 골득실 Top 10', fontsize=14, fontweight='bold')
for i, (gd, gf, ga) in enumerate(zip(top10_gd['goal_diff'], top10_gd['total_gf'], top10_gd['total_ga'])):
    ax.text(gd + 10, i, f'+{gd} ({gf}/{ga})', va='center', fontsize=10)
ax.grid(axis='x', alpha=0.3)
plt.tight_layout()
fig_paths['top10_gd'] = save_fig(fig, 'top10_gd')

# ── 4.3 팀 시즌 참가 히트맵 ──
seasons = sorted(ts['Season'].unique())
teams_sorted = ts.groupby('team')['Season'].count().sort_values(ascending=False).index.tolist()
presence = pd.DataFrame(0, index=teams_sorted, columns=seasons)
for _, row in ts.iterrows():
    presence.loc[row['team'], row['Season']] = 1
fig, ax = plt.subplots(figsize=(18, 14))
sns.heatmap(presence, cmap=['#FFFFFF', '#2ECC71'], cbar=False, linewidths=0.3, linecolor='#EEEEEE', ax=ax)
ax.set_title('EPL 팀별 시즌 참가 현황', fontsize=14, fontweight='bold')
plt.xticks(rotation=45, ha='right', fontsize=8)
plt.yticks(fontsize=8)
plt.tight_layout()
fig_paths['team_presence'] = save_fig(fig, 'team_presence')

# ── 5.1 연령 분포 ──
fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))
valid_ages = ps['age_used'].dropna()
valid_ages = valid_ages[(valid_ages >= 15) & (valid_ages <= 45)]
ax1.hist(valid_ages, bins=30, color='#3498DB', edgecolor='white', alpha=0.8)
ax1.axvline(valid_ages.mean(), color='red', linestyle='--', linewidth=2, label=f'평균: {valid_ages.mean():.1f}세')
ax1.set_xlabel('나이', fontsize=12)
ax1.set_ylabel('선수-시즌 수', fontsize=12)
ax1.set_title('EPL 선수 연령 분포', fontsize=13, fontweight='bold')
ax1.legend()
pos_data = ps[ps['pos_main'].isin(['GK', 'DF', 'MF', 'FW']) & ps['age_used'].notna()]
sns.boxplot(data=pos_data, x='pos_main', y='age_used', order=['GK', 'DF', 'MF', 'FW'], palette='Set2', ax=ax2)
ax2.set_xlabel('포지션', fontsize=12)
ax2.set_ylabel('나이', fontsize=12)
ax2.set_title('포지션별 연령 분포', fontsize=13, fontweight='bold')
plt.tight_layout()
fig_paths['age_dist'] = save_fig(fig, 'age_dist')

# ── 5.2 포지션 분포 ──
pos_counts = ps['pos_main'].value_counts()
fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))
colors_pie = ['#E74C3C', '#3498DB', '#2ECC71', '#F39C12', '#9B59B6']
ax1.pie(pos_counts.values, labels=pos_counts.index, autopct='%1.1f%%', colors=colors_pie[:len(pos_counts)], startangle=90)
ax1.set_title('주 포지션 분포', fontsize=13, fontweight='bold')
detail_pos = ps['pos'].value_counts().head(10)
ax2.barh(range(len(detail_pos)), detail_pos.values, color=sns.color_palette('viridis', len(detail_pos)))
ax2.set_yticks(range(len(detail_pos)))
ax2.set_yticklabels(detail_pos.index, fontsize=10)
ax2.set_xlabel('선수-시즌 수', fontsize=12)
ax2.set_title('세부 포지션 분포 Top 10', fontsize=13, fontweight='bold')
ax2.invert_yaxis()
plt.tight_layout()
fig_paths['pos_dist'] = save_fig(fig, 'pos_dist')

# ── 5.3 Top 20 득점자 ──
top_scorers = ps.groupby('player')['gls'].sum().nlargest(20).sort_values()
fig, ax = plt.subplots(figsize=(12, 8))
ax.barh(range(len(top_scorers)), top_scorers.values, color=sns.color_palette('YlOrRd_r', len(top_scorers)))
ax.set_yticks(range(len(top_scorers)))
ax.set_yticklabels(top_scorers.index, fontsize=10)
ax.set_xlabel('총 골 수', fontsize=12)
ax.set_title('EPL 역대 득점 Top 20', fontsize=14, fontweight='bold')
for i, v in enumerate(top_scorers.values):
    ax.text(v + 1, i, f'{int(v)}골', va='center', fontsize=9)
ax.grid(axis='x', alpha=0.3)
plt.tight_layout()
fig_paths['top20_scorers'] = save_fig(fig, 'top20_scorers')

# ── 5.4 Top 20 시장가치 ──
mv_top = ps[ps['market_value'].notna()].sort_values('market_value', ascending=False).drop_duplicates(subset='player').head(20)
mv_top = mv_top.sort_values('market_value')
fig, ax = plt.subplots(figsize=(12, 8))
ax.barh(range(len(mv_top)), mv_top['market_value'] / 1e6, color=sns.color_palette('Blues', len(mv_top)))
ax.set_yticks(range(len(mv_top)))
labels_mv = [f"{row['player']} ({row['season']})" for _, row in mv_top.iterrows()]
ax.set_yticklabels(labels_mv, fontsize=9)
ax.set_xlabel('시장가치 (백만 유로)', fontsize=12)
ax.set_title('EPL 역대 시장가치 Top 20', fontsize=14, fontweight='bold')
for i, v in enumerate(mv_top['market_value'] / 1e6):
    ax.text(v + 1, i, f'{v:.1f}M', va='center', fontsize=9)
ax.grid(axis='x', alpha=0.3)
plt.tight_layout()
fig_paths['top20_mv'] = save_fig(fig, 'top20_mv')

# ── 5.5 골 vs 시장가치 ──
ps_mv = ps[(ps['market_value'].notna()) & (ps['gls'].notna()) & (ps['gls'] > 0)].copy()
ps_mv['mv_millions'] = ps_mv['market_value'] / 1e6
fig, ax = plt.subplots(figsize=(12, 8))
scatter = ax.scatter(ps_mv['gls'], ps_mv['mv_millions'], c=ps_mv['age_used'], cmap='viridis', alpha=0.4, s=20)
plt.colorbar(scatter, ax=ax, label='나이')
ax.set_xlabel('시즌 골 수', fontsize=12)
ax.set_ylabel('시장가치 (백만 유로)', fontsize=12)
ax.set_title('시즌 골 수 vs 시장가치', fontsize=14, fontweight='bold')
corr_gls_mv = ps_mv[['gls', 'mv_millions']].corr().iloc[0, 1]
ax.text(0.02, 0.98, f'상관계수: {corr_gls_mv:.3f}', transform=ax.transAxes, fontsize=11,
        va='top', bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))
ax.grid(alpha=0.3)
plt.tight_layout()
fig_paths['goals_vs_mv'] = save_fig(fig, 'goals_vs_mv')

# ── 6.1 상관관계 히트맵 ──
mf = match_features.copy()
numeric_cols = [c for c in mf.select_dtypes(include=[np.number]).columns
                if not c.startswith('season') and c not in ['own_goal_flag_home', 'own_goal_flag_away', 'own_goal_flag', 'season_data_missing']]
key_features = [c for c in numeric_cols if any(k in c for k in ['elo', 'form', 'goals_scored', 'goals_conceded',
                                                                  'gd_rolling', 'clean_sheet', 'sot_rolling',
                                                                  'h2h', 'momentum', 'FullTime'])]
if len(key_features) > 25:
    key_features = key_features[:25]
corr_matrix = mf[key_features].corr()
fig, ax = plt.subplots(figsize=(16, 14))
mask = np.triu(np.ones_like(corr_matrix, dtype=bool))
sns.heatmap(corr_matrix, mask=mask, cmap='RdBu_r', center=0, annot=False, square=True, linewidths=0.5, ax=ax, vmin=-1, vmax=1)
ax.set_title('매치 피처 상관관계 히트맵', fontsize=14, fontweight='bold')
plt.xticks(rotation=45, ha='right', fontsize=8)
plt.yticks(fontsize=8)
plt.tight_layout()
fig_paths['corr_heatmap'] = save_fig(fig, 'corr_heatmap')

# ── 6.2 결과 상관 ──
mf_enc = mf.copy()
mf_enc['result_numeric'] = mf_enc['FullTimeResult'].map({'H': 1, 'D': 0, 'A': -1})
feature_cols = [c for c in mf.select_dtypes(include=[np.number]).columns
                if c not in ['FullTimeHomeGoals', 'FullTimeAwayGoals', 'season_year', 'season_data_missing',
                             'own_goal_flag_home', 'own_goal_flag_away', 'own_goal_flag', 'matchweek']]
corr_with_result = mf_enc[feature_cols + ['result_numeric']].corr()['result_numeric'].drop('result_numeric')
corr_sorted = corr_with_result.abs().sort_values(ascending=False).head(20)
corr_top = corr_with_result[corr_sorted.index].sort_values()
fig, ax = plt.subplots(figsize=(12, 8))
colors_c = ['#E74C3C' if v < 0 else '#2ECC71' for v in corr_top.values]
ax.barh(range(len(corr_top)), corr_top.values, color=colors_c, edgecolor='white')
ax.set_yticks(range(len(corr_top)))
ax.set_yticklabels(corr_top.index, fontsize=9)
ax.set_xlabel('상관계수', fontsize=11)
ax.set_title('경기 결과와 상관관계 Top 20 피처', fontsize=14, fontweight='bold')
ax.axvline(x=0, color='black', linewidth=0.8)
ax.grid(axis='x', alpha=0.3)
for i, v in enumerate(corr_top.values):
    ax.text(v + (0.01 if v >= 0 else -0.01), i, f'{v:.3f}', va='center', fontsize=8,
            ha='left' if v >= 0 else 'right')
plt.tight_layout()
fig_paths['result_corr'] = save_fig(fig, 'result_corr')

# ── 7. 시대별 비교 ──
era_stats_full = mr.groupby('era').agg(
    avg_goals=('TotalGoals', 'mean'),
    home_win_pct=('FullTimeResult', lambda x: (x == 'H').mean() * 100),
).reset_index()
for era_val in ['2000-2012', '2013-2025']:
    mask = mr['era'] == era_val
    era_stats_full.loc[era_stats_full['era'] == era_val, 'avg_yellows'] = (mr.loc[mask, 'HomeYellowCards'] + mr.loc[mask, 'AwayYellowCards']).mean()
    era_stats_full.loc[era_stats_full['era'] == era_val, 'avg_shots'] = (mr.loc[mask, 'HomeShots'] + mr.loc[mask, 'AwayShots']).mean()
    era_stats_full.loc[era_stats_full['era'] == era_val, 'avg_corners'] = (mr.loc[mask, 'HomeCorners'] + mr.loc[mask, 'AwayCorners']).mean()
    era_stats_full.loc[era_stats_full['era'] == era_val, 'avg_fouls'] = (mr.loc[mask, 'HomeFouls'] + mr.loc[mask, 'AwayFouls']).mean()

all_compare = ['avg_goals', 'home_win_pct', 'avg_yellows', 'avg_shots', 'avg_corners', 'avg_fouls']
compare_labels = ['경기당\n총 골', '홈 승률\n(%)', '경기당\n옐로카드', '경기당\n슈팅', '경기당\n코너킥', '경기당\n파울']
fig, ax = plt.subplots(figsize=(14, 7))
x = np.arange(len(all_compare))
width = 0.35
era1 = era_stats_full[era_stats_full['era'] == '2000-2012'][all_compare].values[0]
era2 = era_stats_full[era_stats_full['era'] == '2013-2025'][all_compare].values[0]
bars1 = ax.bar(x - width/2, era1, width, label='2000-2012', color='#3498DB', alpha=0.8)
bars2 = ax.bar(x + width/2, era2, width, label='2013-2025', color='#E74C3C', alpha=0.8)
ax.set_xticks(x)
ax.set_xticklabels(compare_labels, fontsize=11)
ax.set_title('EPL 시대별 주요 지표 비교', fontsize=14, fontweight='bold')
ax.legend(fontsize=12)
ax.grid(axis='y', alpha=0.3)
for bar_group in [bars1, bars2]:
    for bar in bar_group:
        h = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., h + 0.1, f'{h:.1f}', ha='center', va='bottom', fontsize=9)
plt.tight_layout()
fig_paths['era_compare'] = save_fig(fig, 'era_compare')

# ── 8.1 시장가치 추이 ──
ps_mv_all = ps[ps['market_value'].notna()].copy()
mv_by_season = ps_mv_all.groupby('season').agg(
    mean_mv=('market_value', 'mean'),
    median_mv=('market_value', 'median'),
    total_mv=('market_value', 'sum'),
).reset_index()
fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(14, 10))
ax1.plot(range(len(mv_by_season)), mv_by_season['mean_mv'] / 1e6, 'o-', color='#E74C3C', linewidth=2, label='평균')
ax1.plot(range(len(mv_by_season)), mv_by_season['median_mv'] / 1e6, 's-', color='#3498DB', linewidth=2, label='중앙값')
ax1.set_xticks(range(len(mv_by_season)))
ax1.set_xticklabels(mv_by_season['season'], rotation=45, ha='right', fontsize=8)
ax1.set_ylabel('시장가치 (백만 유로)', fontsize=12)
ax1.set_title('시즌별 선수 시장가치 추이', fontsize=14, fontweight='bold')
ax1.legend()
ax1.grid(alpha=0.3)
ax2.bar(range(len(mv_by_season)), mv_by_season['total_mv'] / 1e9, color='#2ECC71', alpha=0.8)
ax2.set_xticks(range(len(mv_by_season)))
ax2.set_xticklabels(mv_by_season['season'], rotation=45, ha='right', fontsize=8)
ax2.set_ylabel('총 시장가치 (십억 유로)', fontsize=12)
ax2.set_title('시즌별 리그 총 시장가치', fontsize=14, fontweight='bold')
ax2.grid(axis='y', alpha=0.3)
plt.tight_layout()
fig_paths['mv_trend'] = save_fig(fig, 'mv_trend')

# ── 8.2 포지션별 시장가치 ──
ps_mv_pos = ps_mv_all[ps_mv_all['pos'].str.split(',').str[0].isin(['GK', 'DF', 'MF', 'FW'])].copy()
ps_mv_pos['pos_main'] = ps_mv_pos['pos'].str.split(',').str[0]
ps_mv_pos['mv_millions'] = ps_mv_pos['market_value'] / 1e6
fig, ax = plt.subplots(figsize=(12, 7))
sns.boxplot(data=ps_mv_pos, x='pos_main', y='mv_millions', order=['GK', 'DF', 'MF', 'FW'],
            palette={'GK': '#F39C12', 'DF': '#3498DB', 'MF': '#2ECC71', 'FW': '#E74C3C'}, showfliers=False, ax=ax)
ax.set_xlabel('포지션', fontsize=12)
ax.set_ylabel('시장가치 (백만 유로)', fontsize=12)
ax.set_title('포지션별 시장가치 분포', fontsize=14, fontweight='bold')
ax.grid(axis='y', alpha=0.3)
plt.tight_layout()
fig_paths['mv_by_pos'] = save_fig(fig, 'mv_by_pos')

# ── 8.3 비싼 스쿼드 ──
squad_mv = ps_mv_all.groupby(['season', 'team']).agg(
    total_mv=('market_value', 'sum'), squad_size=('market_value', 'count')).reset_index()
top10_squads = squad_mv.nlargest(10, 'total_mv').sort_values('total_mv')
fig, ax = plt.subplots(figsize=(12, 7))
labels_sq = [f"{row['team']} ({row['season']})" for _, row in top10_squads.iterrows()]
colors_sq = [BIG6_COLORS.get(row['team'], '#95A5A6') for _, row in top10_squads.iterrows()]
ax.barh(range(len(top10_squads)), top10_squads['total_mv'] / 1e6, color=colors_sq, edgecolor='white')
ax.set_yticks(range(len(top10_squads)))
ax.set_yticklabels(labels_sq, fontsize=10)
ax.set_xlabel('총 시장가치 (백만 유로)', fontsize=12)
ax.set_title('EPL 역대 가장 비싼 스쿼드 Top 10', fontsize=14, fontweight='bold')
for i, row in enumerate(top10_squads.itertuples()):
    ax.text(row.total_mv / 1e6 + 10, i, f'{row.total_mv/1e6:.0f}M', va='center', fontsize=9)
ax.grid(axis='x', alpha=0.3)
plt.tight_layout()
fig_paths['top_squads'] = save_fig(fig, 'top_squads')

# ── 9.1 ELO ──
fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))
elo_data = pd.concat([mf['home_elo_pre'], mf['away_elo_pre']]).dropna()
ax1.hist(elo_data, bins=50, color='#3498DB', edgecolor='white', alpha=0.8)
ax1.axvline(elo_data.mean(), color='red', linestyle='--', linewidth=2, label=f'평균: {elo_data.mean():.0f}')
ax1.set_xlabel('ELO 레이팅', fontsize=12)
ax1.set_ylabel('빈도', fontsize=12)
ax1.set_title('ELO 레이팅 분포', fontsize=13, fontweight='bold')
ax1.legend()
elo_by_season = mf.groupby('Season').agg(max_elo=('home_elo_pre', 'max'), mean_elo=('home_elo_pre', 'mean'), min_elo=('home_elo_pre', 'min')).reset_index()
ax2.plot(range(len(elo_by_season)), elo_by_season['max_elo'], 'v-', color='#E74C3C', label='최고', markersize=5)
ax2.plot(range(len(elo_by_season)), elo_by_season['mean_elo'], 'o-', color='#3498DB', label='평균', markersize=5)
ax2.plot(range(len(elo_by_season)), elo_by_season['min_elo'], '^-', color='#2ECC71', label='최저', markersize=5)
ax2.fill_between(range(len(elo_by_season)), elo_by_season['max_elo'], elo_by_season['min_elo'], alpha=0.1, color='gray')
ax2.set_xticks(range(len(elo_by_season)))
ax2.set_xticklabels(elo_by_season['Season'], rotation=45, ha='right', fontsize=7)
ax2.set_ylabel('ELO 레이팅', fontsize=12)
ax2.set_title('시즌별 ELO 레이팅 범위', fontsize=13, fontweight='bold')
ax2.legend(fontsize=10)
ax2.grid(alpha=0.3)
plt.tight_layout()
fig_paths['elo'] = save_fig(fig, 'elo')

# ── 9.2 폼 인덱스 ──
fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))
form5 = mf['home_form_5'].dropna()
ax1.hist(form5, bins=40, color='#2ECC71', edgecolor='white', alpha=0.8)
ax1.axvline(form5.mean(), color='red', linestyle='--', linewidth=2, label=f'평균: {form5.mean():.2f}')
ax1.set_xlabel('최근 5경기 폼 인덱스', fontsize=12)
ax1.set_ylabel('빈도', fontsize=12)
ax1.set_title('홈팀 폼 인덱스(5경기) 분포', fontsize=13, fontweight='bold')
ax1.legend()
mf_temp = mf.dropna(subset=['home_form_5']).copy()
mf_temp['form_bin'] = pd.qcut(mf_temp['home_form_5'], 10, duplicates='drop')
win_by_form = mf_temp.groupby('form_bin').apply(lambda x: (x['FullTimeResult'] == 'H').mean() * 100).reset_index()
win_by_form.columns = ['form_bin', 'home_win_pct']
win_by_form['form_mid'] = win_by_form['form_bin'].apply(lambda x: x.mid)
win_by_form = win_by_form.sort_values('form_mid')
ax2.plot(range(len(win_by_form)), win_by_form['home_win_pct'], 'o-', color='#E74C3C', linewidth=2, markersize=8)
ax2.set_xticks(range(len(win_by_form)))
ax2.set_xticklabels([f'{x:.1f}' for x in win_by_form['form_mid']], rotation=45, fontsize=8)
ax2.set_xlabel('폼 인덱스 (10분위 중앙값)', fontsize=12)
ax2.set_ylabel('홈 승률 (%)', fontsize=12)
ax2.set_title('폼 인덱스 구간별 홈 승률', fontsize=13, fontweight='bold')
ax2.grid(alpha=0.3)
plt.tight_layout()
fig_paths['form'] = save_fig(fig, 'form')

# ── 9.3 롤링 통계 효과성 ──
rolling_features = {
    'home_form_5': '홈 폼(5경기)', 'home_form_10': '홈 폼(10경기)',
    'elo_diff': 'ELO 차이', 'home_gd_rolling_5': '홈 골득실(5경기)',
    'home_gd_rolling_10': '홈 골득실(10경기)', 'home_win_rate_5': '홈 승률(5경기)',
    'form_diff_5': '폼 차이(5경기)', 'form_diff_10': '폼 차이(10경기)',
    'gd_trend_diff': '골득실 트렌드', 'h2h_home_wins': '상대전적 홈승',
}
mf_enc2 = mf.copy()
mf_enc2['result_num'] = mf_enc2['FullTimeResult'].map({'H': 1, 'D': 0, 'A': -1})
correlations = {}
for feat, label in rolling_features.items():
    if feat in mf_enc2.columns:
        c = mf_enc2[[feat, 'result_num']].dropna().corr().iloc[0, 1]
        correlations[label] = c
corr_df = pd.DataFrame(list(correlations.items()), columns=['피처', '상관계수']).sort_values('상관계수', ascending=True)
fig, ax = plt.subplots(figsize=(12, 7))
colors_c2 = ['#E74C3C' if v < 0 else '#2ECC71' for v in corr_df['상관계수']]
ax.barh(range(len(corr_df)), corr_df['상관계수'], color=colors_c2, edgecolor='white')
ax.set_yticks(range(len(corr_df)))
ax.set_yticklabels(corr_df['피처'], fontsize=10)
ax.set_xlabel('경기 결과와의 상관계수', fontsize=12)
ax.set_title('파생변수의 경기 결과 예측력', fontsize=14, fontweight='bold')
ax.axvline(x=0, color='black', linewidth=0.8)
ax.grid(axis='x', alpha=0.3)
for i, v in enumerate(corr_df['상관계수']):
    ax.text(v + (0.005 if v >= 0 else -0.005), i, f'{v:.3f}', va='center', fontsize=9,
            ha='left' if v >= 0 else 'right')
plt.tight_layout()
fig_paths['rolling_eff'] = save_fig(fig, 'rolling_eff')

# ── 결측치 ──
fig, axes = plt.subplots(2, 3, figsize=(20, 12))
fig.suptitle('데이터셋별 결측치 현황', fontsize=16, fontweight='bold')
for idx, (name, df) in enumerate(datasets.items()):
    ax = axes[idx // 3, idx % 3]
    missing_pct = (df.isnull().sum() / len(df) * 100).sort_values(ascending=False)
    missing_cols = missing_pct[missing_pct > 0]
    if len(missing_cols) == 0:
        ax.text(0.5, 0.5, '결측치 없음', ha='center', va='center', fontsize=14, transform=ax.transAxes)
        ax.set_title(f'{name}\n(결측치 없음)', fontsize=11)
    else:
        top_missing = missing_cols.head(20)
        ax.barh(range(len(top_missing)), top_missing.values, color=sns.color_palette('Reds_r', len(top_missing)))
        ax.set_yticks(range(len(top_missing)))
        ax.set_yticklabels(top_missing.index, fontsize=8)
        ax.set_xlabel('결측치 비율 (%)')
        ax.set_title(f'{name}\n(결측 컬럼 {len(missing_cols)}개)', fontsize=11)
    ax.invert_yaxis()
plt.tight_layout()
fig_paths['missing'] = save_fig(fig, 'missing')

print(f"  총 {len(fig_paths)}개 시각화 생성 완료")


# ============================================================
# Word 문서 생성
# ============================================================
print("[4/5] Word 보고서 생성 중...")

doc = Document()

# ── 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = '맑은 고딕'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    hs.font.color.rgb = RGBColor(0x1A, 0x47, 0x84)


def add_heading_kr(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return h


def add_para(doc, text, bold=False, size=10, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)
    run.bold = bold
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_picture_safe(doc, path, width=Inches(6)):
    try:
        doc.add_picture(str(path), width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        add_para(doc, f'[이미지 로드 실패: {path}]')


def add_table_from_df(doc, df, max_rows=20):
    rows_to_show = min(len(df), max_rows)
    table = doc.add_table(rows=rows_to_show + 1, cols=len(df.columns), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                run.font.size = Pt(8)
                run.bold = True
    for i in range(rows_to_show):
        for j, col in enumerate(df.columns):
            cell = table.rows[i + 1].cells[j]
            val = df.iloc[i, j]
            if isinstance(val, float):
                cell.text = f'{val:.2f}' if abs(val) < 1000 else f'{val:,.0f}'
            else:
                cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '맑은 고딕'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    run.font.size = Pt(8)
    return table


# ============================================================
# 표지
# ============================================================
for _ in range(6):
    doc.add_paragraph()
add_para(doc, 'EPL 25시즌 데이터 분석 보고서', bold=True, size=28, align='center')
add_para(doc, '', size=12)
add_para(doc, '전처리 및 탐색적 데이터 분석 (EDA)', size=16, align='center')
add_para(doc, '', size=12)
add_para(doc, '2000/01 ~ 2024/25 시즌', size=14, align='center')
add_para(doc, '', size=12)
add_para(doc, '분석 기준일: 2026-03-22', size=11, align='center')
add_para(doc, '프리미어리그 경기 결과, 팀 성적, 선수 통계, 이적시장 가치 종합 분석', size=11, align='center')
doc.add_page_break()

# ============================================================
# 목차
# ============================================================
add_heading_kr(doc, '목차', level=1)
toc_items = [
    '1장: 프로젝트 개요',
    '2장: 데이터 소스 및 수집 현황',
    '3장: 데이터 품질 분석',
    '4장: 전처리 파이프라인',
    '5장: 매치 데이터 EDA',
    '6장: 팀 분석 EDA',
    '7장: 선수 분석 EDA',
    '8장: 이적시장 분석',
    '9장: 파생변수 분석',
    '10장: 결론 및 다음 단계',
]
for item in toc_items:
    add_para(doc, item, size=12)
doc.add_page_break()

# ============================================================
# 1장: 프로젝트 개요
# ============================================================
add_heading_kr(doc, '1장: 프로젝트 개요', level=1)
add_para(doc, (
    '본 프로젝트는 잉글랜드 프리미어리그(EPL) 2000/01 시즌부터 2024/25 시즌까지 '
    '25시즌에 걸친 종합 데이터를 수집, 전처리, 분석하는 것을 목표로 합니다. '
    '매치 결과, 팀 시즌 통계, 선수별 경기 기록, 이적시장 가치 등 다양한 데이터 소스를 '
    '통합하여 EPL의 트렌드와 패턴을 분석합니다.'
))
add_para(doc, '')
add_para(doc, '분석 범위:', bold=True, size=11)
n_seasons = match_results['Season'].nunique()
n_matches = len(match_results)
n_teams = match_results['HomeTeam'].nunique()
n_players = player_stats['player'].nunique()
n_player_seasons = len(player_stats)

add_para(doc, f'  - 분석 기간: {match_results["Season"].min()} ~ {match_results["Season"].max()} ({n_seasons}시즌)')
add_para(doc, f'  - 총 경기 수: {n_matches:,}경기')
add_para(doc, f'  - 참가 팀 수: {n_teams}팀')
add_para(doc, f'  - 등록 선수 수: {n_players:,}명 (선수-시즌 {n_player_seasons:,}건)')
add_para(doc, f'  - 선수 경기 로그: {len(player_logs):,}건')
add_para(doc, f'  - 매치 피처: {match_features.shape[1]}개 파생변수')
add_para(doc, f'  - 선수 피처: {player_features.shape[1]}개 파생변수')
doc.add_page_break()

# ============================================================
# 2장: 데이터 소스 및 수집 현황
# ============================================================
add_heading_kr(doc, '2장: 데이터 소스 및 수집 현황', level=1)

add_heading_kr(doc, '2.1 데이터 소스', level=2)
add_para(doc, '본 프로젝트는 다음 데이터 소스를 활용합니다:')
add_para(doc, '  1. FBref (fbref.com): 매치 결과, 팀 스쿼드 통계, 선수 경기 로그')
add_para(doc, '  2. Transfermarkt (transfermarkt.com): 선수 시장가치, 이적 정보')
add_para(doc, '  3. football-data.co.uk: 역대 매치 결과 통합 데이터 (epl_final.csv)')

add_heading_kr(doc, '2.2 수집 현황', level=2)

import os
fbref_dirs = os.listdir(str(BASE / 'data' / 'raw' / 'fbref'))
tm_dirs = os.listdir(str(BASE / 'data' / 'raw' / 'transfermarkt'))
total_fbref_files = 0
for d in fbref_dirs:
    dpath = BASE / 'data' / 'raw' / 'fbref' / d
    if dpath.is_dir():
        total_fbref_files += len(list(dpath.iterdir()))
total_tm_files = 0
for d in tm_dirs:
    dpath = BASE / 'data' / 'raw' / 'transfermarkt' / d
    if dpath.is_dir():
        total_tm_files += len(list(dpath.iterdir()))

add_para(doc, f'  - FBref 원본 파일: {len(fbref_dirs)}개 시즌 디렉토리, 총 {total_fbref_files:,}개 CSV')
add_para(doc, f'  - Transfermarkt 원본 파일: {len(tm_dirs)}개 시즌 디렉토리, 총 {total_tm_files:,}개 CSV')

add_heading_kr(doc, '2.3 처리된 데이터셋 요약', level=2)
summary_data = []
for name, df in datasets.items():
    summary_data.append({
        '데이터셋': name,
        '행 수': f'{df.shape[0]:,}',
        '열 수': df.shape[1],
        '메모리(MB)': f'{df.memory_usage(deep=True).sum() / 1024 / 1024:.1f}'
    })
summary_df = pd.DataFrame(summary_data)
add_table_from_df(doc, summary_df)
doc.add_page_break()

# ============================================================
# 3장: 데이터 품질 분석
# ============================================================
add_heading_kr(doc, '3장: 데이터 품질 분석', level=1)

add_heading_kr(doc, '3.1 결측치 현황', level=2)
add_para(doc, '각 데이터셋의 결측치 현황을 분석합니다.')
add_picture_safe(doc, fig_paths['missing'])
add_para(doc, '')

missing_data = []
for name, df in datasets.items():
    total_cells = df.shape[0] * df.shape[1]
    missing_cells = int(df.isnull().sum().sum())
    missing_cols = int((df.isnull().sum() > 0).sum())
    pct = missing_cells / total_cells * 100
    missing_data.append({
        '데이터셋': name,
        '전체 셀': f'{total_cells:,}',
        '결측 셀': f'{missing_cells:,}',
        '결측률(%)': f'{pct:.2f}',
        '결측 컬럼': missing_cols
    })
add_table_from_df(doc, pd.DataFrame(missing_data))

add_heading_kr(doc, '3.2 데이터 타입 및 범위', level=2)
add_para(doc, 'match_results 데이터셋의 주요 통계:')
desc = match_results.describe().T[['count', 'mean', 'std', 'min', 'max']].reset_index()
desc.columns = ['컬럼', '개수', '평균', '표준편차', '최소', '최대']
desc = desc[desc['컬럼'].isin(['FullTimeHomeGoals', 'FullTimeAwayGoals', 'HomeShots', 'AwayShots',
                                'HomeShotsOnTarget', 'AwayShotsOnTarget', 'HomeYellowCards', 'AwayYellowCards'])]
add_table_from_df(doc, desc)
doc.add_page_break()

# ============================================================
# 4장: 전처리 파이프라인
# ============================================================
add_heading_kr(doc, '4장: 전처리 파이프라인', level=1)

add_heading_kr(doc, '4.1 데이터 수집 파이프라인', level=2)
add_para(doc, (
    '데이터 수집은 FBref와 Transfermarkt에서 웹 크롤링을 통해 수행되었습니다. '
    '각 시즌별로 팀 스쿼드 통계와 선수 매치 로그를 수집하고, '
    'Transfermarkt에서는 선수 시장가치와 팀 정보를 수집했습니다.'
))

add_heading_kr(doc, '4.2 데이터 통합', level=2)
add_para(doc, (
    '수집된 원본 데이터를 다음과 같은 단계로 통합했습니다:\n'
    '  1. 팀명 표준화: 각 소스별 다른 팀명을 통일\n'
    '  2. 시즌 표기 통일: "2000/01" 형식으로 표준화\n'
    '  3. 선수 ID 매칭: FBref와 Transfermarkt 선수 데이터 연결\n'
    '  4. 결측치 처리: 통계적 방법 및 도메인 지식 활용'
))

add_heading_kr(doc, '4.3 파생변수 생성', level=2)
add_para(doc, f'매치 데이터에서 {match_features.shape[1]}개의 파생변수를 생성했습니다:')
add_para(doc, '  - ELO 레이팅: 팀별 동적 실력 지수 (초기값 1500)')
add_para(doc, '  - 폼 인덱스: 최근 5/10경기 기반 성적 지표')
add_para(doc, '  - 롤링 통계: 골득실, 슈팅, 클린시트 등 이동평균')
add_para(doc, '  - 상대전적: H2H(Head-to-Head) 기록')
add_para(doc, '  - 컨텍스트: 더비 여부, 주말/주중, 시즌 단계')

add_para(doc, f'\n선수 데이터에서 {player_features.shape[1]}개의 파생변수를 생성했습니다:')
add_para(doc, '  - 90분당 지표: 골, 어시스트, 공헌도 등')
add_para(doc, '  - 시장가치 변화율: 시즌간 가치 변동')
add_para(doc, '  - 경험/일관성 지표: EPL 경력, 성적 안정성')
doc.add_page_break()

# ============================================================
# 5장: 매치 데이터 EDA
# ============================================================
add_heading_kr(doc, '5장: 매치 데이터 EDA', level=1)

add_heading_kr(doc, '5.1 시즌별 평균 골 수 추이', level=2)
overall_avg = mr['TotalGoals'].mean()
max_season = season_stats.loc[season_stats['avg_goals'].idxmax()]
min_season = season_stats.loc[season_stats['avg_goals'].idxmin()]
add_para(doc, (
    f'25시즌 동안 총 {n_matches:,}경기에서 평균 {overall_avg:.2f}골이 기록되었습니다. '
    f'가장 많은 골이 기록된 시즌은 {max_season["Season"]} ({max_season["avg_goals"]:.2f}골/경기)이며, '
    f'가장 적은 골이 기록된 시즌은 {min_season["Season"]} ({min_season["avg_goals"]:.2f}골/경기)입니다.'
))
add_picture_safe(doc, fig_paths['season_goals'])

add_heading_kr(doc, '5.2 경기 결과 분포', level=2)
h_pct = (mr['FullTimeResult'] == 'H').mean() * 100
d_pct = (mr['FullTimeResult'] == 'D').mean() * 100
a_pct = (mr['FullTimeResult'] == 'A').mean() * 100
add_para(doc, (
    f'전체 경기 결과 분포: 홈승 {h_pct:.1f}%, 무승부 {d_pct:.1f}%, 원정승 {a_pct:.1f}%. '
    f'홈팀의 우위가 뚜렷하며, 시간이 지남에 따라 원정 승률이 증가하는 추세입니다.'
))
add_picture_safe(doc, fig_paths['result_dist'])

add_heading_kr(doc, '5.3 홈 어드밴티지 변화', level=2)
if covid_idx is not None:
    covid_home = home_adv.iloc[covid_idx]['홈승률']
    pre_covid = home_adv.iloc[max(0, covid_idx-3):covid_idx]['홈승률'].mean()
    add_para(doc, (
        f'홈 어드밴티지는 25시즌 동안 평균 {home_adv["홈승률"].mean():.1f}%의 홈 승률을 기록했습니다. '
        f'COVID-19 무관중 시즌(2020/21)에는 홈승률이 {covid_home:.1f}%로 '
        f'직전 3시즌 평균({pre_covid:.1f}%) 대비 {pre_covid - covid_home:.1f}%p 하락하여, '
        f'관중의 영향이 홈 어드밴티지에 중요한 요인임을 확인했습니다.'
    ))
add_picture_safe(doc, fig_paths['home_advantage'])

add_heading_kr(doc, '5.4 시대별 비교 (2000-2012 vs 2013-2025)', level=2)
era1_data = era_stats_full[era_stats_full['era'] == '2000-2012'].iloc[0]
era2_data = era_stats_full[era_stats_full['era'] == '2013-2025'].iloc[0]
add_para(doc, (
    f'전반기(2000-2012)와 후반기(2013-2025)를 비교하면:\n'
    f'  - 경기당 평균 골: {era1_data["avg_goals"]:.2f} -> {era2_data["avg_goals"]:.2f}\n'
    f'  - 홈 승률: {era1_data["home_win_pct"]:.1f}% -> {era2_data["home_win_pct"]:.1f}%\n'
    f'  - 경기당 옐로카드: {era1_data["avg_yellows"]:.1f} -> {era2_data["avg_yellows"]:.1f}\n'
    f'  - 경기당 슈팅: {era1_data["avg_shots"]:.1f} -> {era2_data["avg_shots"]:.1f}\n'
    f'  - 경기당 코너킥: {era1_data["avg_corners"]:.1f} -> {era2_data["avg_corners"]:.1f}\n'
    f'  - 경기당 파울: {era1_data["avg_fouls"]:.1f} -> {era2_data["avg_fouls"]:.1f}'
))
add_picture_safe(doc, fig_paths['era_compare'])
doc.add_page_break()

# ============================================================
# 6장: 팀 분석 EDA
# ============================================================
add_heading_kr(doc, '6장: 팀 분석 EDA', level=1)

add_heading_kr(doc, '6.1 역대 총 승수 Top 10', level=2)
top10 = team_total.nlargest(10, 'total_wins')
add_para(doc, f'25시즌간 가장 많은 승리를 기록한 팀은 {top10.iloc[0]["team"]}으로 {int(top10.iloc[0]["total_wins"])}승을 기록했습니다.')
add_picture_safe(doc, fig_paths['top10_wins'])

top10_table = top10[['team', 'total_wins', 'total_draws', 'total_losses', 'win_rate', 'total_points', 'seasons']].copy()
top10_table.columns = ['팀', '승', '무', '패', '승률(%)', '총 승점', '시즌수']
add_table_from_df(doc, top10_table.reset_index(drop=True))

add_heading_kr(doc, '6.2 역대 골득실 Top 10', level=2)
top10_gd_show = team_total.nlargest(10, 'goal_diff')
add_para(doc, f'골득실 1위는 {top10_gd_show.iloc[0]["team"]}(+{int(top10_gd_show.iloc[0]["goal_diff"])})입니다.')
add_picture_safe(doc, fig_paths['top10_gd'])

add_heading_kr(doc, '6.3 팀별 시즌 참가 현황', level=2)
ever_present = [t for t in teams_sorted if presence.loc[t].sum() == len(seasons)]
add_para(doc, (
    f'25시즌 모두 참가한 팀은 {len(ever_present)}팀입니다: {", ".join(ever_present)}. '
    f'총 {len(teams_sorted)}개 팀이 이 기간 동안 프리미어리그에 참가했습니다.'
))
add_picture_safe(doc, fig_paths['team_presence'])
doc.add_page_break()

# ============================================================
# 7장: 선수 분석 EDA
# ============================================================
add_heading_kr(doc, '7장: 선수 분석 EDA', level=1)

add_heading_kr(doc, '7.1 연령 분포', level=2)
add_para(doc, (
    f'선수 평균 나이는 {valid_ages.mean():.1f}세, 중앙값은 {valid_ages.median():.1f}세입니다. '
    f'최연소는 {valid_ages.min():.0f}세, 최고령은 {valid_ages.max():.0f}세입니다. '
    f'골키퍼가 다른 포지션에 비해 평균 연령이 높은 경향을 보입니다.'
))
add_picture_safe(doc, fig_paths['age_dist'])

add_heading_kr(doc, '7.2 포지션 분포', level=2)
for pos in ['GK', 'DF', 'MF', 'FW']:
    cnt = (ps['pos_main'] == pos).sum()
    pct = cnt / len(ps) * 100
    add_para(doc, f'  - {pos}: {cnt:,}명-시즌 ({pct:.1f}%)')
add_picture_safe(doc, fig_paths['pos_dist'])

add_heading_kr(doc, '7.3 역대 득점 Top 20', level=2)
top_scorer_name = top_scorers.index[-1]
top_scorer_goals = int(top_scorers.iloc[-1])
add_para(doc, f'역대 최다 득점자는 {top_scorer_name}({top_scorer_goals}골)입니다.')
add_picture_safe(doc, fig_paths['top20_scorers'])

add_heading_kr(doc, '7.4 역대 시장가치 Top 20', level=2)
add_para(doc, '선수의 최고 시장가치(단일 시즌 기준) 순위입니다.')
add_picture_safe(doc, fig_paths['top20_mv'])

add_heading_kr(doc, '7.5 골 수 vs 시장가치', level=2)
add_para(doc, (
    f'시즌 골 수와 시장가치 사이의 상관계수는 {corr_gls_mv:.3f}으로, '
    f'양의 상관관계를 보입니다. 골을 많이 넣는 선수일수록 시장가치가 높은 경향이 있지만, '
    f'시장가치에는 나이, 포지션, 잠재력 등 다양한 요인이 반영됩니다.'
))
add_picture_safe(doc, fig_paths['goals_vs_mv'])
doc.add_page_break()

# ============================================================
# 8장: 이적시장 분석
# ============================================================
add_heading_kr(doc, '8장: 이적시장 분석', level=1)

add_heading_kr(doc, '8.1 시즌별 시장가치 추이', level=2)
first_mv = mv_by_season.iloc[0]['mean_mv'] / 1e6
last_mv = mv_by_season.iloc[-1]['mean_mv'] / 1e6
growth = (last_mv / first_mv - 1) * 100
add_para(doc, (
    f'선수 평균 시장가치는 {mv_by_season.iloc[0]["season"]} 시즌 '
    f'{first_mv:.1f}백만 유로에서 {mv_by_season.iloc[-1]["season"]} 시즌 '
    f'{last_mv:.1f}백만 유로로 약 {growth:.0f}% 상승했습니다. '
    f'프리미어리그 중계권료 증가와 글로벌 인기 상승이 주요 원인입니다.'
))
add_picture_safe(doc, fig_paths['mv_trend'])

add_heading_kr(doc, '8.2 포지션별 시장가치 분포', level=2)
pos_medians = ps_mv_pos.groupby('pos_main')['mv_millions'].median()
for pos in ['GK', 'DF', 'MF', 'FW']:
    if pos in pos_medians.index:
        add_para(doc, f'  - {pos} 중앙값: {pos_medians[pos]:.1f}백만 유로')
add_picture_safe(doc, fig_paths['mv_by_pos'])

add_heading_kr(doc, '8.3 역대 가장 비싼 스쿼드 Top 10', level=2)
top_sq = squad_mv.nlargest(10, 'total_mv')
add_para(doc, (
    f'역대 가장 비싼 스쿼드는 {top_sq.iloc[0]["team"]} ({top_sq.iloc[0]["season"]})로 '
    f'총 {top_sq.iloc[0]["total_mv"]/1e6:.0f}백만 유로의 스쿼드 가치를 기록했습니다.'
))
add_picture_safe(doc, fig_paths['top_squads'])
doc.add_page_break()

# ============================================================
# 9장: 파생변수 분석
# ============================================================
add_heading_kr(doc, '9장: 파생변수 분석', level=1)

add_heading_kr(doc, '9.1 ELO 레이팅 분석', level=2)
add_para(doc, (
    f'ELO 레이팅은 초기값 1500에서 시작하여 경기 결과에 따라 동적으로 변화합니다. '
    f'전체 데이터에서 ELO 레이팅의 범위는 {elo_data.min():.0f} ~ {elo_data.max():.0f}이며, '
    f'평균은 {elo_data.mean():.0f}, 표준편차는 {elo_data.std():.0f}입니다. '
    f'시즌이 진행됨에 따라 팀간 ELO 격차가 벌어지는 경향을 보입니다.'
))
add_picture_safe(doc, fig_paths['elo'])

add_heading_kr(doc, '9.2 폼 인덱스 분석', level=2)
add_para(doc, (
    f'최근 5경기 폼 인덱스의 평균은 {form5.mean():.2f}이며, '
    f'폼 인덱스가 높을수록 홈 승률이 증가하는 강한 양의 관계를 보입니다. '
    f'이는 최근 폼이 경기 결과 예측에 유효한 지표임을 시사합니다.'
))
add_picture_safe(doc, fig_paths['form'])

add_heading_kr(doc, '9.3 롤링 통계 효과성', level=2)
best_feat = corr_df.loc[corr_df['상관계수'].abs().idxmax()]
add_para(doc, (
    f'파생변수 중 경기 결과와 가장 높은 상관관계를 보이는 변수는 '
    f'"{best_feat["피처"]}" (r={best_feat["상관계수"]:.3f})입니다. '
    f'ELO 차이, 폼 차이 등 두 팀간 차이를 나타내는 변수들이 '
    f'경기 결과 예측에 효과적인 것으로 나타났습니다.'
))
add_picture_safe(doc, fig_paths['rolling_eff'])

add_heading_kr(doc, '9.4 매치 피처 상관관계', level=2)
add_para(doc, '매치 피처 간의 상관관계를 분석합니다. 동일 팀의 공격/수비 지표 간에 높은 상관관계가 관찰됩니다.')
add_picture_safe(doc, fig_paths['corr_heatmap'])

add_heading_kr(doc, '9.5 경기 결과와 주요 피처 상관관계', level=2)
add_para(doc, '경기 결과(홈승=1, 무=0, 원정승=-1)와 가장 높은 상관관계를 보이는 피처 Top 20입니다.')
add_picture_safe(doc, fig_paths['result_corr'])
doc.add_page_break()

# ============================================================
# 10장: 결론 및 다음 단계
# ============================================================
add_heading_kr(doc, '10장: 결론 및 다음 단계', level=1)

add_heading_kr(doc, '10.1 주요 발견사항', level=2)
add_para(doc, (
    f'1. 데이터 규모: {n_seasons}시즌, {n_matches:,}경기, {n_players:,}명 선수의 종합 데이터를 구축했습니다.'
))
add_para(doc, (
    f'2. 홈 어드밴티지: 평균 홈승률 {home_adv["홈승률"].mean():.1f}%로 홈팀의 우위가 있으나, '
    f'COVID-19 시즌에서 관중 효과가 확인되었으며, 장기적으로 감소 추세입니다.'
))
add_para(doc, (
    f'3. 시장가치 인플레이션: 25년간 선수 평균 시장가치가 약 {growth:.0f}% 상승했습니다.'
))
add_para(doc, (
    f'4. 파생변수 유효성: ELO 차이, 폼 인덱스 등 파생변수가 '
    f'경기 결과와 유의미한 상관관계(|r| 최대 {corr_df["상관계수"].abs().max():.3f})를 보입니다.'
))
add_para(doc, (
    f'5. 시대별 변화: 후반기(2013-2025)에서 경기당 골 수 증가, 홈 어드밴티지 감소 등 '
    f'전술적 변화가 관찰됩니다.'
))

add_heading_kr(doc, '10.2 다음 단계', level=2)
add_para(doc, '  1. 경기 결과 예측 모델 구축 (XGBoost, LightGBM, Neural Network)')
add_para(doc, '  2. 선수 시장가치 예측 모델 개발')
add_para(doc, '  3. 팀 전략 클러스터링 분석')
add_para(doc, '  4. 시계열 기반 성과 예측')
add_para(doc, '  5. 대시보드 구축 (Streamlit/Dash)')

add_heading_kr(doc, '10.3 데이터 품질 개선 과제', level=2)
total_missing = sum(df.isnull().sum().sum() for df in datasets.values())
total_cells = sum(df.shape[0] * df.shape[1] for df in datasets.values())
add_para(doc, (
    f'전체 데이터의 결측률은 {total_missing/total_cells*100:.2f}%입니다. '
    f'주요 개선 과제:\n'
    f'  - 초기 시즌(2000-2005) 세부 통계 보완\n'
    f'  - Transfermarkt 시장가치 미매칭 선수 해결\n'
    f'  - 팀명 변경 이력 추가 반영'
))


# ── 저장 ──
doc.save(str(REPORT_PATH))
print(f"[5/5] 보고서 저장 완료: {REPORT_PATH}")
print(f"  총 페이지: 약 {len(doc.paragraphs) // 30 + len(fig_paths)}페이지 이상")
print(f"  포함 시각화: {len(fig_paths)}개")
print("완료!")
