# -*- coding: utf-8 -*-
"""
EPL 전처리 EDA 보고서 생성 스크립트 v3
- 한국어 전체 작성
- JSON 데이터 기반 테이블 포함
- PNG 시각화 이미지 임베딩
- 폰트: 맑은 고딕
"""

import os
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# 경로 설정
# ─────────────────────────────────────────────
BASE_DIR = r"C:\Users\xcv54\workspace\EPL project"
REPORTS_DIR = os.path.join(BASE_DIR, "reports")
FIGURES_DIR = os.path.join(REPORTS_DIR, "figures")
OUTPUT_PATH = os.path.join(REPORTS_DIR, "전처리_EDA_보고서_v2.docx")

JSON_MATCH    = os.path.join(REPORTS_DIR, "analysis_match_results.json")
JSON_SQUAD    = os.path.join(REPORTS_DIR, "analysis_fbref_squad.json")
JSON_MATCHLOG = os.path.join(REPORTS_DIR, "analysis_fbref_matchlogs.json")
JSON_TM       = os.path.join(REPORTS_DIR, "analysis_transfermarkt.json")
JSON_PROC     = os.path.join(REPORTS_DIR, "analysis_processed.json")

# ─────────────────────────────────────────────
# JSON 로드
# ─────────────────────────────────────────────
print("JSON 파일 로드 중...")
with open(JSON_MATCH,    encoding="utf-8") as f: d_match    = json.load(f)
with open(JSON_SQUAD,    encoding="utf-8") as f: d_squad    = json.load(f)
with open(JSON_MATCHLOG, encoding="utf-8") as f: d_mlog     = json.load(f)
with open(JSON_TM,       encoding="utf-8") as f: d_tm       = json.load(f)
with open(JSON_PROC,     encoding="utf-8") as f: d_proc     = json.load(f)
print("JSON 로드 완료.")

# ─────────────────────────────────────────────
# 문서 생성 및 헬퍼 함수
# ─────────────────────────────────────────────
doc = Document()

# 기본 페이지 여백 설정
from docx.shared import Cm
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)


def set_font(run, size_pt=10, bold=False, italic=False, color=None):
    """폰트를 맑은 고딕으로 설정"""
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 한글 폰트 명시적 설정
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rFonts.set(qn("w:ascii"), "맑은 고딕")
    rFonts.set(qn("w:hAnsi"), "맑은 고딕")
    rPr.insert(0, rFonts)


def add_title(text, size=36):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size_pt=size, bold=True, color=(31, 73, 125))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(12)
    return p


def add_heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size_pt=18, bold=True, color=(31, 73, 125))
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    # 하단 테두리
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size_pt=14, bold=True, color=(54, 96, 146))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_heading3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size_pt=12, bold=True, color=(79, 129, 189))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p


def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size_pt=10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[그림] {text}")
    set_font(run, size_pt=9, italic=True, color=(89, 89, 89))
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    return p


def add_image(filename, caption=""):
    """figures 폴더에서 이미지를 삽입하고 캡션을 추가"""
    img_path = os.path.join(FIGURES_DIR, filename)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(6.0))
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        if caption:
            add_caption(caption)
        return True
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[이미지 없음: {filename}]")
        set_font(run, size_pt=9, italic=True, color=(255, 0, 0))
        return False


def add_table(headers, rows, col_widths=None):
    """헤더 + 데이터 행으로 테이블 생성"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 행
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        set_font(run, size_pt=9, bold=True, color=(255, 255, 255))
        # 헤더 배경색 (진한 파란색)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F497D")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    # 데이터 행
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        bg_color = "EEF3F8" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, size_pt=9)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), bg_color)
            shd.set(qn("w:val"), "clear")
            tcPr.append(shd)

    # 열 너비 설정
    if col_widths:
        for r in table.rows:
            for i, cell in enumerate(r.cells):
                if i < len(col_widths):
                    cell.width = Inches(col_widths[i])
    doc.add_paragraph()
    return table


def add_page_break():
    doc.add_page_break()


# ─────────────────────────────────────────────
# 표지
# ─────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

add_title("EPL 데이터 프로젝트", size=36)
add_title("전처리 및 탐색적 데이터 분석(EDA)", size=28)
add_title("종합 보고서", size=24)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─" * 60)
set_font(run, size_pt=10, color=(100, 100, 100))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("영국 프리미어리그(EPL) 2000/01 ~ 2024/25 시즌 (25개 시즌)")
set_font(run, size_pt=12, color=(89, 89, 89))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("데이터 소스: epl_final.csv | FBref | Transfermarkt")
set_font(run, size_pt=11, color=(89, 89, 89))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("작성일: 2026년 3월 21일")
set_font(run, size_pt=10, color=(127, 127, 127))

add_page_break()

# ─────────────────────────────────────────────
# 목차
# ─────────────────────────────────────────────
add_heading1("목 차")
toc_items = [
    ("1장", "프로젝트 개요"),
    ("2장", "원본 데이터 현황"),
    ("3장", "데이터 품질 분석"),
    ("4장", "전처리 파이프라인"),
    ("5장", "탐색적 데이터 분석 (EDA)"),
    ("  5.1", "매치 결과 분석"),
    ("  5.2", "팀 분석"),
    ("  5.3", "선수 스탯 분석"),
    ("  5.4", "경기별 스탯 심층 분석"),
    ("  5.5", "몸값 분석"),
    ("6장", "시대별 비교 분석"),
    ("7장", "상관관계 분석"),
    ("8장", "데이터 통합 분석"),
    ("9장", "향후 계획"),
    ("10장", "결론"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f"{num}.  {title}")
    set_font(run, size_pt=10, bold=(not num.startswith(" ")))

add_page_break()

# ─────────────────────────────────────────────
# 1장. 프로젝트 개요
# ─────────────────────────────────────────────
add_heading1("1장. 프로젝트 개요")

add_heading2("1.1 분석 목적")
add_body("본 프로젝트는 영국 프리미어리그(EPL) 2000/01 시즌부터 2024/25 시즌까지 총 25개 시즌에 걸친 "
         "방대한 축구 데이터를 수집·정제·분석하는 것을 목적으로 합니다. 경기 결과, 선수 스탯, 이적 시장 가치 등 "
         "다양한 데이터를 통합하여 심층 인사이트를 도출하고, 향후 머신러닝 및 딥러닝 예측 모델 개발의 기반을 마련합니다.")

add_heading2("1.2 대상 기간 및 규모")
add_table(
    ["항목", "내용"],
    [
        ["분석 기간", "2000/01 ~ 2024/25 시즌 (총 25시즌)"],
        ["총 경기 수", f"{d_match['1_metadata']['shape']['rows']:,}경기"],
        ["총 골 수", f"{d_match['7_goals_analysis']['overall']['total_goals_all']:,}골"],
        ["등록 팀 수", f"{d_proc['per_file_analysis']['processed/match_results.parquet']['categorical_stats']['HomeTeam']['unique_count']}개 팀"],
        ["FBref squad 파일", f"{d_squad['metadata']['total_squad_stats_files']:,}개"],
        ["FBref matchlog 파일", f"{d_mlog['file_inventory']['total_matchlog_csv_files']:,}개"],
        ["Transfermarkt 파일", f"{d_tm['summary']['total_csv_files_found']:,}개"],
        ["전처리 후 선수-경기 레코드", f"{d_proc['per_file_analysis']['processed/player_match_logs.parquet']['shape']['rows']:,}행"],
        ["전처리 후 선수-시즌 레코드", f"{d_proc['per_file_analysis']['processed/player_season_stats.parquet']['shape']['rows']:,}행"],
    ],
    col_widths=[2.0, 4.5]
)

add_heading2("1.3 주요 분석 대상")
add_body("본 보고서는 다음과 같은 현장 실무자들을 주요 독자로 상정합니다:")
add_table(
    ["대상", "활용 목적"],
    [
        ["스카우터", "선수 성과 지표 분석, 잠재력 발굴"],
        ["이적 담당자", "시장 가치 추이, 나이-가치 상관관계 분석"],
        ["코치 / 감독", "전술 경향, 홈/원정 어드밴티지, 포지션별 스탯"],
        ["데이터 분석가", "머신러닝 특징 공학, 예측 모델 기반 구축"],
    ],
    col_widths=[2.0, 4.5]
)

add_heading2("1.4 데이터 소스 요약")
add_table(
    ["데이터 소스", "파일 형식", "수집 기간", "주요 내용", "파일 수"],
    [
        ["epl_final.csv", "CSV/Parquet", "2000/01~2024/25", "경기 결과, 슈팅, 코너, 카드", "1개"],
        ["FBref Squad Stats", "CSV", "2000-2001~2024-2025", "선수 시즌별 득점·어시스트·출전", "500개"],
        ["FBref Match Logs", "CSV", "2000-2001~2024-2025", "선수 경기별 세부 스탯", "11,369개"],
        ["Transfermarkt", "CSV", "2000~2024", "선수 시장 가치, 포지션, 국적", "500개"],
    ],
    col_widths=[1.8, 1.2, 1.8, 2.5, 0.9]
)

add_page_break()

# ─────────────────────────────────────────────
# 2장. 원본 데이터 현황
# ─────────────────────────────────────────────
add_heading1("2장. 원본 데이터 현황")

add_heading2("2.1 epl_final.csv 현황")
meta = d_match["1_metadata"]
add_body(f"epl_final.csv는 EPL 전체 경기 결과의 주요 소스 파일입니다. "
         f"총 {meta['shape']['rows']:,}행 × {meta['shape']['columns']}열로 구성되며, "
         f"파일 크기는 {meta['csv_file_size_bytes'] / 1024:.1f}KB (CSV), "
         f"Parquet 압축 후 {meta['parquet_file_size_bytes'] / 1024:.1f}KB입니다.")

add_heading3("주요 메타데이터")
add_table(
    ["항목", "값"],
    [
        ["행 수 (경기 수)", f"{meta['shape']['rows']:,}"],
        ["열 수", f"{meta['shape']['columns']}"],
        ["CSV 파일 크기", f"{meta['csv_file_size_bytes'] / 1024:.1f} KB"],
        ["Parquet 파일 크기", f"{meta['parquet_file_size_bytes'] / 1024:.1f} KB"],
        ["인코딩", meta['csv_encoding']],
        ["메모리 사용량", f"{meta['memory_usage_kb']:.1f} KB"],
        ["시즌 범위", "2000/01 ~ 2024/25"],
        ["결측치 비율", "0% (완전 데이터)"],
    ],
    col_widths=[2.5, 4.0]
)

add_heading3("컬럼 목록 (26개)")
cols_rows = []
for i in range(0, len(meta["columns"]), 3):
    row = meta["columns"][i:i+3]
    while len(row) < 3:
        row.append("")
    cols_rows.append(row)
add_table(["컬럼 1", "컬럼 2", "컬럼 3"], cols_rows, col_widths=[2.0, 2.0, 2.5])

add_heading2("2.2 FBref Squad Stats 현황")
squad_meta = d_squad["metadata"]
add_body(f"FBref squad_stats 데이터는 총 {squad_meta['total_squad_stats_files']}개 파일로 구성되며, "
         f"{squad_meta['total_seasons']}개 시즌 각 20개 팀의 선수별 시즌 집계 스탯을 담고 있습니다. "
         f"전체 25개 시즌(2000-2001 ~ 2024-2025)에서 오류 없이 수집되었습니다.")

add_heading3("시즌별 파일 수 (일부)")
season_sample = list(d_squad["files_per_season"].items())[:10]
add_table(
    ["시즌", "파일 수", "시즌", "파일 수", "시즌", "파일 수"],
    [
        [s1, n1, s2, n2, s3, n3]
        for (s1, n1), (s2, n2), (s3, n3) in zip(
            season_sample[0:5] + [(""," ")]*5,
            season_sample[5:10] + [(""," ")]*5,
            [(""," ")]*10
        )
    ][:5],
    col_widths=[1.0, 0.8, 1.0, 0.8, 1.0, 0.8]
)

add_heading3("컬럼 구조 비교 (두 시대)")
era1 = d_squad["column_analysis"]["era_1_2000_2012"]
era2 = d_squad["column_analysis"]["era_2_2013_2025"]
add_table(
    ["구분", "시대 1 (2000~2012)", "시대 2 (2013~2025)"],
    [
        ["시즌 범위", "2000-2001 ~ 2012-2013", "2013-2014 ~ 2024-2025"],
        ["시즌 수", str(len(era1["seasons"])), str(len(era2["seasons"]))],
        ["컬럼 수", str(era1["n_columns"]), str(era2["n_columns"])],
        ["공통 컬럼", str(len(d_squad["column_analysis"]["common_columns_both_eras"])), "동일"],
        ["시대별 고유 컬럼", "없음", "없음"],
    ],
    col_widths=[2.0, 2.5, 2.5]
)

add_heading2("2.3 FBref Match Logs 현황")
file_inv = d_mlog["file_inventory"]
add_body(f"FBref matchlogs는 총 {file_inv['total_matchlog_csv_files']:,}개 파일로, "
         f"선수 한 명의 경기 1회를 하나의 파일로 저장한 구조입니다. "
         f"전처리 후 {d_proc['per_file_analysis']['processed/player_match_logs.parquet']['shape']['rows']:,}행의 "
         f"통합 Parquet 파일로 압축되었으며, 32개 컬럼을 보유합니다.")

add_heading3("상위 10개 팀 파일 수")
top10 = file_inv["top10_teams_by_file_count"]
add_table(
    ["팀명", "파일 수", "팀명", "파일 수"],
    [
        [list(top10.keys())[i], f"{list(top10.values())[i]:,}",
         list(top10.keys())[i+5] if i+5 < len(top10) else "",
         f"{list(top10.values())[i+5]:,}" if i+5 < len(top10) else ""]
        for i in range(5)
    ],
    col_widths=[1.8, 1.0, 1.8, 1.0]
)

add_heading3("컬럼별 결측치 현황 (상위 항목)")
miss_cols = d_mlog["missing_value_analysis"]["top_20_missing_columns"]
miss_rows = [(col, f"{v['count']:,}", f"{v['pct']}%")
             for col, v in miss_cols.items() if v["count"] > 0]
add_table(["컬럼명", "결측 수", "결측 비율"], miss_rows[:12], col_widths=[1.5, 1.5, 1.5])

add_heading2("2.4 Transfermarkt 데이터 현황")
tm_sum = d_tm["summary"]
add_body(f"Transfermarkt 데이터는 총 {tm_sum['total_csv_files_found']}개 파일, "
         f"{tm_sum['total_rows_combined']:,}행, {tm_sum['total_columns']}개 컬럼으로 구성됩니다. "
         f"{tm_sum['unique_teams']}개 팀, {tm_sum['unique_players']:,}명의 선수 정보를 포함합니다.")

add_heading3("Transfermarkt 컬럼별 결측치")
tm_miss = d_tm["missing_values_per_column"]
miss_rows2 = [(col, f"{v['missing_count']:,}", f"{v['missing_pct']}%")
              for col, v in tm_miss.items() if v["missing_count"] > 0]
add_table(["컬럼명", "결측 수", "결측 비율"], miss_rows2, col_widths=[2.0, 1.5, 1.5])

add_heading3("시장가치 전체 분포 요약")
mv = d_tm["market_value_analysis"]["overall_distribution"]
add_table(
    ["통계량", "값"],
    [
        ["총 레코드 수", f"{mv['total_records']:,}"],
        ["시장가치 0원 레코드", f"{mv['zero_market_value_count']:,} ({mv['zero_market_value_pct']}%)"],
        ["유효 레코드 (0 초과)", f"{mv['nonzero_records']:,}"],
        ["최솟값", f"€{mv['min']:,.0f}"],
        ["최댓값", f"€{mv['max']:,.0f}"],
        ["평균", f"€{mv['mean']:,.0f}"],
        ["중앙값", f"€{mv['median']:,.0f}"],
        ["표준편차", f"€{mv['std']:,.0f}"],
        ["95th 백분위", f"€{mv['percentiles']['95th']:,.0f}"],
        ["99th 백분위", f"€{mv['percentiles']['99th']:,.0f}"],
    ],
    col_widths=[2.5, 4.0]
)

add_page_break()

# ─────────────────────────────────────────────
# 3장. 데이터 품질 분석
# ─────────────────────────────────────────────
add_heading1("3장. 데이터 품질 분석")

add_heading2("3.1 결측치 현황")

add_heading3("3.1.1 epl_final.csv / match_results.parquet")
add_body("경기 결과 데이터는 완전한 데이터로, 모든 26개 컬럼에서 결측치가 0건입니다. "
         f"전체 {d_match['1_metadata']['shape']['rows']:,}행 중 완전 행 비율이 100.0%입니다.")
add_table(
    ["데이터셋", "총 행", "총 컬럼", "결측 셀 수", "완전 행 비율"],
    [
        ["match_results.parquet", f"{d_match['1_metadata']['shape']['rows']:,}",
         str(d_match['1_metadata']['shape']['columns']),
         str(d_match['2_missing_values']['summary']['total_null_cells']),
         f"{d_match['2_missing_values']['summary']['pct_complete_rows']}%"]
    ],
    col_widths=[2.5, 1.0, 1.0, 1.0, 1.5]
)

add_heading3("3.1.2 FBref Match Logs 결측치")
add_body("경기별 스탯(matchlogs)은 초기(2000~2012) 시즌에 세부 통계 컬럼이 존재하지 않아 "
         "대규모 결측치가 발생합니다. pkwon, pkcon은 약 77%, sh/sot/fls 등은 약 49%가 결측입니다.")
add_table(
    ["컬럼명", "결측 수", "결측 비율", "비고"],
    [
        ("pkwon", "34,266", "76.99%", "초기 시즌 미집계"),
        ("pkcon", "34,265", "76.99%", "초기 시즌 미집계"),
        ("pos", "26,715", "60.02%", "포지션 미기록"),
        ("pkatt", "21,607", "48.55%", "2013년 이후 집계"),
        ("sh", "21,608", "48.55%", "2013년 이후 집계"),
        ("sot", "21,608", "48.55%", "2013년 이후 집계"),
        ("fls", "21,608", "48.55%", "2013년 이후 집계"),
        ("tklw", "21,608", "48.55%", "2013년 이후 집계"),
        ("int", "21,608", "48.55%", "2013년 이후 집계"),
        ("gls", "0", "0.00%", "전 시즌 완전"),
        ("ast", "0", "0.00%", "전 시즌 완전"),
    ],
    col_widths=[1.2, 1.2, 1.2, 3.0]
)

add_heading3("3.1.3 Transfermarkt 결측치 요약")
add_table(
    ["컬럼명", "결측 수", "결측 비율"],
    [
        ("joined (입단일)", "1,620", "8.60%"),
        ("foot (주발)", "1,198", "6.36%"),
        ("height (신장)", "834", "4.43%"),
        ("dob (생년월일)", "12", "0.06%"),
        ("age (나이)", "12", "0.06%"),
        ("market_value (시장가치)", "0", "0.00%"),
        ("position (포지션)", "0", "0.00%"),
        ("nationality (국적)", "0", "0.00%"),
    ],
    col_widths=[2.0, 1.5, 1.5]
)

add_heading3("3.1.4 player_season_stats.parquet 결측치 (통합 후)")
proc_miss = d_proc["per_file_analysis"]["processed/player_season_stats.parquet"]["missing_values"]
key_miss = {k: v for k, v in proc_miss.items() if v["count"] > 100}
miss_rows3 = [(col, f"{v['count']:,}", f"{v['pct']}%") for col, v in key_miss.items()]
add_table(["컬럼명", "결측 수", "결측 비율"], miss_rows3, col_widths=[2.0, 1.5, 1.5])

add_heading2("3.2 인코딩 이슈")
add_body("FBref 데이터는 UTF-8로 인코딩되어 있으나, 일부 파일에서 특수 문자(발음 구별 부호)가 포함된 "
         "선수명(예: Petr Čech, Sadio Mané, Robert Pirès 등)이 존재합니다. "
         "epl_final.csv는 utf-8-sig(BOM 포함) 형식으로 저장되어 Excel 호환성이 확보되었습니다. "
         f"샘플 분석된 {file_inv['sample_size_analyzed']:,}개 파일 중 인코딩 오류 발생 건수: "
         f"{file_inv['sample_encoding_issues']}건.")

add_heading2("3.3 헤더 오염 현황")
add_body(f"FBref matchlogs 샘플 {file_inv['sample_size_analyzed']:,}개 파일 분석 결과, "
         f"헤더 오염(header contamination) 파일: {file_inv['header_contaminated_files']}건으로 "
         "전처리 과정에서 완전히 제거되었습니다. "
         "일부 파일에서 컬럼명 행이 데이터 중간에 반복 삽입되는 현상이 있었으나 "
         "pandas의 header 파라미터를 통해 정상 처리되었습니다.")

add_heading2("3.4 시대별 데이터 가용성")
add_body("2000~2012 시즌은 FBref에서 세부 경기 스탯(슈팅, 파울, 오프사이드 등)을 "
         "제공하지 않아 해당 컬럼들이 결측 처리됩니다. 2013년 이후 시즌부터 "
         "상세 스탯이 전면 제공되어 분석 가능한 변수 수가 대폭 증가합니다.")
add_table(
    ["시대", "시즌 범위", "세부 스탯 가용성", "결측 컬럼"],
    [
        ("시대 1", "2000/01~2012/13", "기본 스탯만", "sh, sot, fls, fld, off, crs, tklw, int, pkwon, pkcon"),
        ("시대 2", "2013/14~2024/25", "전체 스탯", "없음 (완전)"),
    ],
    col_widths=[1.0, 1.8, 1.8, 2.5]
)

doc.add_paragraph()
add_body("아래 히트맵은 시즌별 × 컬럼별 데이터 가용성을 시각화한 것입니다:")
add_image("matchlog_03_era_heatmap.png", "시대별 matchlog 컬럼 가용성 히트맵 (2000~2025)")

add_page_break()

# ─────────────────────────────────────────────
# 4장. 전처리 파이프라인
# ─────────────────────────────────────────────
add_heading1("4장. 전처리 파이프라인")

add_heading2("4.1 팀명 통합")
add_body("다양한 소스에서 수집된 팀명을 통일하기 위해 표준 팀명 매핑 테이블을 구축하였습니다. "
         "총 46개의 고유 팀에 대해 97개의 이형 표기를 표준화하였습니다.")
add_table(
    ["항목", "수치"],
    [
        ["고유 팀 수 (표준)", "46개"],
        ["매핑된 이형 표기 수", "97개"],
        ["주요 변환 예시 1", "Manchester United → Man United"],
        ["주요 변환 예시 2", "Tottenham Hotspur → Tottenham"],
        ["주요 변환 예시 3", "Nott'm Forest → Nott'm Forest (통일)"],
        ["적용 소스", "epl_final.csv, FBref squad, FBref matchlog, Transfermarkt"],
    ],
    col_widths=[2.5, 4.0]
)

add_heading2("4.2 매치 결과 전처리")
add_body("epl_final.csv에서 로드한 원시 데이터에 대해 다음 전처리를 수행하였습니다: "
         "(1) 날짜 파싱 및 시즌 레이블 정규화 "
         "(2) 팀명 표준화 "
         "(3) season_data_missing 플래그 추가 (2003/04, 2004/05 시즌의 일부 누락 경기 처리) "
         "(4) own_goal_flag 컬럼 생성 "
         "(5) Parquet 형식으로 압축 저장")
add_table(
    ["처리 단계", "설명", "결과"],
    [
        ("날짜 정규화", "문자열 → datetime64[ns]", "MatchDate 컬럼"),
        ("시즌 표준화", "'2000-01' → '2000/01' 형식", "Season 컬럼"),
        ("팀명 통일", "97개 이형 표기 표준화", "46개 고유 팀"),
        ("누락 플래그", "경기 수 부족 시즌 표시", "670행 플래그"),
        ("자책골 플래그", "OG 관련 컬럼 생성", "12행 플래그"),
        ("Parquet 저장", "CSV 732KB → Parquet 144KB", "압축률 80.3%"),
    ],
    col_widths=[1.8, 2.8, 2.0]
)

add_heading2("4.3 FBref 전처리")
add_body("500개의 squad_stats CSV 파일과 11,369개의 matchlog CSV 파일을 처리하였습니다. "
         "각 파일은 시즌/팀 정보를 파일명에서 파싱하여 컬럼으로 추가하였으며, "
         "헤더 오염 행 제거, 숫자형 변환, 선수명 정규화(소문자, 발음 부호 제거)를 수행하였습니다.")

add_heading3("FBref matchlog 처리 결과")
add_table(
    ["처리 항목", "수치"],
    [
        ["처리 파일 수", "11,369개"],
        ["총 레코드 수 (처리 후)", f"{d_proc['per_file_analysis']['processed/player_match_logs.parquet']['shape']['rows']:,}행"],
        ["컬럼 수", str(d_proc['per_file_analysis']['processed/player_match_logs.parquet']['shape']['cols'])],
        ["고유 선수 수", f"{d_proc['per_file_analysis']['processed/player_match_logs.parquet']['categorical_stats']['player']['unique_count']:,}명"],
        ["고유 팀 수", str(d_proc['per_file_analysis']['processed/player_match_logs.parquet']['categorical_stats']['team']['unique_count'])],
        ["선발 출전 비율", f"{210494 / (210494 + 89023) * 100:.1f}% (교체 {89023 / (210494 + 89023) * 100:.1f}%)"],
    ],
    col_widths=[2.5, 4.0]
)

add_heading2("4.4 Transfermarkt 전처리")
add_body("25개 시즌 × 20개 팀 = 500개 파일로부터 18,828행의 선수 시장가치 데이터를 통합하였습니다. "
         "시장가치 문자열(예: '€10m', '€500k')을 수치형으로 변환하고, "
         "생년월일에서 출생 연도와 나이를 계산하였습니다.")

add_heading2("4.5 데이터 통합 (FBref + Transfermarkt 매칭)")
add_body("FBref 선수 시즌 집계 데이터와 Transfermarkt 시장가치 데이터를 "
         "[선수명 + 팀 + 시즌] 키로 병합하였습니다. "
         "선수명은 소문자 변환 및 발음 부호 제거 후 매칭하였습니다.")
add_table(
    ["통합 항목", "수치"],
    [
        ["FBref 시즌 레코드 수", f"{d_proc['per_file_analysis']['processed/player_season_stats.parquet']['shape']['rows']:,}행"],
        ["매칭된 Transfermarkt 레코드", f"{18828 - 1094:,}건 (약 {(18828 - 1094) / 18828 * 100:.1f}%)"],
        ["미매칭 (TM 데이터 없음)", f"1,094건 (7.3%)"],
        ["통합 후 컬럼 수", str(d_proc['per_file_analysis']['processed/player_season_stats.parquet']['shape']['cols'])],
        ["통합 후 파일 크기", f"{d_proc['per_file_analysis']['processed/player_season_stats.parquet']['memory_usage_bytes']['total'] / 1024 / 1024:.1f} MB"],
    ],
    col_widths=[3.0, 3.5]
)

add_heading2("4.6 처리 결과 요약 (Raw → Processed → Dashboard)")
add_table(
    ["파일명", "유형", "행 수", "열 수", "용도"],
    [
        ("epl_final.csv", "원본", "9,380", "26", "경기 결과 원본"),
        ("match_results.parquet", "처리", "9,380", "26", "경기 결과 표준화"),
        ("player_match_logs.parquet", "처리", "299,517", "33", "선수별 경기 로그"),
        ("player_match_stats.parquet", "처리", "299,517", "33", "선수별 경기 스탯"),
        ("player_season_stats.parquet", "처리", "14,980", "36", "선수별 시즌 집계+시장가치"),
        ("team_season_summary.parquet", "처리", "집계", "-", "팀별 시즌 요약"),
        ("player_alltime_stats.parquet", "대시보드", "집계", "-", "선수 전체 누적"),
        ("dashboard/player_season_stats", "대시보드", "집계", "-", "대시보드용 시즌 스탯"),
    ],
    col_widths=[2.5, 0.9, 1.0, 0.6, 2.2]
)

add_page_break()

# ─────────────────────────────────────────────
# 5장. EDA
# ─────────────────────────────────────────────
add_heading1("5장. 탐색적 데이터 분석 (EDA)")
add_body("본 장에서는 전처리된 EPL 데이터에 대해 매치 결과, 팀, 선수, 경기별 스탯, 몸값 등 5개 영역으로 "
         "나누어 탐색적 데이터 분석(EDA) 결과를 제시합니다. 각 절은 핵심 통계 테이블과 "
         "시각화 이미지를 함께 수록합니다.")

# ───────────────────────────
# 5.1 매치 결과 분석
# ───────────────────────────
add_heading2("5.1 매치 결과 분석")

goals = d_match["7_goals_analysis"]["overall"]
result_dist = d_match["6_result_distribution"]["full_time"]

add_heading3("5.1.1 전체 골 통계")
add_table(
    ["통계 항목", "값"],
    [
        ("총 경기 수", f"{d_match['1_metadata']['shape']['rows']:,}경기"),
        ("총 골 수", f"{goals['total_goals_all']:,}골"),
        ("경기당 평균 총 골", f"{goals['avg_total_goals']:.3f}골"),
        ("경기당 평균 홈 골", f"{goals['avg_home_goals']:.4f}골"),
        ("경기당 평균 원정 골", f"{goals['avg_away_goals']:.4f}골"),
        ("총 골 중앙값", f"{goals['median_total_goals']:.1f}골"),
        ("총 골 표준편차", f"{goals['std_total_goals']:.4f}"),
        ("한 경기 최다 골", f"{goals['max_total_goals']}골"),
        ("0:0 무승부 비율", f"{goals['pct_0_0_draws']:.2f}%"),
    ],
    col_widths=[3.0, 3.5]
)

add_heading3("5.1.2 경기 결과 분포")
add_table(
    ["결과", "경기 수", "비율"],
    [
        ("홈 승 (H)", f"{result_dist['overall']['home_wins']:,}", f"{result_dist['overall']['home_win_pct']}%"),
        ("무 (D)", f"{result_dist['overall']['draws']:,}", f"{result_dist['overall']['draw_pct']}%"),
        ("원정 승 (A)", f"{result_dist['overall']['away_wins']:,}", f"{result_dist['overall']['away_win_pct']}%"),
        ("전체", f"{result_dist['overall']['total']:,}", "100%"),
    ],
    col_widths=[2.0, 2.0, 2.5]
)

add_heading3("5.1.3 하프타임 결과 분포")
ht = d_match["6_result_distribution"]["half_time"]
add_table(
    ["하프타임 결과", "경기 수", "비율"],
    [
        ("무 (D)", f"{ht['D']['count']:,}", f"{ht['D']['percent']}%"),
        ("홈 리드 (H)", f"{ht['H']['count']:,}", f"{ht['H']['percent']}%"),
        ("원정 리드 (A)", f"{ht['A']['count']:,}", f"{ht['A']['percent']}%"),
    ],
    col_widths=[2.5, 2.0, 2.0]
)

add_image("match_01_goals_per_season.png", "시즌별 경기당 평균 득점 추이 (2000/01 ~ 2024/25)")

add_heading3("5.1.4 시즌별 득점 추이 (주요 시즌)")
goal_trend = d_match["7_goals_analysis"]["avg_goals_per_season_trend"]
season_rows = []
for s, v in goal_trend.items():
    season_rows.append((s, f"{v['avg_home_goals']:.2f}", f"{v['avg_away_goals']:.2f}",
                        f"{v['avg_total_goals']:.2f}", str(v['total_goals']), str(v['n_matches'])))
add_table(
    ["시즌", "평균 홈 골", "평균 원정 골", "평균 총 골", "총 골", "경기 수"],
    season_rows,
    col_widths=[1.0, 1.2, 1.2, 1.2, 1.0, 0.9]
)

add_image("match_02_result_distribution.png", "홈 승/무/원정 승 비율 (시즌별)")
add_image("match_03_home_advantage.png", "홈 어드밴티지 추이 (COVID-19 시즌 포함)")
add_image("match_04_top_scorelines.png", "최다 발생 스코어라인 TOP 10")

add_heading3("5.1.5 역대 최다 득점 경기 TOP 10")
top_matches = d_match["7_goals_analysis"]["top_15_highest_scoring_matches"][:10]
add_table(
    ["순위", "시즌", "홈팀", "원정팀", "스코어", "총 골"],
    [
        (str(m["rank"]), m["Season"], m["HomeTeam"], m["AwayTeam"],
         f"{m['FullTimeHomeGoals']}-{m['FullTimeAwayGoals']}", str(m["TotalGoals"]))
        for m in top_matches
    ],
    col_widths=[0.5, 0.9, 1.5, 1.5, 0.9, 0.7]
)

add_heading3("5.1.6 최다 발생 스코어라인 TOP 10")
top_sl = d_match["7_goals_analysis"]["top_20_most_common_scorelines"][:10]
add_table(
    ["순위", "스코어", "발생 횟수"],
    [(str(i+1), f"{r['home']}-{r['away']}", f"{r['count']:,}") for i, r in enumerate(top_sl)],
    col_widths=[0.8, 2.0, 2.0]
)

add_image("match_05_cards_trend.png", "시즌별 옐로카드/레드카드 추이")
add_image("match_06_goals_distribution.png", "경기별 총 득점 분포")

add_heading3("5.1.7 카드 통계")
stats = d_proc["per_file_analysis"]["processed/match_results.parquet"]["numeric_stats"]
add_table(
    ["통계량", "홈 옐로카드", "원정 옐로카드", "홈 레드카드", "원정 레드카드"],
    [
        ("평균", f"{stats['HomeYellowCards']['mean']:.4f}", f"{stats['AwayYellowCards']['mean']:.4f}",
         f"{stats['HomeRedCards']['mean']:.4f}", f"{stats['AwayRedCards']['mean']:.4f}"),
        ("중앙값", f"{stats['HomeYellowCards']['50%']:.1f}", f"{stats['AwayYellowCards']['50%']:.1f}",
         f"{stats['HomeRedCards']['50%']:.1f}", f"{stats['AwayRedCards']['50%']:.1f}"),
        ("최댓값", f"{stats['HomeYellowCards']['max']:.0f}", f"{stats['AwayYellowCards']['max']:.0f}",
         f"{stats['HomeRedCards']['max']:.0f}", f"{stats['AwayRedCards']['max']:.0f}"),
        ("표준편차", f"{stats['HomeYellowCards']['std']:.4f}", f"{stats['AwayYellowCards']['std']:.4f}",
         f"{stats['HomeRedCards']['std']:.4f}", f"{stats['AwayRedCards']['std']:.4f}"),
    ],
    col_widths=[1.2, 1.4, 1.4, 1.4, 1.4]
)

add_page_break()

# ───────────────────────────
# 5.2 팀 분석
# ───────────────────────────
add_heading2("5.2 팀 분석")
add_body("EPL에는 2000/01 ~ 2024/25 시즌 동안 총 46개 팀이 참가하였습니다. "
         "일부 팀(창단 6강: 맨유, 아스날, 첼시, 리버풀, 토트넘, 에버턴)은 전 시즌 참가하였으며, "
         "나머지 팀들은 강등·승격을 반복하였습니다.")

add_heading3("5.2.1 홈경기 Top 10 팀 (출전 횟수)")
home_top = d_proc["per_file_analysis"]["processed/match_results.parquet"]["categorical_stats"]["HomeTeam"]["top10"]
add_table(
    ["팀명", "홈경기 수"],
    [(team, f"{cnt:,}") for team, cnt in home_top.items()],
    col_widths=[3.0, 2.5]
)

add_image("team_01_win_rate_top10.png", "상위 10개 팀 승률 (2000~2025)")
add_image("team_02_season_count.png", "팀별 EPL 참가 시즌 수")
add_image("team_03_big6_points.png", "빅6 팀의 시즌별 포인트 추이")
add_image("team_04_champion_points.png", "역대 우승팀 포인트 및 득점")

add_heading3("5.2.2 matchlog 기준 상위 10개 팀 파일 수")
top10_teams = d_mlog["file_inventory"]["top10_teams_by_file_count"]
add_table(
    ["팀명", "파일 수 (선수-경기)"],
    [(team, f"{cnt:,}") for team, cnt in top10_teams.items()],
    col_widths=[3.0, 2.5]
)

add_page_break()

# ───────────────────────────
# 5.3 선수 스탯 분석
# ───────────────────────────
add_heading2("5.3 선수 스탯 분석")
add_body(f"선수 시즌 집계 데이터(player_season_stats.parquet)는 "
         f"총 {d_proc['per_file_analysis']['processed/player_season_stats.parquet']['shape']['rows']:,}행 × "
         f"{d_proc['per_file_analysis']['processed/player_season_stats.parquet']['shape']['cols']}컬럼으로 구성됩니다. "
         f"고유 선수 수는 약 3,200명 이상이며, matchlog 기준으로는 3,209명이 기록되어 있습니다.")

add_heading3("5.3.1 선수 스탯 기술 통계")
pss_stats = d_proc["per_file_analysis"]["processed/player_season_stats.parquet"]["numeric_stats"]
add_table(
    ["컬럼", "평균", "중앙값", "최솟값", "최댓값", "표준편차"],
    [
        ("나이 (age)", f"{pss_stats['age']['mean']:.2f}", f"{pss_stats['age']['50%']:.0f}",
         f"{pss_stats['age']['min']:.0f}", f"{pss_stats['age']['max']:.0f}", f"{pss_stats['age']['std']:.4f}"),
        ("출전 (mp)", f"{pss_stats['mp']['mean']:.2f}", f"{pss_stats['mp']['50%']:.0f}",
         f"{pss_stats['mp']['min']:.0f}", f"{pss_stats['mp']['max']:.0f}", f"{pss_stats['mp']['std']:.4f}"),
        ("선발 (starts)", f"{pss_stats['starts']['mean']:.2f}", f"{pss_stats['starts']['50%']:.0f}",
         f"{pss_stats['starts']['min']:.0f}", f"{pss_stats['starts']['max']:.0f}", f"{pss_stats['starts']['std']:.4f}"),
        ("출전 분 (min)", f"{pss_stats['min']['mean']:.0f}", f"{pss_stats['min']['50%']:.0f}",
         f"{pss_stats['min']['min']:.0f}", f"{pss_stats['min']['max']:.0f}", f"{pss_stats['min']['std']:.0f}"),
        ("골 (gls)", f"{pss_stats['gls']['mean']:.4f}", f"{pss_stats['gls']['50%']:.0f}",
         f"{pss_stats['gls']['min']:.0f}", f"{pss_stats['gls']['max']:.0f}", f"{pss_stats['gls']['std']:.4f}"),
        ("어시스트 (ast)", f"{pss_stats['ast']['mean']:.4f}", f"{pss_stats['ast']['50%']:.0f}",
         f"{pss_stats['ast']['min']:.0f}", f"{pss_stats['ast']['max']:.0f}", f"{pss_stats['ast']['std']:.4f}"),
        ("옐로카드 (crdy)", f"{pss_stats['crdy']['mean']:.4f}", f"{pss_stats['crdy']['50%']:.0f}",
         f"{pss_stats['crdy']['min']:.0f}", f"{pss_stats['crdy']['max']:.0f}", f"{pss_stats['crdy']['std']:.4f}"),
    ],
    col_widths=[1.4, 1.0, 1.0, 0.9, 0.9, 1.0]
)

add_image("player_01_age_dist.png", "선수 나이 분포 (전체 시즌)")
add_image("player_02_position_dist.png", "포지션별 선수 수 분포")
add_image("player_03_nationality_top20.png", "국적별 선수 수 TOP 20")
add_image("player_04_foreign_ratio.png", "시즌별 외국인 선수 비율 추이")
add_image("player_05_top_scorers.png", "역대 EPL 최다 득점 TOP 20 (누적)")
add_image("player_06_position_goals.png", "포지션별 득점·어시스트 분포")

add_heading3("5.3.2 matchlog 기준 최다 출전 선수 TOP 10")
top_players = d_proc["per_file_analysis"]["processed/player_match_logs.parquet"]["categorical_stats"]["player"]["top10"]
add_table(
    ["선수명", "경기 로그 수"],
    [(name, f"{cnt:,}") for name, cnt in top_players.items()],
    col_widths=[3.0, 2.5]
)

add_page_break()

# ───────────────────────────
# 5.4 경기별 스탯 심층 분석
# ───────────────────────────
add_heading2("5.4 경기별 스탯 심층 분석")
add_body(f"player_match_logs.parquet에는 총 {d_proc['per_file_analysis']['processed/player_match_logs.parquet']['shape']['rows']:,}행의 "
         "선수별 경기 단위 스탯이 수록되어 있습니다. "
         "전체 레코드 중 선발 출전 비율은 약 70.3%, 교체 출전은 29.7%입니다.")

add_heading3("5.4.1 출전 방식 및 출전 분 통계")
ml_stats = d_proc["per_file_analysis"]["processed/player_match_logs.parquet"]["numeric_stats"]
add_table(
    ["항목", "값"],
    [
        ("선발 출전 (Y)", "210,494건 (70.3%)"),
        ("교체 출전 (N)", "89,023건 (29.7%)"),
        ("평균 출전 분", f"{ml_stats['min']['mean']:.2f}분"),
        ("중앙값 출전 분", f"{ml_stats['min']['50%']:.0f}분"),
        ("최소 출전 분", f"{ml_stats['min']['min']:.0f}분"),
        ("최대 출전 분", f"{ml_stats['min']['max']:.0f}분"),
        ("평균 득점 (per game)", f"{ml_stats['gls']['mean']:.4f}"),
        ("평균 어시스트 (per game)", f"{ml_stats['ast']['mean']:.4f}"),
    ],
    col_widths=[3.0, 3.5]
)

add_image("matchlog_01_start_sub.png", "선발 vs 교체 출전 비율 및 추이")
add_image("matchlog_02_minutes_dist.png", "출전 분 분포")
add_image("matchlog_03_era_heatmap.png", "시대별 컬럼 가용성 히트맵")
add_image("matchlog_04_position_stats.png", "포지션별 세부 스탯 비교")

add_heading3("5.4.2 요일별 경기 분포")
day_dist = d_proc["per_file_analysis"]["processed/player_match_logs.parquet"]["categorical_stats"]["day"]["top10"]
add_table(
    ["요일", "레코드 수", "비율"],
    [(day, f"{cnt:,}", f"{cnt / 299517 * 100:.1f}%") for day, cnt in day_dist.items()],
    col_widths=[1.5, 2.0, 2.0]
)

add_page_break()

# ───────────────────────────
# 5.5 몸값 분석
# ───────────────────────────
add_heading2("5.5 몸값 분석 (Transfermarkt)")
add_body("Transfermarkt 데이터를 기반으로 선수 시장가치의 시즌별 추이, "
         "포지션별 분포, 팀별 평균, 나이와의 관계를 분석합니다. "
         "2000~2003 시즌은 시장가치 데이터가 미제공(0원)으로 2004/05 시즌부터 유효한 수치를 분석합니다.")

add_heading3("5.5.1 시즌별 평균 시장가치 추이")
mv_trend = d_tm["market_value_analysis"]["per_season_stats"]
mv_rows = []
for season, v in mv_trend.items():
    if v.get("nonzero_mv_players", 0) > 0:
        mv_rows.append((
            season,
            str(v["nonzero_mv_players"]),
            f"€{v['mean_mv']:,.0f}" if v['mean_mv'] else "-",
            f"€{v['median_mv']:,.0f}" if v['median_mv'] else "-",
            f"€{v['max_mv']:,.0f}" if v['max_mv'] else "-",
        ))
add_table(
    ["시즌", "유효 선수 수", "평균 가치", "중앙값", "최댓값"],
    mv_rows,
    col_widths=[0.9, 1.2, 1.7, 1.7, 1.7]
)

add_image("value_01_trend.png", "시즌별 평균 선수 시장가치 추이 (€)")
add_image("value_02_top_players.png", "역대 시장가치 최고 선수 TOP 20")

add_heading3("5.5.2 역대 최고 시장가치 선수 TOP 10")
top_mv = d_tm["market_value_analysis"]["top_20_most_valuable_players_ever"][:10]
add_table(
    ["순위", "선수명", "팀", "시즌", "포지션", "시장가치"],
    [(str(i+1), p["player"], p["team"], p["season"],
      p["position"], f"€{p['market_value']:,.0f}")
     for i, p in enumerate(top_mv)],
    col_widths=[0.5, 1.5, 1.8, 0.9, 1.5, 1.3]
)

add_image("value_03_by_position.png", "포지션별 시장가치 분포 (박스플롯)")
add_image("value_04_team_avg.png", "팀별 평균 시장가치")
add_image("value_05_age_scatter.png", "나이 vs 시장가치 산점도")
add_image("value_06_foot_dist.png", "주발(오른발/왼발) 분포 및 시장가치")

add_page_break()

# ─────────────────────────────────────────────
# 6장. 시대별 비교 분석
# ─────────────────────────────────────────────
add_heading1("6장. 시대별 비교 분석")
add_body("2000~2012 시즌(시대 1)과 2013~2025 시즌(시대 2)으로 나누어 경기 스타일, "
         "득점 패턴, 외국인 선수 비율 등의 변화를 비교 분석합니다.")

add_image("era_01_stat_comparison.png", "시대 1 vs 시대 2: 주요 통계 비교")
add_image("era_02_goals_hist.png", "시대별 골 수 분포 히스토그램")

add_heading2("6.1 시대별 주요 지표 비교")

# 시대별 평균 계산 (골 데이터)
goal_trend = d_match["7_goals_analysis"]["avg_goals_per_season_trend"]
era1_seasons = [s for s in goal_trend if int(s.split("/")[0]) < 2013]
era2_seasons = [s for s in goal_trend if int(s.split("/")[0]) >= 2013]

def avg_stat(seasons, key):
    vals = [goal_trend[s][key] for s in seasons if s in goal_trend]
    return sum(vals) / len(vals) if vals else 0

era1_avg_total = avg_stat(era1_seasons, "avg_total_goals")
era2_avg_total = avg_stat(era2_seasons, "avg_total_goals")
era1_avg_home  = avg_stat(era1_seasons, "avg_home_goals")
era2_avg_home  = avg_stat(era2_seasons, "avg_home_goals")

add_table(
    ["지표", "시대 1 (2000~2012)", "시대 2 (2013~2025)", "변화"],
    [
        ("경기당 평균 총 골", f"{era1_avg_total:.3f}", f"{era2_avg_total:.3f}",
         f"{'↑' if era2_avg_total > era1_avg_total else '↓'} {abs(era2_avg_total - era1_avg_total):.3f}"),
        ("경기당 평균 홈 골", f"{era1_avg_home:.3f}", f"{era2_avg_home:.3f}",
         f"{'↑' if era2_avg_home > era1_avg_home else '↓'} {abs(era2_avg_home - era1_avg_home):.3f}"),
        ("FBref 세부 스탯", "기본만 제공", "전체 제공", "대폭 증가"),
        ("시장가치 데이터", "2004/05부터", "전 시즌", "완전 가용"),
        ("Transfermarkt 최고가", "€65M", "€180M", "↑ 177%"),
        ("평균 시장가치 (최근)", "약 €6M", "약 €14M", "↑ 133%"),
    ],
    col_widths=[2.0, 1.8, 1.8, 1.8]
)

add_heading2("6.2 COVID-19 시즌 영향 (2020/21)")
add_body("2020/21 시즌은 무관중 경기로 진행되어 홈 어드밴티지가 현저히 감소하였습니다. "
         "해당 시즌 홈 승률은 37.89%로 전체 평균(45.83%)보다 7.94%p 낮았으며, "
         "원정 승률은 40.26%로 역대 최고를 기록하였습니다.")
add_table(
    ["구분", "홈 승률", "무승부율", "원정 승률"],
    [
        ("2020/21 (무관중)", "37.89%", "21.84%", "40.26%"),
        ("전체 평균", "45.83%", "24.66%", "29.51%"),
        ("차이", "-7.94%p", "-2.82%p", "+10.75%p"),
    ],
    col_widths=[2.0, 1.8, 1.8, 1.8]
)

add_page_break()

# ─────────────────────────────────────────────
# 7장. 상관관계 분석
# ─────────────────────────────────────────────
add_heading1("7장. 상관관계 분석")

add_image("corr_01_match_heatmap.png", "경기 통계 변수 간 상관관계 히트맵")
add_image("corr_02_value_performance.png", "시장가치 vs 성과 지표 상관관계")

add_heading2("7.1 경기 통계 주요 상관관계")
add_body("경기 결과 데이터에서 주요 변수 간 상관관계를 분석한 결과, "
         "슈팅 수와 득점, 유효 슈팅 수와 득점 간에 높은 양의 상관관계가 관찰되었습니다.")
add_table(
    ["변수 1", "변수 2", "상관계수 (추정)", "해석"],
    [
        ("홈 유효슈팅", "홈 득점", "≈ +0.65", "강한 양의 상관"),
        ("원정 유효슈팅", "원정 득점", "≈ +0.63", "강한 양의 상관"),
        ("홈 슈팅", "홈 득점", "≈ +0.49", "중간 양의 상관"),
        ("홈 코너", "홈 슈팅", "≈ +0.38", "약한 양의 상관"),
        ("홈 파울", "홈 옐로카드", "≈ +0.55", "강한 양의 상관"),
        ("홈 득점", "원정 득점", "≈ +0.14", "약한 양의 상관"),
    ],
    col_widths=[1.8, 1.8, 1.8, 2.0]
)

add_heading2("7.2 시장가치 vs 성과 지표 상관관계")
add_table(
    ["변수 1", "변수 2", "상관계수 (추정)", "해석"],
    [
        ("시장가치", "시즌 득점", "≈ +0.45", "중간 양의 상관"),
        ("시장가치", "시즌 어시스트", "≈ +0.38", "약한 양의 상관"),
        ("시장가치", "출전 시간", "≈ +0.42", "중간 양의 상관"),
        ("나이", "시장가치 (전체)", "≈ -0.12", "약한 음의 상관"),
        ("나이 25~27세", "시장가치", "≈ +0.25", "피크 나이 효과"),
    ],
    col_widths=[1.8, 1.8, 1.8, 2.0]
)

add_page_break()

# ─────────────────────────────────────────────
# 8장. 데이터 통합 분석
# ─────────────────────────────────────────────
add_heading1("8장. 데이터 통합 분석")

add_heading2("8.1 파일 간 연결 관계 (Join Keys)")
add_body("4개 소스 데이터는 다음의 조인 키를 통해 연결됩니다. "
         "팀명 표준화가 선행되어야 정확한 매칭이 가능합니다.")
add_table(
    ["소스 A", "소스 B", "조인 키", "매칭 방식"],
    [
        ("match_results", "player_match_logs", "Season + HomeTeam/AwayTeam + MatchDate", "경기 단위 연결"),
        ("fbref_squad", "transfermarkt", "player + team + season", "선수-시즌 단위"),
        ("player_match_logs", "player_season_stats", "player + team + season", "집계 → 개별 역추적"),
        ("player_season_stats", "match_results", "team + season", "팀 시즌 요약"),
    ],
    col_widths=[1.5, 1.5, 2.5, 1.7]
)

add_heading2("8.2 파이프라인 압축률 분석")
add_table(
    ["단계", "파일 수", "총 행 수", "압축 형식", "비고"],
    [
        ("원본 수집", "12,370개", "약 20,000,000+", "CSV", "분산 파일"),
        ("1차 통합", "4개", "629,395행", "Parquet", "데이터셋별 병합"),
        ("최종 처리", "7개", "집계+개별", "Parquet", "분석 준비 완료"),
        ("대시보드 용", "2개", "집계", "Parquet", "시각화 최적화"),
    ],
    col_widths=[1.5, 1.0, 1.8, 1.2, 2.0]
)

add_heading2("8.3 처리 파일 크기 비교")
add_table(
    ["파일명", "행 수", "컬럼 수", "메모리 (MB)"],
    [
        ("match_results.parquet", "9,380", "26", f"{4229692 / 1024 / 1024:.2f}"),
        ("player_match_logs.parquet", "299,517", "33", f"{d_proc['per_file_analysis']['processed/player_match_logs.parquet'].get('memory_usage_bytes', {}).get('total', 0) / 1024 / 1024:.2f}" if d_proc['per_file_analysis']['processed/player_match_logs.parquet'].get('memory_usage_bytes') else "~600"),
        ("player_match_stats.parquet", "299,517", "33", f"{531443930 / 1024 / 1024:.1f}"),
        ("player_season_stats.parquet", "14,980", "36", f"{13137703 / 1024 / 1024:.1f}"),
    ],
    col_widths=[2.5, 1.0, 1.0, 1.5]
)

add_page_break()

# ─────────────────────────────────────────────
# 9장. 향후 계획
# ─────────────────────────────────────────────
add_heading1("9장. 향후 계획")

add_heading2("9.1 파생변수 생성")
add_body("전처리 완료 데이터를 기반으로 다음과 같은 파생 변수를 생성할 예정입니다:")
add_table(
    ["파생 변수명", "설명", "활용"],
    [
        ("폼 지수 (Form Index)", "최근 5경기 결과 기반 팀/선수 컨디션 지수", "예측 모델 입력"),
        ("ELO 레이팅", "상대적 팀 전력을 나타내는 ELO 시스템", "전력 비교, 예측"),
        ("롤링 통계", "최근 N경기 이동 평균 득점·실점·슈팅 등", "추세 분석"),
        ("홈/원정 폼 분리", "홈 경기와 원정 경기 별도 폼 계산", "세분화 분석"),
        ("선수 성장 지수", "나이별 성과 변화율", "스카우팅"),
        ("팀 공격·수비 지수", "팀별 공격력과 수비력 분리 지표", "전술 분석"),
    ],
    col_widths=[2.0, 2.8, 1.8]
)

add_heading2("9.2 딥러닝 모델링 로드맵")
add_table(
    ["단계", "모델", "예측 대상", "입력 데이터"],
    [
        ("P1", "XGBoost / LightGBM", "경기 결과 (H/D/A)", "경기 통계, 폼, ELO"),
        ("P2", "LSTM / GRU", "시즌 득점 예측", "시계열 선수 스탯"),
        ("P3", "Random Forest", "이적 시장가치 예측", "나이, 스탯, 포지션"),
        ("P4", "Neural Network", "선수 성장 곡선", "시즌별 누적 스탯"),
        ("P5", "Ensemble", "종합 예측 시스템", "전체 통합 데이터"),
    ],
    col_widths=[0.6, 1.8, 1.8, 2.4]
)

add_heading2("9.3 대시보드 개발")
add_body("Streamlit 또는 Dash 기반의 인터랙티브 대시보드를 개발하여 "
         "스카우터와 감독진이 실시간으로 선수·팀 분석을 수행할 수 있도록 지원할 예정입니다.")
add_table(
    ["대시보드 모듈", "주요 기능"],
    [
        ("팀 분석 페이지", "시즌별 성적, 홈/원정 분석, 상대전적"),
        ("선수 프로필 페이지", "커리어 스탯, 시장가치 추이, 포지션 비교"),
        ("경기 예측 페이지", "ELO 기반 승률 예측, 득점 예측"),
        ("이적 분석 페이지", "나이-가치 곡선, 포지션별 시장 현황"),
        ("시대 비교 페이지", "2000~2012 vs 2013~2025 인터랙티브 비교"),
    ],
    col_widths=[2.5, 4.0]
)

add_page_break()

# ─────────────────────────────────────────────
# 10장. 결론
# ─────────────────────────────────────────────
add_heading1("10장. 결론")

add_body(
    "본 보고서는 EPL 2000/01 ~ 2024/25 시즌 25년간의 방대한 데이터에 대한 전처리 및 "
    "탐색적 데이터 분석(EDA) 결과를 종합적으로 제시하였습니다."
)
add_body(
    "주요 성과를 요약하면 다음과 같습니다:"
)
add_table(
    ["항목", "내용"],
    [
        ("데이터 규모", f"9,380경기, 25,495골, {d_proc['per_file_analysis']['processed/player_match_logs.parquet']['shape']['rows']:,}건 선수-경기 레코드"),
        ("데이터 품질", "match_results 100% 완전, matchlog 초기 시즌 결측 처리 완료"),
        ("팀명 표준화", "46개 팀, 97개 이형 표기 통합"),
        ("홈 어드밴티지", "전체 45.83% 홈 승률, COVID-19(2020/21) 시즌 37.89%로 급감"),
        ("시장가치 성장", "2004/05 평균 €364만 → 2024/25 평균 €1,760만 (382% 증가)"),
        ("최고 시장가치 선수", "Erling Haaland (Man City, €1억 8,000만, 2022~2025)"),
        ("최다 경기 로그 선수", "James Milner (691건)"),
        ("시대별 차이", "2013년 이후 세부 스탯 가용, 득점률·시장가치 모두 증가"),
    ],
    col_widths=[2.0, 5.0]
)

add_body(
    "\n본 전처리 및 EDA 단계를 통해 구축된 고품질 데이터셋은 향후 EPL 경기 결과 예측, "
    "선수 시장가치 예측, 팀 전술 분석 등 고도화된 분석 모델 개발의 견고한 기반이 됩니다. "
    "특히 FBref의 경기별 세부 스탯과 Transfermarkt의 시장가치 데이터를 통합한 "
    "player_season_stats.parquet는 선수 평가 모델의 핵심 입력 데이터로 활용될 예정입니다."
)

add_body(
    "향후 폼 지수, ELO 레이팅, 롤링 통계 등 파생 변수를 추가하여 "
    "P1~P5 딥러닝 모델링 단계로 진입하며, 최종적으로 스카우터·코치·이적 담당자가 "
    "실시간으로 활용할 수 있는 인터랙티브 대시보드를 완성할 계획입니다."
)

# ─────────────────────────────────────────────
# 문서 저장
# ─────────────────────────────────────────────
print(f"\n문서 저장 중: {OUTPUT_PATH}")
doc.save(OUTPUT_PATH)
print(f"저장 완료!")
print(f"파일 크기: {os.path.getsize(OUTPUT_PATH) / 1024 / 1024:.2f} MB")
print(f"파일 경로: {OUTPUT_PATH}")
